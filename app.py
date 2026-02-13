import os
import json
import uuid
import re
import zlib
import datetime
import html
import unicodedata
import threading
from typing import Dict, Any, Optional, List, Iterable, Tuple
from functools import wraps
from flask import session
from PIL import Image
import tempfile
from docx.shared import Inches

import requests
from flask import Flask, request, redirect, url_for, jsonify, render_template, abort, send_file

import zipfile
from io import BytesIO
from docx import Document


app = Flask(__name__)

# =========================
# Auth (admin)
# =========================
app.secret_key = os.environ.get("SECRET_KEY", "dev-secret-change-me")

ADMIN_USER = os.environ.get("ADMIN_USER", "")
ADMIN_PASSWORD = os.environ.get("ADMIN_PASSWORD", "")
SECRETARY_USER = os.environ.get("SECRETARY_USER", "")
SECRETARY_PASSWORD = os.environ.get("SECRETARY_PASSWORD", "")
SESSION_DAYS = int(os.environ.get("SESSION_DAYS", "30"))

app.config.update(
    SESSION_COOKIE_NAME="integrale_admin",
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE="Lax",
    SESSION_COOKIE_SECURE=True,  # Render = https
    PERMANENT_SESSION_LIFETIME=datetime.timedelta(days=SESSION_DAYS),
)

def admin_login_required(view):
    @wraps(view)
    def wrapped(*args, **kwargs):
        if not session.get("admin_logged_in"):
            return redirect(url_for("admin_login", next=request.path))
        return view(*args, **kwargs)
    return wrapped

def admin_write_required(view):
    @wraps(view)
    def wrapped(*args, **kwargs):
        if session.get("admin_role") == "viewer":
            abort(403)
        return view(*args, **kwargs)
    return wrapped

@app.context_processor
def inject_read_only():
    admin_notifications = {"notifications": [], "unresolved_total": 0}
    if session.get("admin_logged_in") and _admin_can_view_notifications():
        try:
            admin_notifications = _admin_notifications_payload(load_data())
        except Exception:
            admin_notifications = {"notifications": [], "unresolved_total": 0}
    return {
        "is_admin_logged_in": bool(session.get("admin_logged_in")),
        "is_read_only": session.get("admin_role") == "viewer",
        "admin_notifications": admin_notifications["notifications"],
        "admin_unresolved_total": admin_notifications["unresolved_total"],
        "admin_can_access_notifications": _admin_can_view_notifications(),
        "admin_can_manage_notifications": _admin_can_manage_notifications(),
    }

@app.get("/admin/login")
def admin_login():
    # mini page sans template (pour aller vite)
    next_url = request.args.get("next") or url_for("admin_sessions")
    return f"""
    <!doctype html><html lang="fr"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
    <title>Connexion admin</title></head>
    <body style="font-family:Arial,sans-serif;max-width:420px;margin:60px auto;padding:20px">
      <h2>Connexion</h2>
      <p style="color:#6b7280;font-size:13px;margin-top:-4px">
        Un compte de consultation peut être configuré pour un accès en lecture seule.
      </p>
      <form method="post" action="/admin/login">
        <input type="hidden" name="next" value="{next_url}">
        <div style="margin:10px 0">
          <label>Identifiant</label><br>
          <input name="username" autocomplete="username" style="width:100%;padding:10px">
        </div>
        <div style="margin:10px 0">
          <label>Mot de passe</label><br>
          <input name="password" type="password" autocomplete="current-password" style="width:100%;padding:10px">
        </div>
        <button style="padding:10px 14px">Se connecter</button>
      </form>
    </body></html>
    """

@app.post("/admin/login")
def admin_login_post():
    username = (request.form.get("username") or "").strip()
    password = (request.form.get("password") or "").strip()
    next_url = request.form.get("next") or url_for("admin_sessions")

    # sécurité minimale : si pas configuré, on refuse
    if not (ADMIN_USER and ADMIN_PASSWORD) and not (SECRETARY_USER and SECRETARY_PASSWORD):
        abort(500, "ADMIN_USER/ADMIN_PASSWORD non configurés")

    if username == ADMIN_USER and password == ADMIN_PASSWORD:
        session["admin_logged_in"] = True
        session["admin_role"] = "admin"
        session.permanent = True  # ✅ cookie persistant
        return redirect(next_url)

    if SECRETARY_USER and SECRETARY_PASSWORD:
        if username == SECRETARY_USER and password == SECRETARY_PASSWORD:
            session["admin_logged_in"] = True
            session["admin_role"] = "viewer"
            session.permanent = True
            return redirect(next_url)

    return redirect(url_for("admin_login", next=next_url))

@app.get("/admin/logout")
def admin_logout():
    session.clear()
    return redirect(url_for("admin_login"))

def fr_date(value: str) -> str:
    s = (value or "").strip()
    if not s:
        return ""
    try:
        dt = datetime.datetime.strptime(s[:10], "%Y-%m-%d")
        return dt.strftime("%d/%m/%Y")
    except Exception:
        return value

def fr_datetime(value: str) -> str:
    s = (value or "").strip()
    if not s:
        return ""
    normalized = s.replace("Z", "+00:00")
    try:
        dt = datetime.datetime.fromisoformat(normalized)
        return dt.strftime("%d/%m/%Y à %Hh%M")
    except Exception:
        pass
    for fmt in ("%Y-%m-%d %H:%M", "%Y-%m-%dT%H:%M", "%Y-%m-%dT%H:%M:%S", "%Y-%m-%dT%H:%M:%S.%f"):
        try:
            dt = datetime.datetime.strptime(s[:26], fmt)
            return dt.strftime("%d/%m/%Y à %Hh%M")
        except Exception:
            pass
    return fr_date(s)


def history_datetime(value: str) -> str:
    s = (value or "").strip()
    if not s:
        return ""
    normalized = s.replace("Z", "+00:00")
    try:
        dt = datetime.datetime.fromisoformat(normalized)
        return dt.strftime("%d/%m/%Y %H:%M")
    except Exception:
        pass
    for fmt in ("%Y-%m-%d %H:%M", "%Y-%m-%dT%H:%M", "%Y-%m-%dT%H:%M:%S", "%Y-%m-%dT%H:%M:%S.%f"):
        try:
            dt = datetime.datetime.strptime(s[:26], fmt)
            return dt.strftime("%d/%m/%Y %H:%M")
        except Exception:
            pass
    return fr_date(s)


# ✅ Filtres utilisables dans tous tes templates
app.add_template_filter(fr_date, "frdate")
app.add_template_filter(fr_datetime, "frdatetime")


# =========================
# Persistent disk (Render)
# =========================
PERSIST_DIR = os.environ.get("PERSIST_DIR", "/data")
os.makedirs(PERSIST_DIR, exist_ok=True)
DATA_FILE = os.path.join(PERSIST_DIR, "data.json")

BACKUP_DIR = os.path.join(PERSIST_DIR, "backups")
os.makedirs(BACKUP_DIR, exist_ok=True)
BACKUP_RETENTION = int(os.environ.get("BACKUP_RETENTION", "120"))
BACKUP_MIN_INTERVAL_SECONDS = int(os.environ.get("BACKUP_MIN_INTERVAL_SECONDS", "300"))

_data_lock = threading.RLock()
_last_backup_times: Dict[str, float] = {}

UPLOADS_DIR = os.path.join(PERSIST_DIR, "uploads")
os.makedirs(UPLOADS_DIR, exist_ok=True)

def trainee_upload_dir(session_id: str, trainee_id: str) -> str:
    d = os.path.join(UPLOADS_DIR, session_id, trainee_id)
    os.makedirs(d, exist_ok=True)
    return d


def _cleanup_backups_for(prefix: str) -> None:
    try:
        names = sorted(
            [name for name in os.listdir(BACKUP_DIR) if name.startswith(prefix + ".")],
            reverse=True,
        )
        for old_name in names[BACKUP_RETENTION:]:
            old_path = os.path.join(BACKUP_DIR, old_name)
            if os.path.isfile(old_path):
                os.remove(old_path)
    except Exception:
        pass


def _force_backup_snapshot(path: str) -> None:
    base_name = os.path.basename(path)
    prefix = base_name.replace(".", "_")
    if not os.path.exists(path):
        return
    stamp = datetime.datetime.utcnow().strftime("%Y%m%dT%H%M%SZ")
    backup_path = os.path.join(BACKUP_DIR, f"{prefix}.manual.{stamp}.json")
    try:
        with open(path, "rb") as src, open(backup_path, "wb") as dst:
            dst.write(src.read())
        _cleanup_backups_for(prefix)
    except Exception:
        pass


def _write_json_with_backups(path: str, payload: Dict[str, Any], lock: threading.RLock) -> None:
    with lock:
        now_ts = datetime.datetime.utcnow().timestamp()
        base_name = os.path.basename(path)
        prefix = base_name.replace(".", "_")

        if os.path.exists(path):
            last = _last_backup_times.get(path, 0)
            if now_ts - last >= BACKUP_MIN_INTERVAL_SECONDS:
                stamp = datetime.datetime.utcnow().strftime("%Y%m%dT%H%M%SZ")
                backup_path = os.path.join(BACKUP_DIR, f"{prefix}.{stamp}.json")
                try:
                    with open(path, "rb") as src, open(backup_path, "wb") as dst:
                        dst.write(src.read())
                    _last_backup_times[path] = now_ts
                    _cleanup_backups_for(prefix)
                except Exception:
                    pass

        tmp = path + ".tmp"
        with open(tmp, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False, indent=2)
        os.replace(tmp, path)


# =========================
# Brevo (Sendinblue) config
# =========================
BREVO_API_KEY = os.environ.get("BREVO_API_KEY", "")
BREVO_SENDER_EMAIL = os.environ.get("BREVO_SENDER_EMAIL", "ecole@integraleacademy.com")
BREVO_SENDER_NAME = os.environ.get("BREVO_SENDER_NAME", "Intégrale Academy")
CNAPS_LOOKUP_ENDPOINT = os.environ.get("CNAPS_LOOKUP_ENDPOINT", "")

PUBLIC_STUDENT_PORTAL_BASE = os.environ.get(
    "PUBLIC_STUDENT_PORTAL_BASE",
    "https://gestionstagiaires-r5no.onrender.com"
)

PUBLIC_BASE_URL = os.environ.get(
    "PUBLIC_BASE_URL",
    "https://gestionstagiaires-r5no.onrender.com"
)


CNAPS_STATUS_ENDPOINT = os.environ.get("CNAPS_STATUS_ENDPOINT", "")
HEBERGEMENT_STATUS_ENDPOINT = os.environ.get("HEBERGEMENT_STATUS_ENDPOINT", "")


def normalize_phone_fr(phone: str) -> str:
    p = (phone or "").strip().replace(" ", "").replace(".", "").replace("-", "")
    if not p:
        return ""
    if p.startswith("+"):
        return p
    if p.startswith("00"):
        return "+" + p[2:]
    if p.startswith("0") and len(p) == 10 and p[1:].isdigit():
        return "+33" + p[1:]
    return p

def _collapse_spaces(value: str) -> str:
    return " ".join((value or "").strip().split())


def _parse_no_answer_count(value: Any) -> int:
    try:
        count = int(value)
    except (TypeError, ValueError):
        return 0
    return max(0, min(3, count))


def normalize_last_name(value: str) -> str:
    collapsed = _collapse_spaces(value)
    return collapsed.upper()

def normalize_first_name(value: str) -> str:
    collapsed = _collapse_spaces(value)
    lowered = collapsed.lower()
    return lowered.title()



import base64

def brevo_send_email(
    to_email: str,
    subject: str,
    html: str,
    cc_emails: Optional[List[str]] = None,
) -> bool:
    if not BREVO_API_KEY or not to_email:
        return False

    url = "https://api.brevo.com/v3/smtp/email"
    headers = {
        "accept": "application/json",
        "api-key": BREVO_API_KEY,
        "content-type": "application/json",
    }

    attachments = []  # ✅ pas d'inline CID, Gmail casse souvent

    payload = {
        "sender": {"name": BREVO_SENDER_NAME, "email": BREVO_SENDER_EMAIL},
        "to": [{"email": to_email}],
        "subject": subject,
        "htmlContent": html,
    }

    cc_list = [email for email in (cc_emails or []) if email]
    if cc_list:
        payload["cc"] = [{"email": email} for email in cc_list]

    if attachments:
        payload["attachment"] = attachments

    try:
        r = requests.post(url, headers=headers, json=payload, timeout=12)
        print("[EMAIL] status=", r.status_code)
        print("[EMAIL] response=", r.text)
        return r.status_code in (200, 201, 202)
    except Exception:
        return False


def brevo_send_sms(phone: str, message: str) -> bool:
    phone = normalize_phone_fr(phone)
    if not BREVO_API_KEY or not phone:
        print("[SMS] Missing BREVO_API_KEY or phone")
        return False

    url = "https://api.brevo.com/v3/transactionalSMS/sms"
    headers = {
        "accept": "application/json",
        "api-key": BREVO_API_KEY,
        "content-type": "application/json",
    }

    # (souvent requis selon config Brevo) : nom d’expéditeur SMS
    sms_sender = os.environ.get("BREVO_SMS_SENDER", "").strip()

    payload = {
        "recipient": phone,
        "content": message,
        "type": "transactional",
        "unicodeEnabled": True, 
    }
    if sms_sender:
        payload["sender"] = sms_sender  # ex: "INTEGRALE"

    try:
        r = requests.post(url, headers=headers, json=payload, timeout=12)

        # ✅ logs indispensables (status + réponse Brevo)
        print("[SMS] status=", r.status_code)
        print("[SMS] response=", r.text)

        return r.status_code in (200, 201, 202)
    except Exception as e:
        print("[SMS] exception=", repr(e))
        return False


def notify_elearning_access_available(trainee: Dict[str, Any], session_obj: Dict[str, Any], link: str) -> Dict[str, bool]:
    first_name = (trainee.get("first_name") or "").strip() or "Madame, Monsieur"
    training_name = formation_label(_session_get(session_obj, "training_type", ""))
    date_start = fr_date(_session_get(session_obj, "date_start", ""))
    date_end = fr_date(_session_get(session_obj, "date_end", ""))
    student_space_link = f"{PUBLIC_STUDENT_PORTAL_BASE.rstrip('/')}/espace/{(trainee.get('public_token') or '').strip()}"
    access_link = student_space_link if (trainee.get("public_token") or "").strip() else link

    subject = "Votre accès e-learning est disponible – Intégrale Academy"
    html = mail_layout(f"""
      <h2 style="text-align:center">🚀 Accès e-learning activé</h2>
      <p>Bonjour <strong>{first_name}</strong>,</p>
      <p>
        Bonne nouvelle : votre accès à la <strong>Formation théorique en e-learning</strong>
        est maintenant disponible.
      </p>
      <div style="background:#eff6ff;border:1px solid #bfdbfe;border-radius:12px;padding:14px;margin:16px 0;">
        <p style="margin:0 0 10px 0;">
          <strong>📌 Formation :</strong> {training_name}
          {" — <strong>Dates :</strong> " + date_start + " au " + date_end if (date_start or date_end) else ""}
        </p>
        <p style="margin:0;">
          <strong>🔗 Accéder à votre Espace Stagiaire :</strong><br>
          <a href="{access_link}" style="display:inline-block;margin-top:8px;background:#2563eb;color:#ffffff;text-decoration:none;font-weight:700;padding:10px 16px;border-radius:8px;">
            Accéder à mon Espace Stagiaire
          </a>
        </p>
      </div>
      <p>
        Vous pouvez dès maintenant retrouver vos accès à la formation théorique e-learning
        directement dans votre Espace Stagiaire.
      </p>
    """)

    sms_name = (trainee.get("first_name") or "").strip()
    sms = (
        f"Intégrale Academy ✅ {sms_name + ', ' if sms_name else ''}"
        "Votre accès e-learning Formation VTC est disponible. "
        f"Connectez vous à votre Espace Stagiaire pour suivre votre formation : {access_link}"
    )

    email_ok = brevo_send_email((trainee.get("email") or "").strip(), subject, html)
    sms_ok = brevo_send_sms((trainee.get("phone") or "").strip(), sms)
    return {"email_ok": bool(email_ok), "sms_ok": bool(sms_ok)}




def build_vtc_practice_convocation_email(first_name: str, practice_training_date: str) -> Tuple[str, str]:
    trainee_first_name = (first_name or "").strip() or "Prénom"
    practice_date_fr = fr_date(practice_training_date) or "DATE FORMATION PRATIQUE"

    subject = "Convocation formation pratique"
    html = mail_layout(f"""
      <div style="background:linear-gradient(135deg,#eff6ff,#f0fdf4);border:1px solid #dbeafe;border-radius:14px;padding:18px;">
        <h2 style="margin:0 0 12px 0;color:#0f172a;">Convocation formation pratique</h2>
        <p style="margin:0 0 10px 0;">Bonjour {trainee_first_name},</p>

        <p style="margin:0 0 10px 0;">Je reviens vers vous concernant votre parcours Chauffeur VTC.</p>

        <p style="margin:0 0 10px 0;">Tout d’abord, félicitations pour votre réussite à l’examen théorique 👏 C’est une étape importante vers l’obtention de votre carte professionnelle !</p>

        <p style="margin:0 0 10px 0;">Vous avez normalement reçu un message de la Chambre de Métiers et de l'Artisanat vous demandant de préciser le centre de formation ainsi que l’établissement mettant à disposition le véhicule à doubles commandes pour l’épreuve pratique. Merci d’indiquer : Intégrale Sécurité Formations.</p>

        <div style="background:#ffffff;border:1px solid #bbf7d0;border-radius:12px;padding:12px 14px;margin:12px 0;">
          <p style="margin:0;">Votre formation pratique est prévue le {practice_date_fr}, de 08h30 à 12h00, dans nos locaux :<br>
          Intégrale Academy<br>
          54 chemin du Carreou<br>
          83480 PUGET-SUR-ARGENS</p>
        </div>

        <p style="margin:0 0 8px 0;">Au cours de cette matinée, nous vous préparerons concrètement à l’examen pratique :</p>
        <ul style="margin:0 0 12px 18px;padding:0;">
          <li>déroulement détaillé de l’épreuve,</li>
          <li>mise en situation professionnelle,</li>
          <li>examen blanc,</li>
          <li>prise en main du véhicule à doubles commandes,</li>
          <li>conseils méthodologiques pour optimiser votre passage devant le jury.</li>
        </ul>

        <p style="margin:0 0 10px 0;">Vous recevrez prochainement par mail votre convocation à la formation pratique, ainsi que le document officiel de prêt du véhicule à doubles commandes.<br>
        ⚠️ Il est impératif de présenter ce document le jour de l’examen : en son absence, le jury peut prononcer un ajournement.</p>

        <p style="margin:0 0 10px 0;">Nous restons à votre disposition si vous avez la moindre question.<br>
        À très bientôt pour la préparation finale ! 🚗</p>
      </div>
    """)
    return subject, html


def build_vtc_practice_convocation_sms(first_name: str, practice_training_date: str) -> str:
    trainee_first_name = (first_name or "").strip()
    practice_date_fr = fr_date(practice_training_date) or "DATE FORMATION PRATIQUE"
    greeting = f"Bonjour {trainee_first_name}, " if trainee_first_name else "Bonjour, "
    return (
        "Intégrale Academy 🚗 "
        f"{greeting}félicitations pour votre réussite à l'examen théorique. "
        f"Votre formation pratique VTC est prévue le {practice_date_fr} de 08h30 à 12h00 "
        "à Intégrale Academy, 54 chemin du Carreou 83480 Puget-sur-Argens."
    )


def _normalize_cmar_identifier(value: str) -> str:
    raw = (value or "").strip().upper()
    if not raw:
        return ""
    return "".join(ch for ch in raw if ch.isalnum())


def _extract_cmar_identifiers_from_pdf(file_bytes: bytes) -> List[str]:
    if not file_bytes:
        return []

    def _decode_pdf_literal_strings(blob: bytes) -> List[str]:
        out: List[str] = []
        i = 0
        n = len(blob)
        while i < n:
            if blob[i] != 0x28:  # (
                i += 1
                continue

            i += 1
            depth = 1
            buf = bytearray()
            while i < n and depth > 0:
                ch = blob[i]

                if ch == 0x5C:  # backslash
                    i += 1
                    if i >= n:
                        break
                    esc = blob[i]
                    simple = {
                        0x6E: 0x0A,  # \n
                        0x72: 0x0D,  # \r
                        0x74: 0x09,  # \t
                        0x62: 0x08,  # \b
                        0x66: 0x0C,  # \f
                        0x28: 0x28,  # \(
                        0x29: 0x29,  # \)
                        0x5C: 0x5C,  # \\
                    }
                    if esc in simple:
                        buf.append(simple[esc])
                        i += 1
                        continue

                    if 0x30 <= esc <= 0x37:
                        oct_digits = bytes([esc])
                        i += 1
                        for _ in range(2):
                            if i < n and 0x30 <= blob[i] <= 0x37:
                                oct_digits += bytes([blob[i]])
                                i += 1
                            else:
                                break
                        buf.append(int(oct_digits, 8) & 0xFF)
                        continue

                    buf.append(esc)
                    i += 1
                    continue

                if ch == 0x28:  # (
                    depth += 1
                    buf.append(ch)
                    i += 1
                    continue

                if ch == 0x29:  # )
                    depth -= 1
                    if depth > 0:
                        buf.append(ch)
                    i += 1
                    continue

                buf.append(ch)
                i += 1

            if buf:
                out.append(buf.decode("latin-1", errors="ignore"))

        return out

    def _decode_pdf_hex_strings(blob: bytes) -> List[str]:
        out: List[str] = []
        for m in re.finditer(rb"<([0-9A-Fa-f\s]{4,})>", blob):
            raw = re.sub(rb"\s+", b"", m.group(1))
            if len(raw) % 2 == 1:
                raw += b"0"
            try:
                out.append(bytes.fromhex(raw.decode("ascii", errors="ignore")).decode("latin-1", errors="ignore"))
            except Exception:
                continue
        return out

    chunks: List[bytes] = [file_bytes]
    for m in re.finditer(rb"stream\r?\n(.*?)\r?\nendstream", file_bytes, flags=re.S):
        stream_data = m.group(1)
        chunks.append(stream_data)

        if b"/FlateDecode" in file_bytes[max(0, m.start() - 250):m.start()]:
            try:
                inflated = zlib.decompress(stream_data)
                chunks.append(inflated)
            except Exception:
                pass

    text_sources: List[str] = []
    for chunk in chunks:
        text_sources.append(chunk.decode("latin-1", errors="ignore"))
        text_sources.extend(_decode_pdf_literal_strings(chunk))
        text_sources.extend(_decode_pdf_hex_strings(chunk))

    content = "\n".join(text_sources).upper()
    candidates = set()

    for token in re.findall(r"\b(?:CMAR\s*[:\-]?)?([A-Z0-9\-]{4,})\b", content):
        normalized = _normalize_cmar_identifier(token)
        if not normalized:
            continue
        if normalized.startswith("CMAR"):
            normalized = normalized[4:]
        if normalized.isdigit() and len(normalized) < 6:
            continue
        if any(ch.isdigit() for ch in normalized) and len(normalized) >= 6:
            candidates.add(normalized)

    return sorted(candidates)


def _send_vtc_theory_exam_notification(session_obj: Dict[str, Any], trainee: Dict[str, Any], send_email: bool = True) -> Dict[str, Any]:
    practice_training_date = (
        _session_get(session_obj, "practice_training_date", "")
        or _session_get(session_obj, "exam_practice_date", "")
        or _session_get(session_obj, "exam_date", "")
    )

    first_name = (trainee.get("first_name") or "").strip()
    email = (trainee.get("email") or "").strip()
    phone = (trainee.get("phone") or "").strip()

    subject, html = build_vtc_practice_convocation_email(first_name, practice_training_date)
    sms = build_vtc_practice_convocation_sms(first_name, practice_training_date)

    email_ok = brevo_send_email(email, subject, html) if (send_email and email) else False
    sms_ok = brevo_send_sms(phone, sms) if phone else False

    trainee["vtc_theory_exam_sent_at"] = _now_iso()
    trainee["vtc_theory_exam_email_ok"] = bool(email_ok)
    trainee["vtc_theory_exam_sms_ok"] = bool(sms_ok)
    trainee["updated_at"] = _now_iso()

    return {
        "email_ok": bool(email_ok),
        "sms_ok": bool(sms_ok),
        "sent_at": trainee.get("vtc_theory_exam_sent_at") or "",
    }


def mail_layout(inner_html: str) -> str:
    # ✅ logo en URL HTTPS (fiable dans Gmail)
    logo_src = f"{PUBLIC_BASE_URL.rstrip('/')}/static/logo-integrale.png"

    return f"""
    <div style="font-family:Arial,sans-serif;max-width:640px;margin:auto;background:#f7f7f7;padding:18px;border-radius:12px">
      <div style="background:white;padding:18px;border-radius:12px">
        <div style="text-align:center;margin-bottom:18px">
          <img src="{logo_src}" alt="Intégrale Academy"
               style="height:60px;width:auto;display:block;margin:0 auto;border:0;outline:none;text-decoration:none">
        </div>

        {inner_html}

        <p style="margin-top:30px;color:#666;font-size:13px;text-align:center">
          Intégrale Academy
        </p>
      </div>
    </div>
    """


def build_vtc_onboarding_email(first_name: str, form_link: str) -> Tuple[str, str]:
    first_name = (first_name or "").strip()
    greeting = f"Bonjour <strong>{first_name}</strong>," if first_name else "Bonjour,"
    subject = "Votre inscription Chauffeur VTC – Intégrale Academy"

    html = mail_layout(f"""
      <p>{greeting}</p>
      <p>
        Je fais suite à votre inscription en formation <strong>Chauffeur VTC</strong>.
      </p>
      <p>Je vous remercie pour votre confiance !</p>
      <p>
        Vous pouvez à présent accéder à votre Espace Stagiaire en cliquant ici :
      </p>
      <p style="text-align:center;margin:18px 0;">
        <a href="{form_link}"
           style="display:inline-block;background:#1f8f4a;color:white;padding:12px 18px;border-radius:10px;text-decoration:none;font-weight:bold">
          👉 Accéder à mon Espace Stagiaire
        </a>
      </p>

      <p>Dans votre Espace Stagiaire vous allez retrouver :</p>

      <p>
        1️⃣ Les indications pour créer votre compte Chambre des métiers (exament3p) :
        ce compte vous permettra de déposer vos documents officiels nécessaires pour
        l'inscription à l'examen théorique et l'examen pratique
      </p>
      <p>
        2️⃣ Dès que votre compte Chambre des métiers sera créé, vous devrez nous indiquer,
        dans votre Espace Stagiaire, votre identifiant et votre mot de passe Chambre des métiers,
        afin que nous puissions nous connecter et procéder au paiement des frais d'examen
        (⚠️ Veillez à ne pas régler les frais d'inscriptions, ils sont inclus dans votre formation)
      </p>
      <p>
        3️⃣ Les accès à votre formation théorique en e-learning
      </p>

      <p>
        Nous restons à votre disposition pour tout renseignement complémentaire.<br>
        Excellente journée à vous.
      </p>

      <p style="margin-top:18px;">
        Bien cordialement,<br>
        <strong>Clément VAILLANT</strong><br>
        Directeur Général – Intégrale Sécurité Formations<br>
        04 22 47 07 68<br>
        <a href="https://www.integraleacademy.com" target="_blank" rel="noopener">www.integraleacademy.com</a>
      </p>
    """)
    return subject, html


def build_vtc_onboarding_sms(first_name: str, form_link: str) -> str:
    first_name = (first_name or "").strip()
    greeting = f"Bonjour {first_name}, " if first_name else "Bonjour, "
    return (
        "Intégrale Academy 🚖 "
        f"{greeting}"
        "votre inscription en formation Chauffeur VTC est confirmée. "
        f"Accédez à votre Espace Stagiaire : {form_link} "
        "Vous y retrouverez les étapes Chambre des métiers et vos accès e-learning. "
        "Besoin d'aide ? 04 22 47 07 68."
    )


def _parse_iso_datetime(value: str) -> Optional[datetime.datetime]:
    raw = (value or "").strip()
    if not raw:
        return None
    normalized = raw.replace("Z", "+00:00")
    try:
        return datetime.datetime.fromisoformat(normalized)
    except Exception:
        return None


def build_vtc_credentials_reminder_email(first_name: str, form_link: str) -> Tuple[str, str]:
    first_name = (first_name or "").strip()
    greeting = f"Bonjour <strong>{first_name}</strong>," if first_name else "Bonjour,"
    subject = "Relance – Identifiants Chambre des métiers manquants"

    html = mail_layout(f"""
      <h2 style="text-align:center;color:#b91c1c">⏰ Relance – Formation Chauffeur VTC</h2>
      <p>{greeting}</p>
      <p>
        Je me permets de revenir vers vous concernant votre formation <strong>Chauffeur VTC</strong>.
      </p>
      <p>
        À ce jour, nous n'avons toujours pas reçu vos identifiants Chambre des métiers
        (<strong>exament3p</strong>) afin que nous puissions procéder au paiement de vos frais d'examen.
      </p>
      <p>
        Nous vous rappelons que vous devez nous faire parvenir vos identifiants via votre
        <strong>Espace Stagiaire</strong> (les envois par mail et tout autre moyen de communication
        ne sont pas pris en compte).
      </p>
      <p style="text-align:center;margin:18px 0;">
        <a href="{form_link}"
           style="display:inline-block;background:#1f8f4a;color:white;padding:12px 18px;border-radius:10px;text-decoration:none;font-weight:bold">
          👉 Accéder à mon Espace Stagiaire
        </a>
      </p>

      <p>
        Je vous remercie par avance.
      </p>

      <p style="margin-top:18px;">
        Bien cordialement,<br>
        <strong>Clément VAILLANT</strong><br>
        Directeur Intégrale Academy
      </p>
    """)
    return subject, html


def build_vtc_credentials_reminder_sms(first_name: str, form_link: str) -> str:
    first_name = (first_name or "").strip()
    greeting = f"Bonjour {first_name}, " if first_name else "Bonjour, "
    return (
        "Intégrale Academy ⏰ "
        f"{greeting}nous n'avons pas reçu vos identifiants Chambre des métiers (exament3p). "
        "Merci de les transmettre uniquement via votre Espace Stagiaire : "
        f"{form_link}"
    )


def _send_vtc_credentials_reminder(data: Dict[str, Any], session_obj: Dict[str, Any], trainee: Dict[str, Any], details: str) -> bool:
    link = f"{PUBLIC_STUDENT_PORTAL_BASE.rstrip('/')}/espace/{(trainee.get('public_token') or '').strip()}"
    first_name = (trainee.get("first_name") or "").strip()
    subject, html_content = build_vtc_credentials_reminder_email(first_name, link)

    trainee_email = (trainee.get("email") or "").strip()
    trainee_phone = (trainee.get("phone") or "").strip()

    email_ok = brevo_send_email(
        trainee_email,
        subject,
        html_content,
        cc_emails=["clement@integraleacademy.com"],
    ) if trainee_email else False
    sms_ok = brevo_send_sms(trainee_phone, build_vtc_credentials_reminder_sms(first_name, link)) if trainee_phone else False

    copy_subject = f"Copie relance VTC identifiants envoyée – {first_name} {(trainee.get('last_name') or '').strip()}".strip()
    copy_html = mail_layout(f"""
      <h2>Copie relance VTC</h2>
      <p><strong>Stagiaire :</strong> {_format_trainee_name(trainee.get('first_name', ''), trainee.get('last_name', ''))}</p>
      <p><strong>Session :</strong> {_session_get(session_obj, 'name', '') or '—'}</p>
      <p><strong>Email stagiaire :</strong> {trainee_email or 'Non renseigné'}</p>
      <p><strong>Téléphone stagiaire :</strong> {trainee_phone or 'Non renseigné'}</p>
      <p><strong>Email stagiaire envoyé :</strong> {'Oui' if email_ok else 'Non'}</p>
      <p><strong>SMS envoyé :</strong> {'Oui' if sms_ok else 'Non'}</p>
      <p><strong>Contexte :</strong> {details}</p>
    """)
    copy_email_ok = brevo_send_email("clement@integraleacademy.com", copy_subject, copy_html)

    trainee["vtc_cm_reminder_sent_at"] = _now_iso()
    trainee["vtc_cm_reminder_email_ok"] = bool(email_ok)
    trainee["vtc_cm_reminder_sms_ok"] = bool(sms_ok)
    trainee["vtc_cm_reminder_copy_email_ok"] = bool(copy_email_ok)
    trainee["updated_at"] = _now_iso()

    phone_followups = trainee.get("phone_followups")
    if not isinstance(phone_followups, list):
        phone_followups = []
    phone_followups.insert(0, {
        "type": "RELANCE VTC IDENTIFIANTS",
        "details": details,
        "at": _now_iso(),
        "status": "ENVOYÉE",
        "comment": "Relance envoyée (mail + SMS) pour identifiants Chambre des métiers manquants.",
    })
    trainee["phone_followups"] = phone_followups

    trainee_display_name = _format_trainee_name(trainee.get("first_name", ""), trainee.get("last_name", ""))
    add_admin_notification(
        data,
        f"⏰ Relance VTC identifiants envoyée à {trainee_display_name}",
        meta={
            "type": "vtc_credentials_reminder",
            "session_id": session_obj.get("id"),
            "trainee_id": trainee.get("id"),
        },
    )
    return bool(email_ok or sms_ok or copy_email_ok)


def _send_vtc_credentials_missing_reminders(data: Dict[str, Any]) -> bool:
    changed = False
    now_utc = datetime.datetime.utcnow().replace(tzinfo=datetime.timezone.utc)
    session_list = data.get("sessions") or []

    for session_obj in session_list:
        if session_obj.get("archived"):
            continue

        training_type = (_session_get(session_obj, "training_type", "") or "").upper()
        if "VTC" not in training_type:
            continue

        trainees = _session_trainees_list(session_obj)
        for trainee in trainees:
            if (trainee.get("vtc_cm_login") or "").strip() and (trainee.get("vtc_cm_password") or "").strip():
                continue
            if (trainee.get("vtc_cm_submitted_at") or "").strip():
                continue
            if (trainee.get("vtc_cm_reminder_sent_at") or "").strip():
                continue

            created_at = _parse_iso_datetime(trainee.get("created_at") or "")
            if not created_at:
                continue
            if created_at.tzinfo is None:
                created_at = created_at.replace(tzinfo=datetime.timezone.utc)

            days_since_creation = (now_utc - created_at).days
            if days_since_creation < 7:
                continue

            _send_vtc_credentials_reminder(data, session_obj, trainee, "Relance automatique J+7")
            changed = True

        session_obj["trainees"] = trainees
        session_obj.pop("stagiaires", None)

    return changed

# =========================
# Helpers
# =========================

# =========================
# Public trainee "mini-login" (nom + date naissance)
# =========================
import unicodedata
import threading
import re

def _norm_lastname(s: str) -> str:
    s = (s or "").strip().lower()
    s = unicodedata.normalize("NFD", s)
    s = "".join(ch for ch in s if unicodedata.category(ch) != "Mn")  # remove accents
    s = re.sub(r"[^a-z0-9]+", "", s)  # keep only alnum, no spaces
    return s

def _birth_to_ddmmyyyy(value: str) -> str:
    """
    Convertit une date stockée (ex: '1993-09-16' ou '16/09/1993' ou '16091993')
    en 'DDMMYYYY'. Si impossible, renvoie ''.
    """
    v = (value or "").strip()
    if not v:
        return ""

    # déjà au bon format
    digits = re.sub(r"\D+", "", v)
    if len(digits) == 8:
        # peut être DDMMYYYY ou YYYYMMDD -> on tente de deviner
        # si commence par 19/20 => probablement YYYYMMDD
        if digits.startswith(("19", "20")):
            return digits[6:8] + digits[4:6] + digits[0:4]  # DDMMYYYY
        return digits  # DDMMYYYY

    # formats classiques
    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%Y/%m/%d"):
        try:
            dt = datetime.datetime.strptime(v[:10], fmt)
            return dt.strftime("%d%m%Y")
        except Exception:
            pass

    return ""

def _public_is_authed(token: str) -> bool:
    # ✅ si admin connecté, on bypass toujours
    if session.get("admin_logged_in"):
        return True
    return bool(session.get(f"public_auth_{token}"))


def _now_iso() -> str:
    return datetime.datetime.utcnow().isoformat() + "Z"

def _mark_public_login(data: Dict[str, Any], session_data: Dict[str, Any], trainee: Dict[str, Any]) -> None:
    if not trainee.get("public_has_logged_in"):
        trainee["public_has_logged_in"] = True
    trainee["public_last_login_at"] = _now_iso()
    session_data["trainees"] = _session_trainees_list(session_data)
    session_data.pop("stagiaires", None)
    save_data(data)


def _normalize_cnaps_status(value: Optional[str]) -> str:
    return (value or "").strip().upper()

def _cnaps_is_accepted(value: Optional[str]) -> bool:
    normalized = _normalize_cnaps_status(value)
    if not normalized:
        return False
    if normalized in {"ACCEPTE", "ACCEPTÉ"}:
        return True
    return normalized.startswith("CARTE PROFESSIONNELLE OK")


def record_cnaps_status_change(t: Dict[str, Any], new_status: Optional[str]) -> None:
    normalized = (new_status or "").strip()
    if not normalized:
        return
    history = t.get("cnaps_history")
    if not isinstance(history, list):
        history = []
    last_status = history[-1].get("status") if history else ""
    if _normalize_cnaps_status(last_status) != _normalize_cnaps_status(normalized):
        history.append({"status": normalized, "date": _now_iso()})
    t["cnaps_history"] = history


def record_cnaps_pre_request(t: Dict[str, Any]) -> None:
    history = t.get("cnaps_history")
    if not isinstance(history, list):
        history = []
    history.append({
        "status": "Demande de numéro PRE faite au CNAPS",
        "date": _now_iso(),
        "kind": "pre_request",
    })
    t["cnaps_history"] = history


def ensure_cnaps_history(t: Dict[str, Any]) -> None:
    history = t.get("cnaps_history")
    if not isinstance(history, list):
        history = []
    t["cnaps_history"] = history
    current_status = (t.get("cnaps") or "INCONNU").strip()
    if current_status:
        record_cnaps_status_change(t, current_status)


def load_data() -> Dict[str, Any]:
    if not os.path.exists(DATA_FILE):
        base = {
            "sessions": [],
            "positioning_tests": [],
            "notifications_edof": [],
            "notifications_financement_refuse": [],
            "notifications_prelevements": [],
            "notifications_phone_relances": [],
            "notifications_cnaps_pre_relances": [],
            "notifications_test_fr": [],
            "notifications_convention_unsigned": [],
            "notifications_admin": [],
        }
        save_data(base)
        return base
    try:
        with open(DATA_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)

        # ✅ Assure que tous les stagiaires ont un public_token
        changed = False

        if ensure_public_tokens(data):
            changed = True

        # ✅ IMPORTANT : normalise en "trainees" partout (sinon admin/public désynchronisés)
        if normalize_sessions_schema(data):
            changed = True

        if "positioning_tests" not in data:
            data["positioning_tests"] = []
            changed = True
        if "notifications_edof" not in data:
            data["notifications_edof"] = []
            changed = True
        if "notifications_prelevements" not in data:
            data["notifications_prelevements"] = []
            changed = True
        if "notifications_financement_refuse" not in data:
            data["notifications_financement_refuse"] = []
            changed = True
        if "notifications_phone_relances" not in data:
            data["notifications_phone_relances"] = []
            changed = True
        if "notifications_cnaps_pre_relances" not in data:
            data["notifications_cnaps_pre_relances"] = []
            changed = True
        if "notifications_test_fr" not in data:
            data["notifications_test_fr"] = []
            changed = True
        if "notifications_convention_unsigned" not in data:
            data["notifications_convention_unsigned"] = []
            changed = True
        if "notifications_admin" not in data:
            data["notifications_admin"] = []
            changed = True

        if _send_vtc_credentials_missing_reminders(data):
            changed = True

        if changed:
            save_data(data)

        return data


    except Exception:
        try:
            backup = DATA_FILE + ".corrupt." + str(int(datetime.datetime.utcnow().timestamp()))
            os.replace(DATA_FILE, backup)
        except Exception:
            pass
        base = {
            "sessions": [],
            "positioning_tests": [],
            "notifications_edof": [],
            "notifications_financement_refuse": [],
            "notifications_prelevements": [],
            "notifications_phone_relances": [],
            "notifications_cnaps_pre_relances": [],
            "notifications_test_fr": [],
            "notifications_convention_unsigned": [],
            "notifications_admin": [],
        }
        save_data(base)
        return base




def save_data(data: Dict[str, Any]) -> None:
    _write_json_with_backups(DATA_FILE, data, _data_lock)


def _notification_id(prefix: str) -> str:
    return f"{prefix}-{uuid.uuid4().hex[:10].upper()}"


def add_notification(data: Dict[str, Any], bucket: str, label: str, meta: Optional[dict] = None) -> dict:
    notifications = data.setdefault(bucket, [])
    prefix_map = {
        "notifications_edof": "EDOF",
        "notifications_financement_refuse": "FTR",
        "notifications_prelevements": "PREL",
        "notifications_phone_relances": "REL",
        "notifications_cnaps_pre_relances": "PRE",
        "notifications_test_fr": "TFR",
        "notifications_convention_unsigned": "CNS",
        "notifications_admin": "ADM",
    }
    entry = {
        "id": _notification_id(prefix_map.get(bucket, "NOTI")),
        "label": label,
        "created_at": _now_iso(),
        "done": False,
    }
    if meta:
        entry["meta"] = meta
    notifications.insert(0, entry)
    return entry


def _notifications_bucket_key(bucket: str) -> Optional[str]:
    return {
        "edof": "notifications_edof",
        "financement_refuse": "notifications_financement_refuse",
        "prelevements": "notifications_prelevements",
        "relances": "notifications_phone_relances",
        "cnaps_pre": "notifications_cnaps_pre_relances",
        "test_fr": "notifications_test_fr",
        "convention_unsigned": "notifications_convention_unsigned",
    }.get(bucket)


def _secretariat_notifications_payload(data: Dict[str, Any]) -> Dict[str, Any]:
    def _clean(value: Any) -> str:
        return str(value or "").strip()

    def _test_fr_contact(meta: Dict[str, Any]) -> Dict[str, str]:
        session_id = _clean(meta.get("session_id"))
        trainee_id = _clean(meta.get("trainee_id"))
        if not session_id or not trainee_id:
            return {}

        session_obj = _find_session_by_id(data, session_id)
        if not session_obj:
            return {}

        trainee = next(
            (item for item in _session_trainees_list(session_obj) if _clean(item.get("id")) == trainee_id),
            None,
        )
        if not trainee:
            return {}

        return {
            "first_name": _clean(trainee.get("first_name")),
            "last_name": _clean(trainee.get("last_name")),
            "phone": _clean(trainee.get("phone")),
        }

    def _with_created_fr(items: List[Dict[str, Any]], bucket: str) -> List[Dict[str, Any]]:
        out = []
        for item in items:
            cloned = dict(item)
            cloned["created_fr"] = fr_datetime(item.get("created_at") or "")
            if bucket == "test_fr":
                meta = dict(cloned.get("meta") or {})
                fallback = _test_fr_contact(meta)
                for key in ("first_name", "last_name", "phone"):
                    if not _clean(meta.get(key)) and fallback.get(key):
                        meta[key] = fallback[key]
                if meta:
                    cloned["meta"] = meta
            out.append(cloned)
        return out

    notifications = {
        "edof": _with_created_fr(list(data.get("notifications_edof", [])), "edof"),
        "financement_refuse": _with_created_fr(list(data.get("notifications_financement_refuse", [])), "financement_refuse"),
        "prelevements": _with_created_fr(list(data.get("notifications_prelevements", [])), "prelevements"),
        "relances": _with_created_fr(list(data.get("notifications_phone_relances", [])), "relances"),
        "cnaps_pre": _with_created_fr(list(data.get("notifications_cnaps_pre_relances", [])), "cnaps_pre"),
        "test_fr": _with_created_fr(list(data.get("notifications_test_fr", [])), "test_fr"),
        "convention_unsigned": _with_created_fr(list(data.get("notifications_convention_unsigned", [])), "convention_unsigned"),
    }
    unresolved_total = 0
    for items in notifications.values():
        unresolved_total += sum(1 for item in items if not item.get("done"))
    return {
        "notifications": notifications,
        "unresolved_total": unresolved_total,
    }


def _session_period_label(session_obj: Dict[str, Any]) -> str:
    training = formation_label(_session_get(session_obj, "training_type", ""))
    start = fr_date(_session_get(session_obj, "date_start", ""))
    end = fr_date(_session_get(session_obj, "date_end", ""))
    if training and start and end:
        return f"{training} du {start} au {end}"
    if training:
        return training
    return ""


def _find_session_by_id(data: Dict[str, Any], session_id: str) -> Optional[Dict[str, Any]]:
    target = (session_id or "").strip()
    if not target:
        return None
    for session_obj in (data.get("sessions", []) or []):
        if (session_obj.get("id") or "").strip() == target:
            return session_obj
    return None


def _history_sort_key(value: str) -> float:
    s = (value or "").strip()
    if not s:
        return 0.0
    normalized = s.replace("Z", "+00:00")
    try:
        return datetime.datetime.fromisoformat(normalized).timestamp()
    except Exception:
        pass
    for fmt in ("%Y-%m-%d %H:%M", "%Y-%m-%dT%H:%M", "%Y-%m-%dT%H:%M:%S", "%Y-%m-%dT%H:%M:%S.%f"):
        try:
            return datetime.datetime.strptime(s[:26], fmt).timestamp()
        except Exception:
            pass
    return 0.0


def build_trainee_history_entries(trainee: Dict[str, Any]) -> List[Dict[str, str]]:
    entries: List[Dict[str, str]] = []

    def _add(kind: str, label: str, at: str, details: str = "") -> None:
        if not (at or "").strip():
            return
        emoji = {
            "mail": "📧",
            "relance": "🔁",
            "appel": "📞",
            "sms": "📱",
            "action": "📝",
        }.get(kind, "📌")
        entries.append({
            "kind": kind,
            "emoji": emoji,
            "label": label,
            "at": at,
            "at_display": history_datetime(at),
            "details": (details or "").strip(),
        })

    field_events = [
        ("access_sent_at", "Espace stagiaire envoyé", "mail"),
        ("docs_notified_at", "Notification envoi des documents", "mail"),
        ("docs_last_nonconform_notified_at", "Notification documents non conformes", "mail"),
        ("docs_last_relance_at", "Relance documents", "relance"),
        ("test_fr_last_notified_at", "Lien test de français envoyé", "mail"),
        ("test_fr_last_relance_at", "Relance test de français", "relance"),
        ("convention_unsigned_notified_at", "Relance convention non signée", "relance"),
        ("cnaps_pre_relance_last_sent_at", "Relance PRE envoyée", "relance"),
        ("elearning_link_sent_at", "Lien e-learning envoyé", "mail"),
        ("vtc_cm_reminder_sent_at", "Relance identifiants VTC envoyée", "relance"),
    ]
    for field_name, label, kind in field_events:
        _add(kind, label, trainee.get(field_name) or "")

    for followup in (trainee.get("phone_followups") or []):
        event_type = (followup.get("type") or "Suivi").strip()
        details = (followup.get("details") or "").strip()
        comment = (followup.get("comment") or "").strip()

        source_text = f"{event_type} {details}".upper()
        kind = "action"
        if "SMS" in source_text:
            kind = "sms"
        elif "MAIL" in source_text:
            kind = "mail"
        elif "APPEL" in source_text or "RÉPONSE SECRÉTAIRE" in source_text or "REPONSE SECRETAIRE" in source_text:
            kind = "appel"
        elif "RELANCE" in source_text:
            kind = "relance"

        full_details = " · ".join([chunk for chunk in (details, comment) if chunk])
        _add(kind, event_type, followup.get("at") or "", full_details)

    entries.sort(key=lambda item: _history_sort_key(item.get("at") or ""), reverse=True)
    return entries


def _admin_notification_details(data: Dict[str, Any], item: Dict[str, Any]) -> List[str]:
    meta = item.get("meta") or {}
    details: List[str] = []

    session_obj = _find_session_by_id(data, meta.get("session_id") or "")
    period = ""
    if session_obj:
        period = _session_period_label(session_obj)
    if not period:
        training = (meta.get("training") or "").strip()
        date_start = fr_date(meta.get("date_start") or "")
        date_end = fr_date(meta.get("date_end") or "")
        if training and date_start and date_end:
            period = f"{training} du {date_start} au {date_end}"
        elif training:
            period = training
    if period:
        details.append(period)

    comment = (meta.get("comment") or meta.get("last_comment") or "").strip()
    if comment:
        details.append(f"Commentaire : {comment}")

    call_status = (meta.get("call_status") or "").strip()
    if call_status and "appel" in call_status.lower():
        details.append(call_status)

    return details


def _admin_notification_trainee_url(data: Dict[str, Any], item: Dict[str, Any]) -> str:
    meta = item.get("meta") or {}
    session_id = (meta.get("session_id") or "").strip()
    trainee_id = (meta.get("trainee_id") or "").strip()
    if not session_id or not trainee_id:
        return ""

    session_obj = _find_session_by_id(data, session_id)
    if not session_obj:
        return ""

    trainee_exists = any(
        (trainee.get("id") or "").strip() == trainee_id
        for trainee in _session_trainees_list(session_obj)
    )
    if not trainee_exists:
        return ""

    return url_for("admin_trainee_page", session_id=session_id, trainee_id=trainee_id)


def _admin_notifications_payload(data: Dict[str, Any]) -> Dict[str, Any]:
    notifications = []
    unresolved_total = 0
    for item in list(data.get("notifications_admin", [])):
        cloned = dict(item)
        cloned["created_fr"] = fr_datetime(item.get("created_at") or "")
        cloned["details"] = _admin_notification_details(data, item)
        cloned["trainee_url"] = _admin_notification_trainee_url(data, item)
        notifications.append(cloned)
        if not item.get("done"):
            unresolved_total += 1
    return {
        "notifications": notifications,
        "unresolved_total": unresolved_total,
    }


def _admin_can_view_notifications() -> bool:
    return bool(session.get("admin_logged_in"))


def _admin_can_manage_notifications() -> bool:
    return session.get("admin_role") == "admin"


def _format_trainee_name(first_name: str, last_name: str) -> str:
    return f"{normalize_first_name(first_name)} {normalize_last_name(last_name)}".strip()


def add_admin_notification(data: Dict[str, Any], label: str, meta: Optional[dict] = None) -> dict:
    return add_notification(data, "notifications_admin", label, meta=meta)


def _find_prelevement_request(data: Dict[str, Any], entry_id: str):
    for s in data.get("sessions", []) or []:
        for t in _session_trainees_list(s):
            for req in (t.get("financement_rejected_requests") or []):
                if (req.get("id") or "").strip() == (entry_id or "").strip():
                    return s, t, req
    return None, None, None


def _find_phone_followup_entry(data: Dict[str, Any], followup_id: str):
    for s in data.get("sessions", []) or []:
        for t in _session_trainees_list(s):
            for req in (t.get("phone_followups") or []):
                if (req.get("id") or "").strip() == (followup_id or "").strip():
                    return s, t, req
    return None, None, None


def positioning_test_public_sections() -> List[Dict[str, Any]]:
    public_sections = []
    for section in POSITIONING_TEST_SECTIONS:
        public_sections.append(
            {
                "id": section["id"],
                "title": section["title"],
                "questions": [
                    {
                        "id": q["id"],
                        "text": q["text"],
                        "options": q["options"],
                    }
                    for q in section["questions"]
                ],
            }
        )
    return public_sections


def positioning_test_score(answers: Dict[str, Any]) -> Dict[str, Any]:
    score = 0
    total = POSITIONING_TEST_TOTAL
    for section in POSITIONING_TEST_SECTIONS:
        for question in section["questions"]:
            qid = question["id"]
            if qid not in answers:
                continue
            try:
                selected = int(answers[qid])
            except (TypeError, ValueError):
                continue
            if selected == question["correct_index"]:
                score += 1
    score_over_20 = round((score / total) * 20, 2) if total else 0
    return {"score": score, "total": total, "score_over_20": score_over_20}


def find_session(data: Dict[str, Any], session_id: str) -> Optional[Dict[str, Any]]:
    for s in data.get("sessions", []):
        if s.get("id") == session_id:
            return s
    return None


def ensure_public_tokens(data):
    changed = False

    for session in data.get("sessions", []):
        trainees = session.get("trainees") or session.get("stagiaires") or []

        for trainee in trainees:
            if "public_token" not in trainee or not trainee["public_token"]:
                trainee["public_token"] = uuid.uuid4().hex
                changed = True

    return changed



def find_trainee(session: Dict[str, Any], trainee_id: str) -> Optional[Dict[str, Any]]:
    for t in session.get("trainees", []):
        if t.get("id") == trainee_id:
            return t
    return None


def _session_get(s: Dict[str, Any], key: str, fallback: str = "") -> str:
    """
    Backward compatible getter: support old FR keys if needed.
    """
    if key in s and s.get(key) not in (None, ""):
        return s.get(key)

    # old keys from previous versions
    fr_map = {
        "name": "nom",
        "date_start": "date_debut",
        "date_end": "date_fin",
        "exam_date": "date_examen",
        "exam_theory_date": "date_examen_theorique",
        "exam_practice_date": "date_examen_pratique",
        "training_type": "type_formation",
        "trainees": "stagiaires",
    }
    fr_key = fr_map.get(key)
    if fr_key and fr_key in s and s.get(fr_key) not in (None, ""):
        return s.get(fr_key)

    return fallback


def _session_trainees_list(s: Dict[str, Any]) -> List[Dict[str, Any]]:
    if "trainees" in s and isinstance(s.get("trainees"), list):
        return s.get("trainees", [])
    if "stagiaires" in s and isinstance(s.get("stagiaires"), list):
        # convert on the fly (non destructif)
        out = []
        for st in s.get("stagiaires", []):
            out.append(_convert_old_stagiaire_to_trainee(st))
        return out
    return []


def _convert_old_stagiaire_to_trainee(st: Dict[str, Any]) -> Dict[str, Any]:
    # best-effort mapping
    return {
        "id": st.get("id") or ("TRN-" + uuid.uuid4().hex[:8].upper()),
        "personal_id": st.get("id") or "",
        "last_name": st.get("nom") or "",
        "first_name": st.get("prenom") or "",
        "email": st.get("email") or "",
        "phone": st.get("telephone") or "",
        "comment": st.get("commentaire") or "",
        "cnaps": (st.get("cnaps") or "INCONNU"),
        "convention_status": _map_convention_to_enum(st.get("convention")),
        "test_fr_status": _map_testfr_to_enum(st.get("test_francais")),
        "dossier_status": "complete" if (st.get("dossier") == "complet") else "incomplete",
        "financement_status": _map_financement_to_enum(st.get("financement")),
        "vae_status": _map_vae_to_enum(st.get("vae")),
        "hosting_status": _map_hosting_to_enum(st.get("hebergement")),
        "documents": st.get("documents") or [],
        "public_token": st.get("public_token") or "",
        "created_at": st.get("created_at") or "",
        "updated_at": st.get("updated_at") or "",
        "phone_followups": st.get("phone_followups") or [],
    }


def _map_convention_to_enum(v: Optional[str]) -> str:
    v = (v or "").strip().lower()
    if v in ("signée", "signee", "signed"):
        return "signed"
    if "signature" in v or v in ("en cours de signature", "signing"):
        return "signing"
    return "soon"


def _map_testfr_to_enum(v: Optional[str]) -> str:
    v = (v or "").strip().lower()
    if v in ("validé", "valide", "validated"):
        return "validated"
    if v in ("relancé", "relance", "relancé(e)", "relancee"):
        return "relance"
    if v in ("en cours", "in progress", "in_progress", "en_cours"):
        return "in_progress"
    return "soon"


def _map_financement_to_enum(v: Optional[str]) -> str:
    v = (v or "").strip().lower()
    if v in ("validé", "valide", "validated"):
        return "validated"
    if "validation" in v or v in ("en cours de validation", "in_review"):
        return "in_review"
    return "soon"


def _map_vae_to_enum(v: Optional[str]) -> str:
    v = (v or "").strip().lower()
    if v in ("validé", "valide", "validated"):
        return "validated"
    if v in ("en cours", "in_progress", "in progress"):
        return "in_progress"
    return "soon"


def _map_hosting_to_enum(v: Optional[str]) -> str:
    v = (v or "").strip().lower()
    if v in ("réservé", "reserve", "reserved"):
        return "reserved"
    return "unknown"




# =========================
# Conformity logic (matching your enums)
# =========================

def trainee_is_conform(t: Dict[str, Any], training_type: str) -> bool:
    if t.get("convention_status") != "signed":
        return False
    if t.get("test_fr_status") != "validated":
        return False
    if t.get("dossier_status") != "complete":
        return False
    if t.get("financement_status") != "validated":
        return False
    if training_type == "DIRIGEANT VAE":
        if t.get("vae_status") != "validated":
            return False
    return True


def session_is_conform(session: Dict[str, Any]) -> bool:
    training_type = _session_get(session, "training_type", "")
    trainees = _session_trainees_list(session)
    if not trainees:
        return False
    return all(trainee_is_conform(t, training_type) for t in trainees)

def normalize_sessions_schema(data: Dict[str, Any]) -> bool:
    changed = False
    for s in data.get("sessions", []):
        # Si pas de trainees, on convertit depuis stagiaires
        if "trainees" not in s or not isinstance(s.get("trainees"), list):
            s["trainees"] = _session_trainees_list(s)
            changed = True

        # On supprime l’ancienne clé pour éviter 2 sources
        if "stagiaires" in s:
            s.pop("stagiaires", None)
            changed = True

    return changed


def compute_stats(session: Dict[str, Any]) -> Dict[str, Any]:
    training_type = _session_get(session, "training_type", "")
    trainees = _session_trainees_list(session)
    conform_count = sum(1 for t in trainees if trainee_is_conform(t, training_type))
    total = len(trainees)
    cnaps_accepted_count = sum(1 for t in trainees if _cnaps_is_accepted(t.get("cnaps")))
    return {
        "total": total,
        "conform_count": conform_count,
        "non_conform_count": total - conform_count,
        "session_is_conform": (total > 0 and conform_count == total),
        "cnaps_accepted_count": cnaps_accepted_count,
    }


def _session_jury_entries(session: Dict[str, Any]) -> List[Dict[str, Any]]:
    raw = session.get("juries")
    if raw is None:
        raw = session.get("jurys")
    if raw is None:
        raw = session.get("jury")
    if isinstance(raw, dict):
        if isinstance(raw.get("items"), list):
            raw = raw.get("items")
        else:
            raw = [raw]
    if not isinstance(raw, list):
        return []
    return [item for item in raw if isinstance(item, dict)]


def _normalize_jury_status(value: str) -> str:
    raw = (value or "").strip()
    if not raw:
        return ""
    cleaned = unicodedata.normalize("NFD", raw)
    cleaned = "".join(ch for ch in cleaned if unicodedata.category(ch) != "Mn")
    cleaned = cleaned.upper()
    cleaned = " ".join(cleaned.replace("_", " ").replace("-", " ").split())
    if cleaned in ("EN ATTENTE", "EN ATTENTE DE REPONSE", "ATTENTE", "PENDING", "A CONFIRMER"):
        return "pending"
    if cleaned in ("PRESENT", "PRESENTE", "PRESENTS", "PRESENTES"):
        return "present"
    if cleaned in ("ABSENT", "ABSENTE", "ABSENTS", "ABSENTES"):
        return "absent"
    return ""


# =========================
# CNAPS / Hosting fetchers
# =========================

def fetch_cnaps_status_by_name(nom: str, prenom: str) -> Optional[str]:
    if not CNAPS_LOOKUP_ENDPOINT:
        return None

    def _normalize_cnaps_name(value: str) -> str:
        raw = (value or "").strip()
        if not raw:
            return ""
        cleaned = unicodedata.normalize("NFD", raw)
        cleaned = "".join(ch for ch in cleaned if unicodedata.category(ch) != "Mn")
        cleaned = " ".join(cleaned.replace("-", " ").split())
        return cleaned.upper()

    nom = _normalize_cnaps_name(nom)
    prenom = _normalize_cnaps_name(prenom)
    if not nom or not prenom:
        return None

    try:
        r = requests.get(CNAPS_LOOKUP_ENDPOINT, params={"nom": nom, "prenom": prenom}, timeout=10)
        if r.status_code != 200:
            return None
        data = r.json()
        return data.get("statut_cnaps") or data.get("status")
    except Exception:
        return None


import time

def fetch_hebergement_status(email: str) -> Optional[str]:
    if not HEBERGEMENT_STATUS_ENDPOINT:
        return None

    email = (email or "").strip().lower()
    if not email:
        return None

    def _is_truthy(v) -> bool:
        if v is True:
            return True
        if isinstance(v, (int, float)) and v == 1:
            return True
        if isinstance(v, str) and v.strip().lower() in ("true", "1", "yes", "y", "ok", "oui"):
            return True
        return False

    def _norm(s: str) -> str:
        return (s or "").strip().lower().replace("é", "e").replace("è", "e").replace("ê", "e")

    for attempt in range(2):
        try:
            r = requests.get(
                HEBERGEMENT_STATUS_ENDPOINT,
                params={"email": email},
                timeout=8
            )

            print("[HEBERGEMENT] status=", r.status_code, "email=", email)
            if r.status_code != 200:
                print("[HEBERGEMENT] body=", r.text[:400])
                time.sleep(0.3)
                continue

            data = r.json()
            print("[HEBERGEMENT] json=", data)

            # 1) cas idéal : bool clair
            if _is_truthy(data.get("reserved")):
                return "reserved"

            # 2) cas fréquent : champ texte
            candidates = [
                data.get("status"),
                data.get("hosting_status"),
                data.get("hebergement"),
                data.get("value"),
                data.get("result"),
            ]
            for c in candidates:
                if isinstance(c, str):
                    cc = _norm(c)
                    if cc in ("reserved", "reserve", "reserver", "reservé", "reservee", "ok", "oui"):
                        return "reserved"
                    if cc in ("unknown", "inconnu", "non", "no", "false"):
                        # on ne downgrade pas agressivement
                        return None

            # 3) si rien de concluant -> on ne touche pas l'existant
            return None

        except Exception as e:
            print("[HEBERGEMENT] exception=", repr(e))
            time.sleep(0.3)

    return None




# =========================
# UI enums (for template)
# =========================

FORMATION_TYPES = ["APS", "A3P", "DIRIGEANT initial", "DIRIGEANT VAE", "SSIAP 1", "CHEF DE POSTE", "VTC"]
FORMATION_PRICE_DEFAULTS = {
    "A3P": 4200,
    "APS": 1650,
    "DIRIGEANT INITIAL": 4300,
    "DIRIGEANT VAE": 4300,
}

ENUMS = {
    "convention": ["soon", "signing", "signed"],
    "test_fr": ["soon", "in_progress", "validated", "relance"],
    "dossier": ["complete", "incomplete"],
    "financement": ["soon", "in_review", "validated"],
    "vae": ["soon", "in_progress", "validated"],
}

# =========================
# Libellés longs (pour mails/SMS)
# =========================
FORMATION_LONG_LABELS = {
    "APS": "Agent de Prévention et de Sécurité (APS)",
    "A3P": "Agent de Protection Physique des Personnes (A3P)",
    "VTC": "Chauffeur VTC",
    "SSIAP 1": "Service de Sécurité Incendie et d’Assistance à Personnes – niveau 1 (SSIAP 1)",
    "CHEF DE POSTE": "Chef de Poste en sécurité privée / CPSP",
    "DIRIGEANT": "Dirigeant d'une entreprise de sécurité privée (DESP)",
    "DIRIGEANT INITIAL": "Dirigeant d'une entreprise de sécurité privée (DESP)",
    "DIRIGEANT VAE": "Dirigeant d'une entreprise de sécurité privée (DESP) – VAE",
}

EDOF_TRAININGS = {
    "A3P": {
        "label": "Agent de protection physique des personnes (A3P)",
        "calendly": "https://calendly.com/integraleacademy/apr",
    },
    "APS": {
        "label": "Agent de prévention et de sécurité (APS)",
        "calendly": "https://calendly.com/integraleacademy/aps",
    },
    "VTC": {
        "label": "Chauffeur VTC",
        "calendly": "https://calendly.com/integraleacademy/chauffeurvtc",
    },
    "DESP": {
        "label": "Dirigeant d'entreprise de sécurité privée (DESP)",
        "calendly": "https://calendly.com/integraleacademy/dirigeant",
    },
}

FINANCEMENT_REFUSE_TRAININGS = {
    "A3P": "Agent de protection physique des personnes (A3P)",
    "APS": "Agent de prévention et de sécurité (APS)",
    "VTC": "Chauffeur VTC",
    "DIRIGEANT": "Dirigeant d'entreprise de sécurité privée (DESP)",
}

def formation_label(training_type: str) -> str:
    """
    Retourne un libellé lisible pour les mails/SMS.
    Si non trouvé, renvoie le training_type brut.
    """
    tt = (training_type or "").strip()
    key = tt.upper()
    return FORMATION_LONG_LABELS.get(key, tt)


def default_training_price(training_type: str) -> Optional[int]:
    key = (training_type or "").strip().upper()
    return FORMATION_PRICE_DEFAULTS.get(key)


# =========================
# Documents requis par formation
# =========================

REQUIRED_DOCS = {
    "COMMON": [
        {"key": "id", "label": "Passeport OU Carte d’identité recto/verso OU Titre de séjour", "accept": "application/pdf"},
        {"key": "photo", "label": "Photo d’identité officielle (photo de face de votre visage sur fond neutre)", "accept": "image/jpeg,image/png"},
        {"key": "carte_vitale_doc", "label": "Carte vitale", "accept": "application/pdf"},
        {"key": "cnaps_doc", "label": "Autorisation CNAPS ou Carte professionnelle CNAPS (en cours de validité)", "accept": "application/pdf"},
    ],
    "A3P_ONLY": [
        {"key": "permis", "label": "Permis de conduire (obligatoire sauf si vous n’avez pas le permis)", "accept": "application/pdf"},
        {"key": "certif_med", "label": "Certificat médical (-3 mois)", "accept": "application/pdf"},
        {"key": "assurance_rc", "label": "Attestation d’assurance responsabilité civile", "accept": "application/pdf"},
    ],
    "DIRIGEANT_VAE_ONLY": [
        {
            "key": "cv",
            "label": "CV",
            "accept": "application/pdf,image/jpeg,image/png",
        },
        {
            "key": "highest_diploma",
            "label": "Diplôme le plus élevé",
            "accept": "application/pdf,image/jpeg,image/png",
        },
        {
            "key": "candidate_info_sheet",
            "label": "Fiche de renseignement candidat à compléter",
            "accept": "application/pdf,image/jpeg,image/png",
        },
    ],
}

def required_docs_for_training(training_type: str) -> List[Dict[str, Any]]:
    tt = (training_type or "").strip().upper()
    docs = list(REQUIRED_DOCS["COMMON"])

    # Pour les parcours dirigeant (initial / VAE),
    # ne pas demander le justificatif CNAPS à l'inscription.
    if tt in {"DIRIGEANT INITIAL", "DIRIGEANT VAE"}:
        docs = [d for d in docs if d.get("key") != "cnaps_doc"]

    if tt == "A3P":
        docs += list(REQUIRED_DOCS["A3P_ONLY"])
    if tt == "DIRIGEANT VAE":
        docs += list(REQUIRED_DOCS["DIRIGEANT_VAE_ONLY"])
    return docs

def ensure_documents_schema_for_trainee(t: Dict[str, Any], training_type: str) -> bool:
    """
    S'assure que t["documents"] contient tous les docs requis pour la formation,
    sans écraser fichiers/statuts existants. Supprime l'ancien doc 'dom' (domicile).
    """
    required = required_docs_for_training(training_type)
    existing = t.get("documents") or []
    changed = False

    # index existant
    by_key = {d.get("key"): d for d in existing if isinstance(d, dict) and d.get("key")}

    out = []
    for rd in required:
        k = rd["key"]
        if k in by_key:
            d = by_key[k]
            if not d.get("label"):
                d["label"] = rd["label"]; changed = True
            if "accept" not in d:
                d["accept"] = rd.get("accept", ""); changed = True
            if "status" not in d:
                d["status"] = "NON DÉPOSÉ"; changed = True
            if "comment" not in d:
                d["comment"] = ""; changed = True
            if "file" not in d:
                d["file"] = ""; changed = True
            if "files" not in d or not isinstance(d.get("files"), list):
                d["files"] = []
                changed = True
            out.append(d)
        else:
            out.append({
                "key": k,
                "label": rd["label"],
                "accept": rd.get("accept", ""),
                "status": "NON DÉPOSÉ",
                "comment": "",
                "file": "",
                "files": [],
            })
            changed = True

    # 🔥 on vire dom (plus utilisé)
    if "dom" in by_key:
        changed = True

    t["documents"] = out
    return changed


def _ensure_livret2_document_entry(t: Dict[str, Any]) -> Dict[str, Any]:
    """Garantit la présence du document technique livret_2 pour les uploads publics VAE."""
    docs = t.get("documents")
    if not isinstance(docs, list):
        docs = []
        t["documents"] = docs

    existing = next((d for d in docs if isinstance(d, dict) and d.get("key") == "livret_2"), None)
    if existing:
        if "files" not in existing or not isinstance(existing.get("files"), list):
            existing["files"] = []
        if "file" not in existing:
            existing["file"] = ""
        if "status" not in existing:
            existing["status"] = "NON DÉPOSÉ"
        if "accept" not in existing:
            existing["accept"] = "application/pdf,image/jpeg,image/png"
        if "label" not in existing:
            existing["label"] = "Livret 2"
        return existing

    livret2_doc = {
        "key": "livret_2",
        "label": "Livret 2",
        "accept": "application/pdf,image/jpeg,image/png",
        "status": "NON DÉPOSÉ",
        "comment": "",
        "file": "",
        "files": [],
    }
    docs.append(livret2_doc)
    return livret2_doc

def allowed_doc_keys_for_training(training_type: str) -> set:
    keys = {d["key"] for d in required_docs_for_training(training_type)}
    if (training_type or "").strip().upper() == "DIRIGEANT VAE":
        keys.add("livret_2")
    return keys

def dossier_is_complete(trainee: Dict[str, Any], training_type: str) -> bool:
    """
    Complet si TOUS les docs requis sont CONFORME,
    sauf permis si trainee a coché no_permis=True (A3P).
    """
    docs = trainee.get("documents") or []
    if not docs:
        return False

    by_key = {d.get("key"): d for d in docs if isinstance(d, dict)}

    tt = (training_type or "").strip().upper()
    no_permis = bool(trainee.get("no_permis"))  # checkbox "je n'ai pas le permis"

    for rd in required_docs_for_training(training_type):
        k = rd["key"]

        # permis optionnel si no_permis
        if tt == "A3P" and k == "permis" and no_permis:
            continue

        d = by_key.get(k)
        if not d:
            return False

        st = (d.get("status") or "").strip().upper()
        if st != "CONFORME":
            return False

    return True


def required_docs_are_deposited(trainee: Dict[str, Any], training_type: str) -> bool:
    """Vrai si tous les documents requis sont déposés (peu importe leur conformité)."""
    docs = trainee.get("documents") or []
    if not docs:
        return False

    by_key = {d.get("key"): d for d in docs if isinstance(d, dict)}
    tt = (training_type or "").strip().upper()
    no_permis = bool(trainee.get("no_permis"))

    for rd in required_docs_for_training(training_type):
        k = rd["key"]

        if tt == "A3P" and k == "permis" and no_permis:
            continue

        d = by_key.get(k)
        if not d:
            return False

        # La fiche candidat (DIRIGEANT VAE) est un formulaire HTML sauvegardé,
        # pas forcément un upload de fichier binaire.
        if k == "candidate_info_sheet":
            st = (d.get("status") or "").strip().upper()
            candidate_sheet_saved = bool(trainee.get("candidate_sheet_saved_at") or trainee.get("candidate_sheet"))
            if st in {"A CONTRÔLER", "A CONTROLER", "CONFORME", "NON CONFORME"}:
                continue
            if candidate_sheet_saved and st not in {"", "NON DÉPOSÉ", "NON DEPOSE"}:
                continue
            return False

        files = d.get("files") if isinstance(d.get("files"), list) else []
        has_files = any(f for f in files)
        has_file = bool((d.get("file") or "").strip())
        if not (has_files or has_file):
            return False

    return True

    
import re

def infos_is_complete(t: Dict[str, Any]) -> bool:
    # Champs obligatoires
    required = [
        "birth_date",
        "birth_city",
        "birth_country",
        "nationality",
        "address",
        "zip_code",
        "city",
    ]
    for k in required:
        if not (t.get(k) or "").strip():
            return False

    # Sécu : 15 chiffres
    secu_digits = re.sub(r"\D+", "", (t.get("carte_vitale") or ""))
    if len(secu_digits) != 15:
        return False

    # PRE : format PRE-083-2025-12-01-20250000000 ou PRE-2025-12-01-20250000000
    pre = (t.get("pre_number") or "").strip().upper().replace(" ", "")
    if not re.match(r"^(PRE|CAR)-(?:\d{3}-)?\d{4}-\d{2}-\d{2}-\d{11,}$", pre):
        return False

    return True

def dossier_is_complete_total(trainee: Dict[str, Any], training_type: str) -> bool:
    # ✅ complet seulement si infos OK + tous docs CONFORME
    # ✅ OU si forçage admin
    if trainee.get("force_dossier_complete"):
        return True
    return infos_is_complete(trainee) and dossier_is_complete(trainee, training_type)


# =========================
# Pages (templates)
# =========================

POSITIONING_TEST_SECTIONS = [
    {
        "id": "section-1",
        "title": "Section 1 – Institutions françaises, Principes généraux de République et économie (14 questions)",
        "questions": [
            {
                "id": "s1q1",
                "text": "Quelle est la devise de la République Française ?",
                "options": [
                    "Liberté, égalité, fraternité",
                    "Travail, Famille, Patrie",
                    "Unité, justice, progrès",
                ],
                "correct_index": 0,
            },
            {
                "id": "s1q2",
                "text": "Quelle est le régime politique de la France ?",
                "options": [
                    "Dictature",
                    "Monarchie absolue",
                    "République parlementaire",
                    "République présidentielle",
                ],
                "correct_index": 2,
            },
            {
                "id": "s1q3",
                "text": "Qui est le chef de l’Etat en France ?",
                "options": [
                    "Le premier Ministre",
                    "Le Président de la République",
                    "Le Président de l’Assemblée nationale",
                ],
                "correct_index": 1,
            },
            {
                "id": "s1q4",
                "text": "Quel est le principal symbole de la République française ?",
                "options": [
                    "La Marseillaise",
                    "Le coq",
                    "Le drapeau tricolore",
                ],
                "correct_index": 2,
            },
            {
                "id": "s1q5",
                "text": "Quel document est considéré comme la loi fondamentale de la République ?",
                "options": [
                    "La Constitution",
                    "Le Code civil",
                    "La Déclaration des Droits de l'Homme et du Citoyen",
                ],
                "correct_index": 2,
            },
            {
                "id": "s1q6",
                "text": "Qui propose les lois en France ?",
                "options": [
                    "Les députés et les sénateurs",
                    "Le Président uniquement",
                    "Le Premier ministre uniquement",
                ],
                "correct_index": 0,
            },
            {
                "id": "s1q7",
                "text": "Qui est le chef du gouvernement français ?",
                "options": [
                    "Le Président",
                    "Le Premier ministre",
                    "Le ministre de l’Intérieur",
                ],
                "correct_index": 1,
            },
            {
                "id": "s1q8",
                "text": "A quoi fait référence la date du 14 juillet, jour férié en France ?",
                "options": [
                    "Fin de la Première guerre mondiale",
                    "Prise de la Bastille",
                    "Fin de la Seconde guerre mondiale",
                    "Fête des feux d’artifices",
                ],
                "correct_index": 1,
            },
            {
                "id": "s1q9",
                "text": "Quel est le principe de la laïcité ?",
                "options": [
                    "Les religions sont interdites",
                    "La séparation de l’Eglise et de l’Etat",
                    "L’Etat soutient toutes les religions",
                ],
                "correct_index": 1,
            },
            {
                "id": "s1q10",
                "text": "Parmi les organismes publics cités ci-dessous, lequel est une collectivité territoriale ?",
                "options": [
                    "Le ministère de l’Économie",
                    "La mairie de Nice",
                    "Le centre des finances publiques de Fréjus",
                ],
                "correct_index": 1,
            },
            {
                "id": "s1q11",
                "text": "Comment s’appelle l’élu à la tête d’un département ?",
                "options": [
                    "Le préfet",
                    "Le Président du conseil général",
                    "Le Maire",
                    "Le Commissaire",
                ],
                "correct_index": 1,
            },
            {
                "id": "s1q12",
                "text": "Les décisions du préfet sont des :",
                "options": [
                    "Décrets",
                    "Arrêtés",
                    "Ordonnances",
                    "Lois",
                ],
                "correct_index": 1,
            },
            {
                "id": "s1q13",
                "text": "Les ministres sont :",
                "options": [
                    "Désignés par les députés",
                    "Désignés par les députés et les sénateurs",
                    "Nommés par le Premier ministre",
                    "Nommés par le Président de la République",
                ],
                "correct_index": 2,
            },
            {
                "id": "s1q14",
                "text": "Comment s’appellent les élus de l’Assemblée nationale ?",
                "options": [
                    "Les sénateurs",
                    "Les députés",
                    "Les ministres",
                    "Les aristocrates",
                ],
                "correct_index": 1,
            },
        ],
    },
    {
        "id": "section-2",
        "title": "Section 2 – Police, justice, sécurité (16 questions)",
        "questions": [
            {
                "id": "s2q1",
                "text": "Quelle est la principale mission de la police nationale ?",
                "options": [
                    "Protéger la sécurité intérieure",
                    "Assurer la sécurité routière",
                    "Contrôler les frontières",
                ],
                "correct_index": 0,
            },
            {
                "id": "s2q2",
                "text": "Qu’est-ce qui distingue la gendarmerie de la police nationale ?",
                "options": [
                    "La gendarmerie est civile, la police est militaire",
                    "La gendarmerie est militaire, la police est civile",
                    "Aucune différence",
                ],
                "correct_index": 1,
            },
            {
                "id": "s2q3",
                "text": "Qui supervise la police nationale ?",
                "options": [
                    "Le ministère de la Défense",
                    "Le ministère de l’Intérieur",
                    "Le ministère des Transports",
                ],
                "correct_index": 1,
            },
            {
                "id": "s2q4",
                "text": "Quel est le rôle principal des agents de sécurité privée ?",
                "options": [
                    "Contrôler les billets à l’entrée d’un événement",
                    "Protéger les biens et les personnes",
                    "Faire la circulation",
                ],
                "correct_index": 1,
            },
            {
                "id": "s2q5",
                "text": "Quel numéro dois-je composer pour appeler les forces de l’ordre ?",
                "options": [
                    "15",
                    "17",
                    "18",
                ],
                "correct_index": 1,
            },
            {
                "id": "s2q6",
                "text": "Les agents de sécurité privée ont-ils les mêmes droits que les forces de l’ordre ?",
                "options": [
                    "Oui",
                    "Non",
                    "Cela dépend de la situation",
                ],
                "correct_index": 1,
            },
            {
                "id": "s2q7",
                "text": "Quelle est la plus haute juridiction en France ?",
                "options": [
                    "La Cour d’Appel",
                    "La Cour de cassation",
                    "Le Conseil d’Etat",
                ],
                "correct_index": 2,
            },
            {
                "id": "s2q8",
                "text": "Quel est le rôle de la Cour d’Assises ?",
                "options": [
                    "Juger les affaires civiles",
                    "Juger les crimes graves",
                    "Juger les infractions mineures",
                ],
                "correct_index": 1,
            },
            {
                "id": "s2q9",
                "text": "Quel est le tribunal chargé de régler les conflits du travail ?",
                "options": [
                    "Le Tribunal de grande instance",
                    "Les Prud’hommes",
                    "Le Tribunal administratif",
                ],
                "correct_index": 1,
            },
            {
                "id": "s2q10",
                "text": "Qui supervise la Police Municipale ?",
                "options": [
                    "Le ministère de la Défense",
                    "Le ministère de l’Intérieur",
                    "Le Maire",
                ],
                "correct_index": 2,
            },
            {
                "id": "s2q11",
                "text": "Quel est le rôle principal de l'armée ?",
                "options": [
                    "Protéger les citoyens",
                    "Défendre le pays contre les agressions extérieures",
                    "Appliquer la loi",
                ],
                "correct_index": 1,
            },
            {
                "id": "s2q12",
                "text": "Quel type d’opération la gendarmerie exerce-t-elle principalement ?",
                "options": [
                    "Mission de secours",
                    "Opération de maintien de l’ordre",
                    "Aucune intervention, uniquement des enquêtes",
                ],
                "correct_index": 1,
            },
            {
                "id": "s2q13",
                "text": "Quelle opération de l'armée française a pour but de mener des actions de lutte contre le terrorisme sur le territoire national ?",
                "options": [
                    "Sentinelle",
                    "Barkhane",
                    "Serval",
                ],
                "correct_index": 0,
            },
            {
                "id": "s2q14",
                "text": "Quel est le système d'alerte d'urgence en cas de menace terroriste en France ?",
                "options": [
                    "Plan Antiterrorisme",
                    "Plan Urgence Sécurité",
                    "Plan Vigipirate",
                ],
                "correct_index": 2,
            },
            {
                "id": "s2q15",
                "text": "Quelle est la mission des forces de l’ordre lors d’une manifestation ?",
                "options": [
                    "Arrêter tous les manifestants",
                    "Préserver l’ordre public tout en garantissant le droit de manifester",
                    "Disperser la manifestation sans complexe",
                ],
                "correct_index": 1,
            },
            {
                "id": "s2q16",
                "text": "En cas d'agression, que doit faire un agent de sécurité privée avant tout ?",
                "options": [
                    "Appeler la police",
                    "Intervenir directement",
                    "Ignorer la situation",
                ],
                "correct_index": 0,
            },
        ],
    },
    {
        "id": "section-3",
        "title": "Section 3 : Compréhension de la langue française (12 questions)",
        "questions": [
            {
                "id": "s3q1",
                "text": "Dans la phrase \"Les agents de sécurité ont procédé à des palpations de sécurité\", quel est le temps verbal ?",
                "options": [
                    "Passé composé",
                    "Présent de l’indicatif",
                    "Futur",
                ],
                "correct_index": 0,
            },
            {
                "id": "s3q2",
                "text": "Quelle est la forme correcte au futur du verbe \"aller\" à la deuxième personne du pluriel ?",
                "options": [
                    "Vous allez",
                    "Vous irez",
                    "Tu iras",
                ],
                "correct_index": 1,
            },
            {
                "id": "s3q3",
                "text": "Quelle phrase est grammaticalement correcte ?",
                "options": [
                    "Elle prit les livres",
                    "Elle a pris les livres",
                    "Elle a pris le livres",
                ],
                "correct_index": 1,
            },
            {
                "id": "s3q4",
                "text": "Dans la phrase « Il a appelé la police », quel est le sujet ?",
                "options": [
                    "Police",
                    "Il",
                    "Appelé",
                ],
                "correct_index": 1,
            },
            {
                "id": "s3q5",
                "text": "Quel est la définition du mot « éloquent » ?",
                "options": [
                    "Qui parle peu",
                    "Qui sait bien s’exprimer",
                    "Qui est difficile à comprendre",
                ],
                "correct_index": 1,
            },
            {
                "id": "s3q6",
                "text": "Quel est l’antonyme du mot « heureux » ?",
                "options": [
                    "Content",
                    "Triste",
                    "Energique",
                ],
                "correct_index": 1,
            },
            {
                "id": "s3q7",
                "text": "Quel est la bonne conjugaison du verbe « avoir » à la première personne du singulier au présent de l’indicatif ?",
                "options": [
                    "Je suis",
                    "J’avais",
                    "J’ai",
                ],
                "correct_index": 2,
            },
            {
                "id": "s3q8",
                "text": "Lequel de ces mots est de genre féminin ?",
                "options": [
                    "Entracte",
                    "Ovule",
                    "Oasis",
                    "Eloge",
                ],
                "correct_index": 2,
            },
            {
                "id": "s3q9",
                "text": "Dans cette phrase « Pierre porte une parka noire », quel est l’adjectif ?",
                "options": [
                    "Parka",
                    "Pierre",
                    "Noire",
                ],
                "correct_index": 2,
            },
            {
                "id": "s3q10",
                "text": "Dans la phrase \"Elle parle doucement\", quel est le complément ?",
                "options": [
                    "Elle",
                    "Doucement",
                    "Parle",
                ],
                "correct_index": 1,
            },
            {
                "id": "s3q11",
                "text": "Quelle phrase est correcte ?",
                "options": [
                    "Je suis allé au cinéma hier",
                    "Je suis allée au cinémas hier",
                    "Je suis allé au cinémas hier",
                ],
                "correct_index": 0,
            },
            {
                "id": "s3q12",
                "text": "Quel pronom personnel peut remplacer le groupe nominal \"les enfants\" dans la phrase suivante : \"Les enfants jouent\" ?",
                "options": [
                    "Ils",
                    "Leur",
                    "Les",
                ],
                "correct_index": 0,
            },
        ],
    },
    {
        "id": "section-4",
        "title": "Section 4 : Logique et calculs (14 questions)",
        "questions": [
            {
                "id": "s4q1",
                "text": "Si une tarte est coupée en 8 parts égales et que 2 parts sont mangées, quelle fraction de la tarte reste-t-il ?",
                "options": [
                    "1/2",
                    "3/8",
                    "6/8",
                ],
                "correct_index": 2,
            },
            {
                "id": "s4q2",
                "text": "Dans une pièce, il y a 5 chaises : une rouge, une bleue, une verte, une jaune et une noire. Si on enlève la chaise verte et qu’on ajoute 2 chaises oranges, combien de chaises noires reste-t-il ?",
                "options": [
                    "3",
                    "4",
                    "6",
                    "1",
                ],
                "correct_index": 3,
            },
            {
                "id": "s4q3",
                "text": "Lors d'une course, si Paul court plus vite que tous les autres et que seuls deux coureurs dépassent Paul, quel est son rang à l'arrivée ?",
                "options": [
                    "Dernier",
                    "Premier",
                    "Troisième",
                ],
                "correct_index": 2,
            },
            {
                "id": "s4q4",
                "text": "Si un train part de Paris à 14h00, qu’il ne subit aucun retard et qu'il roule à une vitesse constante de 90 km/h, à quelle heure arrivera-t-il s’il parcours 180 km ?",
                "options": [
                    "15h00",
                    "15h30",
                    "16h00",
                ],
                "correct_index": 2,
            },
            {
                "id": "s4q5",
                "text": "Si une pizza coûte 12 euros et qu'un client en achète 3, quel est le coût total pour le client après application d’une remise fidélité de 3 euros ?",
                "options": [
                    "24 euros",
                    "36 euros",
                    "33 euros",
                    "48 euros",
                ],
                "correct_index": 2,
            },
            {
                "id": "s4q6",
                "text": "Dans une famille, le père a 40 ans et la mère a 38 ans. Leur fils a 8 ans. Dans combien d'années le fils aura-t-il la moitié de l'âge de son père ?",
                "options": [
                    "8",
                    "12",
                    "16",
                    "20",
                ],
                "correct_index": 1,
            },
            {
                "id": "s4q7",
                "text": "Quelle est la suite logique qui suit ce modèle ? 2, 4, 8, 16, ?",
                "options": [
                    "24",
                    "30",
                    "32",
                    "40",
                ],
                "correct_index": 2,
            },
            {
                "id": "s4q8",
                "text": "Dans un sac, il y a 5 billes rouges, 3 billes bleues et 2 billes vertes. Si l’on tire une bille au hasard, quelle est la probabilité d’en tirer une bleue ?",
                "options": [
                    "1/5",
                    "1/3",
                    "3/10",
                    "3/5",
                ],
                "correct_index": 2,
            },
            {
                "id": "s4q9",
                "text": "Si 4 x 3 = 12, quelle est la valeur de 12 ÷ 4 ?",
                "options": [
                    "2",
                    "3",
                    "4",
                    "5",
                ],
                "correct_index": 1,
            },
            {
                "id": "s4q10",
                "text": "Si une boîte a la forme d'un cube, combien de faces a-t-elle ?",
                "options": [
                    "4",
                    "5",
                    "6",
                    "8",
                ],
                "correct_index": 2,
            },
            {
                "id": "s4q11",
                "text": "Si \"X\" est une abeille et \"Y\" est un insecte, alors :",
                "options": [
                    "X est un insecte",
                    "Y est une abeille",
                    "X n’est pas un insecte",
                ],
                "correct_index": 1,
            },
            {
                "id": "s4q12",
                "text": "Si vous avez un livre qui compte 300 pages et que vous en lisez 50 par jour, combien de jours vous faudra-t-il pour le terminer ?",
                "options": [
                    "6 jours",
                    "5 jours",
                    "7 jours",
                ],
                "correct_index": 0,
            },
            {
                "id": "s4q13",
                "text": "Si A est plus grand que B, et B est plus grand que C, qu'est-ce que cela implique pour A et C ?",
                "options": [
                    "A est égal à C",
                    "A est plus petit que C",
                    "A est plus grand que C",
                ],
                "correct_index": 2,
            },
            {
                "id": "s4q14",
                "text": "Si un film dure 120 minutes et commence à 14h00, à quelle heure se termine-t-il ?",
                "options": [
                    "15h00",
                    "15h30",
                    "16h00",
                ],
                "correct_index": 2,
            },
        ],
    },
    {
        "id": "section-5",
        "title": "Section 5 : Economie et culture générale (4 questions)",
        "questions": [
            {
                "id": "s5q1",
                "text": "Qu’est-ce qu’un ménage ?",
                "options": [
                    "Le nettoyage de locaux",
                    "Un ensemble de personnes vivant ensemble",
                    "Une méthode de cuisine",
                ],
                "correct_index": 1,
            },
            {
                "id": "s5q2",
                "text": "Dans une économie, la rémunération du travail est généralement appelée :",
                "options": [
                    "Salaire",
                    "Dividende",
                    "Intérêt",
                ],
                "correct_index": 0,
            },
            {
                "id": "s5q3",
                "text": "Quel indicateur économique permet de mesurer le niveau de vie d'une population ?",
                "options": [
                    "Taux de chômage",
                    "PIB par habitant",
                    "Taux d’inflation",
                ],
                "correct_index": 1,
            },
            {
                "id": "s5q4",
                "text": "Que signifie la \"consommation\" dans le contexte économique ?",
                "options": [
                    "L’acquisition de biens et de services par les ménages",
                    "La production de biens et de services",
                    "L’épargne et l’investissement",
                ],
                "correct_index": 0,
            },
        ],
    },
]

POSITIONING_TEST_TOTAL = sum(
    len(section["questions"]) for section in POSITIONING_TEST_SECTIONS
)

@app.get("/")
def home():
    return redirect(url_for("admin_sessions"))


@app.get("/test-positionnement")
def positioning_test_public():
    return render_template(
        "positioning_test_public.html",
        sections=positioning_test_public_sections(),
        total_questions=POSITIONING_TEST_TOTAL,
        total_over_20=20,
    )


@app.get("/positioning_test_public.html")
def positioning_test_public_legacy():
    return redirect(url_for("positioning_test_public"), code=301)


@app.get("/admin/sessions")
@admin_login_required
def admin_sessions():
    data = load_data()
    out_sessions = []
    for s in data.get("sessions", []):
        if bool(s.get("archived")):
            continue

        trainees = _session_trainees_list(s)
        st = compute_stats(s)

        # ✅ docs fin de formation : nb de stagiaires COMPLETS / nb stagiaires
        done_total = 0
        for t in trainees:
            _, _, ok = deliverables_progress(t)
            if ok:
                done_total += 1

        public_logged_in_total = sum(
            1 for t in trainees if bool(t.get("public_has_logged_in"))
        )
        cmar_registered_total = sum(
            1 for t in trainees if bool((t.get("vtc_cm_submitted_at") or "").strip())
        )
        
        total_total = len(trainees)
        dossier_complete_total = 0
        for t in trainees:
            if dossier_is_complete_total(t, _session_get(s, "training_type", "")):
                dossier_complete_total += 1
        session_dossier_complete = (total_total > 0 and dossier_complete_total == total_total)

        jury_entries = _session_jury_entries(s)
        jury_notified = bool(s.get("juries_notified_at") or s.get("jury_notified_at"))
        jury_counts = {"pending": 0, "present": 0, "absent": 0}
        for jury in jury_entries:
            status_raw = jury.get("status") or jury.get("state") or jury.get("response")
            status_key = _normalize_jury_status(status_raw)
            if status_key:
                jury_counts[status_key] += 1
                jury_notified = True
            if jury.get("notified_at") or jury.get("notified"):
                jury_notified = True

        date_start_raw = _session_get(s, "date_start", "")
        date_end_raw = _session_get(s, "date_end", "")
        today = datetime.date.today()
        status_key = "upcoming"
        try:
            dt_start = datetime.datetime.strptime(date_start_raw[:10], "%Y-%m-%d").date()
        except (ValueError, TypeError):
            dt_start = None
        try:
            dt_end = datetime.datetime.strptime(date_end_raw[:10], "%Y-%m-%d").date()
        except (ValueError, TypeError):
            dt_end = None

        if dt_end and dt_end < today:
            status_key = "ended"
        elif dt_start and dt_start <= today:
            status_key = "ongoing"

        status_label = {
            "ended": "Formation terminée",
            "ongoing": "Formation en cours",
            "upcoming": "Prochainement",
        }[status_key]

        training_type_raw = (_session_get(s, "training_type", "") or "").strip().upper()
        if training_type_raw.startswith("APS"):
            training_type_class = "aps"
        elif training_type_raw.startswith("A3P"):
            training_type_class = "a3p"
        elif "VTC" in training_type_raw:
            training_type_class = "vtc"
        elif training_type_raw.startswith("DIRIGEANT"):
            training_type_class = "dirigeant"
        else:
            training_type_class = "other"

        out_sessions.append({
            "id": s.get("id"),
            "name": _session_get(s, "name", ""),
            "training_type": _session_get(s, "training_type", ""),
            "date_start": _session_get(s, "date_start", ""),
            "date_end": _session_get(s, "date_end", ""),
            "exam_date": _session_get(s, "exam_date", ""),
            "exam_theory_date": _session_get(s, "exam_theory_date", ""),
            "exam_practice_date": _session_get(s, "exam_practice_date", ""),
            "practice_training_date": _session_get(s, "practice_training_date", ""),
            "total": st["total"],
            "session_is_conform": st["session_is_conform"],
            "session_dossier_complete": session_dossier_complete,

            # ✅ new
            "deliverables_done": done_total,
            "deliverables_total": total_total,
            "public_logged_in_total": public_logged_in_total,
            "cmar_registered_total": cmar_registered_total,
            "status_label": status_label,
            "status_key": status_key,
            "training_type_class": training_type_class,
            "jury_notified": jury_notified,
            "jury_pending": jury_counts["pending"],
            "jury_present": jury_counts["present"],
            "jury_absent": jury_counts["absent"],
        })

    return render_template(
        "admin_sessions.html",
        sessions=out_sessions,
        formation_types=FORMATION_TYPES,
    )


@app.get("/admin/gestion-secretariat")
@admin_login_required
def admin_secretariat():
    data = load_data()
    payload = _secretariat_notifications_payload(data)
    notifications = payload["notifications"]
    return render_template(
        "admin_secretariat.html",
        edof_notifications=notifications["edof"],
        financement_refuse_notifications=notifications["financement_refuse"],
        prelevement_notifications=notifications["prelevements"],
        phone_notifications=notifications["relances"],
        cnaps_pre_notifications=notifications["cnaps_pre"],
        test_fr_notifications=notifications["test_fr"],
        convention_unsigned_notifications=notifications["convention_unsigned"],
        unresolved_total=payload["unresolved_total"],
    )


@app.get("/api/secretariat/notifications")
@admin_login_required
def api_secretariat_notifications():
    data = load_data()
    payload = _secretariat_notifications_payload(data)
    return jsonify({"ok": True, **payload})


@app.get("/api/admin/notifications")
@admin_login_required
def api_admin_notifications():
    if not _admin_can_view_notifications():
        return jsonify({"ok": False, "error": "forbidden"}), 403
    data = load_data()
    payload = _admin_notifications_payload(data)
    return jsonify({"ok": True, **payload})


@app.post("/api/admin/notifications/<notification_id>/toggle")
@admin_login_required
@admin_write_required
def api_admin_notification_toggle(notification_id: str):
    if not _admin_can_manage_notifications():
        return jsonify({"ok": False, "error": "forbidden"}), 403

    data = load_data()
    notifications = data.get("notifications_admin", [])
    entry = next((item for item in notifications if item.get("id") == notification_id), None)
    if not entry:
        return jsonify({"ok": False, "error": "notification_not_found"}), 404

    done = bool(entry.get("done"))
    entry["done"] = not done
    entry["done_at"] = _now_iso() if entry["done"] else ""

    save_data(data)
    payload = _admin_notifications_payload(data)
    return jsonify({"ok": True, "done": bool(entry.get("done")), **payload})


@app.post("/api/admin/notifications/<notification_id>/delete")
@admin_login_required
@admin_write_required
def api_admin_notification_delete(notification_id: str):
    if not _admin_can_manage_notifications():
        return jsonify({"ok": False, "error": "forbidden"}), 403

    data = load_data()
    notifications = data.get("notifications_admin", [])
    entry = next((item for item in notifications if item.get("id") == notification_id), None)
    if not entry:
        return jsonify({"ok": False, "error": "notification_not_found"}), 404

    data["notifications_admin"] = [item for item in notifications if item.get("id") != notification_id]
    save_data(data)
    payload = _admin_notifications_payload(data)
    return jsonify({"ok": True, **payload})


@app.post("/admin/edof/submit")
@admin_login_required
def admin_edof_submit():
    payload = request.get_json(silent=True) or {}
    last_name = (payload.get("last_name") or "").strip()
    first_name = (payload.get("first_name") or "").strip()
    phone = (payload.get("phone") or "").strip()
    email = (payload.get("email") or "").strip()
    training_key = (payload.get("training") or "").strip()

    if not all([last_name, first_name, phone, email, training_key]):
        return jsonify({"ok": False, "error": "missing_fields"}), 400

    training = EDOF_TRAININGS.get(training_key)
    if not training:
        return jsonify({"ok": False, "error": "invalid_training"}), 400

    training_label = training["label"]
    calendly_url = training["calendly"]

    admin_subject = f"🟦 Demande CPF (EDOF) – {first_name} {last_name}".strip()
    admin_html = mail_layout(f"""
      <p>Cette personne a fait une demande d'inscription en formation depuis son Compte Personnel de Formation (CPF), il faudrait la rappeler pour lui donner tous les renseignements et prendre un RDV téléphonique pour finaliser son inscription.</p>

      <div style="background:#f3f4f6;border:1px solid #e5e7eb;border-radius:12px;padding:14px;margin:14px 0">
        <p style="margin:0 0 8px 0"><strong>Nom :</strong> {last_name}</p>
        <p style="margin:0 0 8px 0"><strong>Prénom :</strong> {first_name}</p>
        <p style="margin:0 0 8px 0"><strong>Email :</strong> {email}</p>
        <p style="margin:0 0 8px 0"><strong>Téléphone :</strong> {phone}</p>
        <p style="margin:0"><strong>Formation :</strong> {training_label}</p>
      </div>
    """)

    admin_email_ok = brevo_send_email(
        "clement@integraleacademy.com",
        admin_subject,
        admin_html,
        cc_emails=["znaw83@gmail.com"],
    )

    data = load_data()
    add_notification(
        data,
        "notifications_edof",
        f"{first_name} {last_name} • {training_label} • {phone} • {email}",
        meta={
            "first_name": first_name,
            "last_name": last_name,
            "training": training_label,
            "phone": phone,
            "email": email,
        },
    )
    save_data(data)

    user_subject = f"Votre demande d'inscription – {training_label}".strip()
    user_html = mail_layout(f"""
      <p>Bonjour {first_name},</p>

      <p>Je me permets de revenir vers vous concernant votre demande d'inscription en formation <strong>{training_label}</strong> depuis votre Compte Personnel de Formation (CPF). Je vous remercie pour votre demande !</p>

      <p>Afin que nous puissions finaliser ensemble votre inscription vous serait-il possible de nous contacter au 04 22 47 07 68 ?</p>

      <p>Vous pouvez également réserver un RDV téléphonique en cliquant ici :</p>

      <p style="text-align:center;margin:18px 0">
        <a href="{calendly_url}"
           style="display:inline-block;background:#2563eb;color:#fff;padding:12px 16px;border-radius:10px;text-decoration:none;font-weight:800">
          Réserver un RDV téléphonique
        </a>
      </p>

      <p>Je vous remercie par avance et je vous souhaite une bonne journée,</p>

      <p>A très bientôt !</p>

      <p>Clément VAILLANT<br>Directeur Intégrale Academy</p>
    """)

    email_ok = brevo_send_email(email, user_subject, user_html)

    sms = (
        f"Bonjour {first_name},\n"
        f"Je me permets de revenir vers vous concernant votre demande d'inscription en formation {training_label} "
        "depuis votre Compte Personnel de Formation (CPF). Je vous remercie pour votre demande !\n\n"
        "Afin que nous puissions finaliser ensemble votre inscription vous serait-il possible de nous contacter au 04 22 47 07 68 ?\n\n"
        "Vous pouvez également réserver un RDV téléphonique en cliquant ici :\n"
        f"{calendly_url}\n\n"
        "Je vous remercie par avance et je vous souhaite une bonne journée,\n\n"
        "A très bientôt !\n"
        "Clément VAILLANT\n"
        "Directeur Intégrale Academy"
    ).strip()

    sms_ok = brevo_send_sms(phone, sms)

    return jsonify({
        "ok": True,
        "admin_email_ok": bool(admin_email_ok),
        "email_ok": bool(email_ok),
        "sms_ok": bool(sms_ok),
    })


@app.post("/admin/financement-refuse/submit")
@admin_login_required
def admin_financement_refuse_submit():
    payload = request.get_json(silent=True) or {}
    last_name = (payload.get("last_name") or "").strip()
    first_name = (payload.get("first_name") or "").strip()
    phone = (payload.get("phone") or "").strip()
    email = (payload.get("email") or "").strip()
    training_key = (payload.get("training") or "").strip().upper()

    if not all([last_name, first_name, phone, email, training_key]):
        return jsonify({"ok": False, "error": "missing_fields"}), 400

    training_label = FINANCEMENT_REFUSE_TRAININGS.get(training_key)
    if not training_label:
        return jsonify({"ok": False, "error": "invalid_training"}), 400

    full_name = f"{first_name} {last_name}".strip()

    user_subject = f"Financement France Travail refusé – {training_label}"
    user_html = mail_layout(f"""
      <p>Bonjour {first_name},</p>

      <p>Je me permets de revenir vers vous concernant votre formation <strong>{training_label}</strong>.</p>

      <p>Suite à la demande de financement que nous avons envoyée, je suis au regret de vous annoncer que France Travail a refusé de financer votre formation.</p>

      <p>Si vous souhaitez malgré tout vous inscrire en formation, nous vous remercions de bien vouloir nous contacter au <strong>04 22 47 07 68</strong>.</p>

      <p>Nous pouvons vous proposer un paiement en plusieurs fois par prélèvement tous les mois, jusqu'à la date d'examen.</p>

      <p>Nous restons à votre disposition pour tous renseignements complémentaires et nous vous souhaitons une bonne journée.</p>

      <p>La Team Intégrale Academy</p>
    """)

    sms = (
        f"Bonjour {first_name}, je me permets de revenir vers vous concernant votre formation {training_label}. "
        "Suite à la demande de financement que nous avons envoyée, je suis au regret de vous annoncer que France Travail a refusé de financer votre formation. "
        "Si vous souhaitez malgré tout vous inscrire en formation, merci de nous contacter au 04 22 47 07 68. "
        "Nous pouvons vous proposer un paiement en plusieurs fois par prélèvement tous les mois, jusqu'à la date d'examen. "
        "Nous restons à votre disposition pour tous renseignements complémentaires et nous vous souhaitons une bonne journée. "
        "La Team Intégrale Academy"
    ).strip()

    email_ok = brevo_send_email(email, user_subject, user_html)
    sms_ok = brevo_send_sms(phone, sms)

    data = load_data()
    add_notification(
        data,
        "notifications_financement_refuse",
        f"{full_name} • {training_label} • {phone} • {email}",
        meta={
            "first_name": first_name,
            "last_name": last_name,
            "training": training_label,
            "phone": phone,
            "email": email,
        },
    )
    save_data(data)

    return jsonify({
        "ok": True,
        "email_ok": bool(email_ok),
        "sms_ok": bool(sms_ok),
    })


@app.post("/api/secretariat/notifications/<bucket>/<notification_id>/toggle")
@admin_login_required
def api_secretariat_notification_toggle(bucket: str, notification_id: str):
    bucket_key = _notifications_bucket_key(bucket)
    if not bucket_key:
        return jsonify({"ok": False, "error": "invalid_bucket"}), 400

    data = load_data()
    notifications = data.get(bucket_key, [])
    entry = next((item for item in notifications if item.get("id") == notification_id), None)
    if not entry:
        return jsonify({"ok": False, "error": "not_found"}), 404

    entry["done"] = not bool(entry.get("done"))
    entry["done_at"] = _now_iso() if entry["done"] else ""
    save_data(data)

    payload = _secretariat_notifications_payload(data)
    return jsonify({"ok": True, "done": bool(entry.get("done")), **payload})


@app.post("/api/secretariat/notifications/<bucket>/<notification_id>/delete")
@admin_login_required
def api_secretariat_notification_delete(bucket: str, notification_id: str):
    bucket_key = _notifications_bucket_key(bucket)
    if not bucket_key:
        return jsonify({"ok": False, "error": "invalid_bucket"}), 400

    data = load_data()
    notifications = data.get(bucket_key, [])
    entry = next((item for item in notifications if item.get("id") == notification_id), None)
    if not entry:
        return jsonify({"ok": False, "error": "not_found"}), 404

    if not entry.get("done"):
        return jsonify({"ok": False, "error": "not_done"}), 400

    data[bucket_key] = [item for item in notifications if item.get("id") != notification_id]
    save_data(data)

    payload = _secretariat_notifications_payload(data)
    return jsonify({"ok": True, **payload})


@app.post("/api/secretariat/notifications/prelevements/<notification_id>/new-date")
@admin_login_required
@admin_write_required
def api_secretariat_prelevement_new_date(notification_id: str):
    payload = request.get_json(silent=True) or {}
    new_date = (payload.get("new_date") or "").strip()
    if not new_date:
        return jsonify({"ok": False, "error": "missing_new_date"}), 400

    data = load_data()
    notification = next(
        (item for item in data.get("notifications_prelevements", []) if item.get("id") == notification_id),
        None,
    )
    if not notification:
        return jsonify({"ok": False, "error": "notification_not_found"}), 404

    entry_id = ((notification.get("meta") or {}).get("entry_id") or "").strip()
    if not entry_id:
        return jsonify({"ok": False, "error": "entry_not_found"}), 404

    s, t, req = _find_prelevement_request(data, entry_id)
    if not s or not t or not req:
        return jsonify({"ok": False, "error": "entry_not_found"}), 404

    if req.get("new_date"):
        return jsonify({"ok": False, "error": "already_set"}), 400

    req["status"] = "DONE"
    req["responded_at"] = _now_iso()
    req["new_date"] = new_date
    req["new_date_source"] = "SECRETARIAT_DASHBOARD"

    _send_prelevement_new_date_email(t, s, req, new_date)

    trainee_display_name = _format_trainee_name(t.get("first_name", ""), t.get("last_name", ""))
    add_admin_notification(
        data,
        f"🟢{trainee_display_name} - Nouveau prélèvement proposé le {fr_date(new_date) or new_date}",
        meta={
            "type": "prelevement_new_date",
            "source": "secretariat_dashboard",
            "session_id": s.get("id"),
            "trainee_id": t.get("id"),
            "entry_id": entry_id,
            "comment": (req.get("comment") or "").strip(),
        },
    )

    notification.setdefault("meta", {})["new_date"] = new_date
    notification["done"] = True
    notification["done_at"] = _now_iso()

    s["trainees"] = _session_trainees_list(s)
    s.pop("stagiaires", None)
    save_data(data)

    refreshed_payload = _secretariat_notifications_payload(data)
    return jsonify({"ok": True, "new_date": new_date, **refreshed_payload})


@app.post("/api/secretariat/notifications/prelevements/<notification_id>/call-result")
@admin_login_required
def api_secretariat_prelevement_result(notification_id: str):
    payload = request.get_json(silent=True) or {}
    outcome = (payload.get("outcome") or "").strip().upper()
    comment = (payload.get("comment") or "").strip()
    if outcome not in ("CALLED", "NO_ANSWER"):
        return jsonify({"ok": False, "error": "invalid_outcome"}), 400

    data = load_data()
    notification = next(
        (item for item in data.get("notifications_prelevements", []) if item.get("id") == notification_id),
        None,
    )
    if not notification:
        return jsonify({"ok": False, "error": "notification_not_found"}), 404

    notification_meta = notification.setdefault("meta", {})
    previous_no_answer = _parse_no_answer_count(notification_meta.get("no_answer_count"))
    if outcome == "NO_ANSWER":
        no_answer_count = min(3, previous_no_answer + 1)
        display = {
            1: "1er appel pas de réponse",
            2: "2ème appel pas de réponse",
            3: "3ème appel pas de réponse",
        }[no_answer_count]
        notification["done"] = no_answer_count >= 3
        notification["done_at"] = _now_iso() if notification.get("done") else ""
    else:
        no_answer_count = 0
        display = "Personne jointe"
        notification["done"] = True
        notification["done_at"] = _now_iso()

    notification_meta["call_status"] = display
    notification_meta["no_answer_count"] = no_answer_count
    if comment:
        notification_meta["last_comment"] = comment

    trainee_display_name = _format_trainee_name(
        notification_meta.get("first_name", ""),
        notification_meta.get("last_name", ""),
    )
    call_icon = "🟢" if outcome == "CALLED" else ({1: "🟡", 2: "🟠", 3: "🔴"}.get(no_answer_count, "🟡"))
    call_label = (
        f"{call_icon}Prélèvement rejeté {trainee_display_name} - personne appelée"
        if outcome == "CALLED"
        else f"{call_icon}Prélèvement rejeté {trainee_display_name} - {display}"
    )
    add_admin_notification(
        data,
        call_label,
        meta={
            "type": "prelevement_call_result",
            "outcome": outcome,
            "no_answer_count": no_answer_count,
            "session_id": notification_meta.get("session_id"),
            "trainee_id": notification_meta.get("trainee_id"),
            "comment": comment,
            "call_status": display,
        },
    )

    save_data(data)
    refreshed_payload = _secretariat_notifications_payload(data)
    return jsonify({
        "ok": True,
        "done": bool(notification.get("done")),
        "call_status": display,
        "no_answer_count": no_answer_count,
        **refreshed_payload,
    })


@app.post("/api/secretariat/notifications/relances/<notification_id>/call-result")
@admin_login_required
def api_secretariat_relance_result(notification_id: str):
    payload = request.get_json(silent=True) or {}
    outcome = (payload.get("outcome") or "").strip().upper()
    comment = (payload.get("comment") or "").strip()
    if outcome not in ("CALLED", "NO_ANSWER"):
        return jsonify({"ok": False, "error": "invalid_outcome"}), 400

    data = load_data()
    notification = next(
        (item for item in data.get("notifications_phone_relances", []) if item.get("id") == notification_id),
        None,
    )
    if not notification:
        return jsonify({"ok": False, "error": "notification_not_found"}), 404

    followup_id = ((notification.get("meta") or {}).get("followup_id") or "").strip()
    if not followup_id:
        return jsonify({"ok": False, "error": "followup_not_found"}), 404

    s, t, entry = _find_phone_followup_entry(data, followup_id)
    if not s or not t or not entry:
        return jsonify({"ok": False, "error": "followup_not_found"}), 404

    previous_no_answer = _parse_no_answer_count(entry.get("no_answer_count"))
    if outcome == "NO_ANSWER":
        no_answer_count = min(3, previous_no_answer + 1)
        detail = "❌ Pas pu joindre"
        display = {
            1: "1er appel pas de réponse",
            2: "2ème appel pas de réponse",
            3: "3ème appel pas de réponse",
        }[no_answer_count]
    else:
        no_answer_count = 0
        detail = "✅ Appelé"
        display = "Personne jointe"

    t.setdefault("phone_followups", [])
    t["phone_followups"].insert(0, {
        "id": "PHN-REP-" + uuid.uuid4().hex[:8].upper(),
        "type": "RÉPONSE SECRÉTAIRE",
        "at": _now_iso(),
        "details": detail,
        "comment": comment,
        "ref": entry.get("id", ""),
    })

    entry["status"] = "DONE" if outcome == "CALLED" else "PENDING"
    entry["done_at"] = _now_iso() if outcome == "CALLED" else ""
    entry["done_outcome"] = outcome
    entry["no_answer_count"] = no_answer_count

    notification_meta = notification.setdefault("meta", {})
    notification_meta["call_status"] = display
    notification_meta["no_answer_count"] = no_answer_count
    if comment:
        notification_meta["last_comment"] = comment

    trainee_display_name = _format_trainee_name(
        notification_meta.get("first_name", ""),
        notification_meta.get("last_name", ""),
    )
    call_icon = "🟢" if outcome == "CALLED" else ({1: "🟡", 2: "🟠", 3: "🔴"}.get(no_answer_count, "🟡"))
    call_label = (
        f"{call_icon}Relance téléphonique {trainee_display_name} a été appelé"
        if outcome == "CALLED"
        else f"{call_icon}Relance téléphonique {trainee_display_name} {display}"
    )
    add_admin_notification(
        data,
        call_label,
        meta={
            "type": "relance_call_result",
            "outcome": outcome,
            "no_answer_count": no_answer_count,
            "session_id": s.get("id") if s else None,
            "trainee_id": t.get("id") if t else None,
            "comment": comment,
            "call_status": display,
        },
    )

    if outcome == "CALLED":
        notification["done"] = True
        notification["done_at"] = _now_iso()
    else:
        notification["done"] = no_answer_count >= 3
        notification["done_at"] = _now_iso() if notification["done"] else ""

    s["trainees"] = _session_trainees_list(s)
    s.pop("stagiaires", None)
    save_data(data)

    refreshed_payload = _secretariat_notifications_payload(data)
    return jsonify({
        "ok": True,
        "done": bool(notification.get("done")),
        "call_status": display,
        "no_answer_count": no_answer_count,
        **refreshed_payload,
    })


@app.post("/api/secretariat/notifications/cnaps_pre/<notification_id>/call-result")
@admin_login_required
def api_secretariat_cnaps_pre_result(notification_id: str):
    payload = request.get_json(silent=True) or {}
    outcome = (payload.get("outcome") or "").strip().upper()
    comment = (payload.get("comment") or "").strip()
    if outcome not in ("CALLED", "NO_ANSWER"):
        return jsonify({"ok": False, "error": "invalid_outcome"}), 400

    data = load_data()
    notification = next(
        (item for item in data.get("notifications_cnaps_pre_relances", []) if item.get("id") == notification_id),
        None,
    )
    if not notification:
        return jsonify({"ok": False, "error": "notification_not_found"}), 404

    notification_meta = notification.setdefault("meta", {})
    previous_no_answer = _parse_no_answer_count(notification_meta.get("no_answer_count"))
    if outcome == "NO_ANSWER":
        no_answer_count = min(3, previous_no_answer + 1)
        display = {
            1: "1er appel pas de réponse",
            2: "2ème appel pas de réponse",
            3: "3ème appel pas de réponse",
        }[no_answer_count]
        notification["done"] = no_answer_count >= 3
        notification["done_at"] = _now_iso() if notification.get("done") else ""
    else:
        no_answer_count = 0
        display = "Personne jointe"
        notification["done"] = True
        notification["done_at"] = _now_iso()

    notification_meta["call_status"] = display
    notification_meta["no_answer_count"] = no_answer_count
    if comment:
        notification_meta["last_comment"] = comment

    trainee_display_name = _format_trainee_name(
        notification_meta.get("first_name", ""),
        notification_meta.get("last_name", ""),
    )
    call_icon = "🟢" if outcome == "CALLED" else ({1: "🟡", 2: "🟠", 3: "🔴"}.get(no_answer_count, "🟡"))
    call_label = (
        f"{call_icon}Relance PRE CNAPS {trainee_display_name} - personne appelée"
        if outcome == "CALLED"
        else f"{call_icon}Relance PRE CNAPS {trainee_display_name} - {display}"
    )
    add_admin_notification(
        data,
        call_label,
        meta={
            "type": "cnaps_pre_call_result",
            "outcome": outcome,
            "no_answer_count": no_answer_count,
            "session_id": notification_meta.get("session_id"),
            "trainee_id": notification_meta.get("trainee_id"),
            "comment": comment,
            "call_status": display,
        },
    )

    save_data(data)
    refreshed_payload = _secretariat_notifications_payload(data)
    return jsonify({
        "ok": True,
        "done": bool(notification.get("done")),
        "call_status": display,
        "no_answer_count": no_answer_count,
        **refreshed_payload,
    })


@app.post("/api/secretariat/notifications/financement_refuse/<notification_id>/call-result")
@admin_login_required
def api_secretariat_financement_refuse_result(notification_id: str):
    payload = request.get_json(silent=True) or {}
    outcome = (payload.get("outcome") or "").strip().upper()
    comment = (payload.get("comment") or "").strip()
    if outcome not in ("CALLED", "NO_ANSWER"):
        return jsonify({"ok": False, "error": "invalid_outcome"}), 400

    data = load_data()
    notification = next(
        (item for item in data.get("notifications_financement_refuse", []) if item.get("id") == notification_id),
        None,
    )
    if not notification:
        return jsonify({"ok": False, "error": "notification_not_found"}), 404

    notification_meta = notification.setdefault("meta", {})
    previous_no_answer = _parse_no_answer_count(notification_meta.get("no_answer_count"))
    if outcome == "NO_ANSWER":
        no_answer_count = min(3, previous_no_answer + 1)
        display = {
            1: "1er appel pas de réponse",
            2: "2ème appel pas de réponse",
            3: "3ème appel pas de réponse",
        }[no_answer_count]
        notification["done"] = no_answer_count >= 3
        notification["done_at"] = _now_iso() if notification.get("done") else ""
    else:
        no_answer_count = 0
        display = "Personne jointe"
        notification["done"] = True
        notification["done_at"] = _now_iso()

    notification_meta["call_status"] = display
    notification_meta["no_answer_count"] = no_answer_count
    if comment:
        notification_meta["last_comment"] = comment

    trainee_display_name = _format_trainee_name(
        notification_meta.get("first_name", ""),
        notification_meta.get("last_name", ""),
    )
    call_icon = "🟢" if outcome == "CALLED" else ({1: "🟡", 2: "🟠", 3: "🔴"}.get(no_answer_count, "🟡"))
    call_label = (
        f"{call_icon}Financement refusé France Travail {trainee_display_name} a été appelé"
        if outcome == "CALLED"
        else f"{call_icon}Financement refusé France Travail {trainee_display_name} {display}"
    )
    add_admin_notification(
        data,
        call_label,
        meta={
            "type": "financement_refuse_call_result",
            "outcome": outcome,
            "no_answer_count": no_answer_count,
            "session_id": notification_meta.get("session_id"),
            "trainee_id": notification_meta.get("trainee_id"),
            "comment": comment,
            "call_status": display,
        },
    )

    save_data(data)
    refreshed_payload = _secretariat_notifications_payload(data)
    return jsonify({
        "ok": True,
        "done": bool(notification.get("done")),
        "call_status": display,
        "no_answer_count": no_answer_count,
        **refreshed_payload,
    })


@app.post("/api/secretariat/notifications/edof/<notification_id>/call-result")
@admin_login_required
def api_secretariat_edof_result(notification_id: str):
    payload = request.get_json(silent=True) or {}
    outcome = (payload.get("outcome") or "").strip().upper()
    comment = (payload.get("comment") or "").strip()
    if outcome not in ("CALLED", "NO_ANSWER"):
        return jsonify({"ok": False, "error": "invalid_outcome"}), 400

    data = load_data()
    notification = next(
        (item for item in data.get("notifications_edof", []) if item.get("id") == notification_id),
        None,
    )
    if not notification:
        return jsonify({"ok": False, "error": "notification_not_found"}), 404

    notification_meta = notification.setdefault("meta", {})
    previous_no_answer = _parse_no_answer_count(notification_meta.get("no_answer_count"))
    if outcome == "NO_ANSWER":
        no_answer_count = min(3, previous_no_answer + 1)
        display = {
            1: "1er appel pas de réponse",
            2: "2ème appel pas de réponse",
            3: "3ème appel pas de réponse",
        }[no_answer_count]
        notification["done"] = no_answer_count >= 3
        notification["done_at"] = _now_iso() if notification.get("done") else ""
    else:
        no_answer_count = 0
        display = "Personne jointe"
        notification["done"] = True
        notification["done_at"] = _now_iso()

    notification_meta["call_status"] = display
    notification_meta["no_answer_count"] = no_answer_count
    if comment:
        notification_meta["last_comment"] = comment

    trainee_display_name = _format_trainee_name(
        notification_meta.get("first_name", ""),
        notification_meta.get("last_name", ""),
    )
    call_icon = "🟢" if outcome == "CALLED" else ({1: "🟡", 2: "🟠", 3: "🔴"}.get(no_answer_count, "🟡"))
    call_label = (
        f"{call_icon}Pré-inscription CPF {trainee_display_name} a été appelé"
        if outcome == "CALLED"
        else f"{call_icon}Pré-inscription CPF {trainee_display_name} {display}"
    )
    add_admin_notification(
        data,
        call_label,
        meta={
            "type": "edof_call_result",
            "outcome": outcome,
            "no_answer_count": no_answer_count,
            "session_id": notification_meta.get("session_id"),
            "trainee_id": notification_meta.get("trainee_id"),
            "comment": comment,
            "call_status": display,
        },
    )

    save_data(data)
    refreshed_payload = _secretariat_notifications_payload(data)
    return jsonify({
        "ok": True,
        "done": bool(notification.get("done")),
        "call_status": display,
        "no_answer_count": no_answer_count,
        **refreshed_payload,
    })


@app.post("/api/secretariat/notifications/test_fr/<notification_id>/call-result")
@admin_login_required
def api_secretariat_test_fr_result(notification_id: str):
    payload = request.get_json(silent=True) or {}
    outcome = (payload.get("outcome") or "").strip().upper()
    comment = (payload.get("comment") or "").strip()
    if outcome not in ("CALLED", "NO_ANSWER"):
        return jsonify({"ok": False, "error": "invalid_outcome"}), 400

    data = load_data()
    notification = next(
        (item for item in data.get("notifications_test_fr", []) if item.get("id") == notification_id),
        None,
    )
    if not notification:
        return jsonify({"ok": False, "error": "notification_not_found"}), 404

    notification_meta = notification.setdefault("meta", {})
    previous_no_answer = _parse_no_answer_count(notification_meta.get("no_answer_count"))
    if outcome == "NO_ANSWER":
        no_answer_count = min(3, previous_no_answer + 1)
        display = {
            1: "1er appel pas de réponse",
            2: "2ème appel pas de réponse",
            3: "3ème appel pas de réponse",
        }[no_answer_count]
        notification["done"] = no_answer_count >= 3
        notification["done_at"] = _now_iso() if notification.get("done") else ""
    else:
        no_answer_count = 0
        display = "Personne jointe"
        notification["done"] = True
        notification["done_at"] = _now_iso()

    notification_meta["call_status"] = display
    notification_meta["no_answer_count"] = no_answer_count
    if comment:
        notification_meta["last_comment"] = comment

    trainee_display_name = _format_trainee_name(
        notification_meta.get("first_name", ""),
        notification_meta.get("last_name", ""),
    )
    call_icon = "🟢" if outcome == "CALLED" else ({1: "🟡", 2: "🟠", 3: "🔴"}.get(no_answer_count, "🟡"))
    call_label = (
        f"{call_icon}Test de français {trainee_display_name} - personne appelée"
        if outcome == "CALLED"
        else f"{call_icon}Test de français {trainee_display_name} - {display}"
    )
    add_admin_notification(
        data,
        call_label,
        meta={
            "type": "test_fr_call_result",
            "outcome": outcome,
            "no_answer_count": no_answer_count,
            "session_id": notification_meta.get("session_id"),
            "trainee_id": notification_meta.get("trainee_id"),
            "comment": comment,
            "call_status": display,
        },
    )

    save_data(data)
    refreshed_payload = _secretariat_notifications_payload(data)
    return jsonify({
        "ok": True,
        "done": bool(notification.get("done")),
        "call_status": display,
        "no_answer_count": no_answer_count,
        **refreshed_payload,
    })


@app.post("/api/secretariat/notifications/convention_unsigned/<notification_id>/call-result")
@admin_login_required
def api_secretariat_convention_unsigned_result(notification_id: str):
    payload = request.get_json(silent=True) or {}
    outcome = (payload.get("outcome") or "").strip().upper()
    comment = (payload.get("comment") or "").strip()
    if outcome not in ("CALLED", "NO_ANSWER"):
        return jsonify({"ok": False, "error": "invalid_outcome"}), 400

    data = load_data()
    notification = next(
        (item for item in data.get("notifications_convention_unsigned", []) if item.get("id") == notification_id),
        None,
    )
    if not notification:
        return jsonify({"ok": False, "error": "notification_not_found"}), 404

    notification_meta = notification.setdefault("meta", {})
    previous_no_answer = _parse_no_answer_count(notification_meta.get("no_answer_count"))
    if outcome == "NO_ANSWER":
        no_answer_count = min(3, previous_no_answer + 1)
        display = {
            1: "1er appel pas de réponse",
            2: "2ème appel pas de réponse",
            3: "3ème appel pas de réponse",
        }[no_answer_count]
        notification["done"] = no_answer_count >= 3
        notification["done_at"] = _now_iso() if notification.get("done") else ""
    else:
        no_answer_count = 0
        display = "Personne jointe"
        notification["done"] = True
        notification["done_at"] = _now_iso()

    notification_meta["call_status"] = display
    notification_meta["no_answer_count"] = no_answer_count
    if comment:
        notification_meta["last_comment"] = comment

    trainee_display_name = _format_trainee_name(
        notification_meta.get("first_name", ""),
        notification_meta.get("last_name", ""),
    )
    call_icon = "🟢" if outcome == "CALLED" else ({1: "🟡", 2: "🟠", 3: "🔴"}.get(no_answer_count, "🟡"))
    call_label = (
        f"{call_icon}Convention non signée {trainee_display_name} - personne appelée"
        if outcome == "CALLED"
        else f"{call_icon}Convention non signée {trainee_display_name} - {display}"
    )
    add_admin_notification(
        data,
        call_label,
        meta={
            "type": "convention_unsigned_call_result",
            "outcome": outcome,
            "no_answer_count": no_answer_count,
            "session_id": notification_meta.get("session_id"),
            "trainee_id": notification_meta.get("trainee_id"),
            "comment": comment,
            "call_status": display,
        },
    )

    save_data(data)
    refreshed_payload = _secretariat_notifications_payload(data)
    return jsonify({
        "ok": True,
        "done": bool(notification.get("done")),
        "call_status": display,
        "no_answer_count": no_answer_count,
        **refreshed_payload,
    })


@app.get("/admin/test-positionnement")
@admin_login_required
def admin_positioning_tests():
    data = load_data()
    entries = list(data.get("positioning_tests", []))
    entries.sort(key=lambda e: e.get("created_at") or "", reverse=True)
    return render_template("admin_positioning_tests.html", entries=entries)

@app.get("/admin/test-positionnement/<test_id>")
@admin_login_required
def admin_positioning_test_detail(test_id: str):
    data = load_data()
    entries = list(data.get("positioning_tests", []))
    entry = next((e for e in entries if e.get("id") == test_id), None)
    if not entry:
        abort(404)
    return render_template(
        "admin_positioning_test_detail.html",
        entry=entry,
        sections=POSITIONING_TEST_SECTIONS,
    )



@app.get("/admin/sessions/<session_id>/trainees")
@admin_login_required
def admin_trainees(session_id: str):
    data = load_data()
    s = find_session(data, session_id)
    if not s:
        abort(404)

    # normalize session view
    session_view = {
        "id": s.get("id"),
        "name": _session_get(s, "name", ""),
        "training_type": _session_get(s, "training_type", ""),
        "date_start": _session_get(s, "date_start", ""),
        "date_end": _session_get(s, "date_end", ""),
        "exam_date": _session_get(s, "exam_date", ""),
        "exam_theory_date": _session_get(s, "exam_theory_date", ""),
        "exam_practice_date": _session_get(s, "exam_practice_date", ""),
        "practice_training_date": _session_get(s, "practice_training_date", ""),
    }

    trainees = _session_trainees_list(s)

    # refresh CNAPS (best-effort) using last_name/first_name
    for t in trainees:
        ln = normalize_last_name(t.get("last_name") or "")
        fn = normalize_first_name(t.get("first_name") or "")

        if ln:
            t["last_name"] = ln
        if fn:
            t["first_name"] = fn

        current_cnaps = t.get("cnaps") or ""

        # ✅ si déjà validé manuellement, on ne touche pas
        if _normalize_cnaps_status(current_cnaps) == "CARTE PROFESSIONNELLE OK":
            pass
        else:
            if ln and fn:
                cn = fetch_cnaps_status_by_name(ln, fn)

                # ✅ n'écrase jamais avec INCONNU
                if cn:
                    cn_u = str(cn).strip().upper()
                    if cn_u not in ("INCONNU", "UNKNOWN", ""):
                        if _normalize_cnaps_status(cn_u) != _normalize_cnaps_status(current_cnaps):
                            t["cnaps"] = cn_u
                            record_cnaps_status_change(t, cn_u)

        # valeur par défaut si vide
        if not (t.get("cnaps") or "").strip():
            t["cnaps"] = "INCONNU"
            record_cnaps_status_change(t, t["cnaps"])

        # hosting only for A3P
        if session_view["training_type"] == "A3P":
            email = (t.get("email") or "").strip().lower()
        
            hb = fetch_hebergement_status(email) if email else None
        
            # ✅ règle anti-bug : on ne downgrade JAMAIS "reserved"
            current = (t.get("hosting_status") or "unknown").strip().lower()
        
            if hb == "reserved":
                t["hosting_status"] = "reserved"
            elif current == "reserved":
                # on garde reserved quoi qu'il arrive
                t["hosting_status"] = "reserved"
            else:
                # sinon on garde l'ancien si on n'a pas mieux
                t["hosting_status"] = current if current else "unknown"
        else:
            t.pop("hosting_status", None)


    # persist normalized trainees back into storage
    s["trainees"] = trainees
    s.pop("stagiaires", None)
    save_data(data)
    stats = compute_stats(s)
    show_hosting = (session_view["training_type"] == "A3P")
    show_vae = (session_view["training_type"] == "DIRIGEANT VAE")
    is_vtc = ("VTC" in (session_view["training_type"] or "").upper())

    # ✅ docs fin de formation par stagiaire (pour surlignage + n/3 + étiquettes)
    for t in trainees:
        t.setdefault("deliverables", {})
        d_done, d_total, d_ok = deliverables_progress(t)
        t["deliverables_done"] = d_done
        t["deliverables_total"] = d_total
        t["deliverables_ok"] = d_ok
        t["deliverables_text"] = f"{d_done}/{d_total}"

        # ✅ étiquettes ligne (admin_trainees.html)
        dv = t.get("deliverables") or {}
        t["has_sst"] = bool((dv.get("carte_sst") or "").strip())
        t["has_attestation"] = bool((dv.get("attestation_fin_formation") or "").strip())
        t["has_diplome"] = bool((dv.get("diplome") or "").strip())

        badges = []
        if t["has_sst"]:
            badges.append("SST")
        if t["has_attestation"]:
            badges.append("ATTESTATIONS")
        if t["has_diplome"]:
            badges.append("DIPLÔME")
        t["badges"] = badges




    return render_template(
        "admin_trainees.html",
        session=session_view,
        trainees=trainees,
        stats=stats,
        show_hosting=show_hosting,
        show_vae=show_vae,
        is_vtc=is_vtc,
        enums=ENUMS,
    )


# =========================
# FICHE STAGIAIRE (HTML)
# =========================



# =========================
# API - Sessions (used by your modal JS)
# =========================

@app.post("/api/sessions/create")
@admin_login_required
@admin_write_required
def api_create_session():
    data = load_data()
    payload = request.get_json(silent=True) or {}

    name = (payload.get("name") or "").strip()
    training_type = (payload.get("training_type") or "").strip()
    date_start = (payload.get("date_start") or "").strip()
    date_end = (payload.get("date_end") or "").strip()
    exam_date = (payload.get("exam_date") or "").strip()
    exam_theory_date = (payload.get("exam_theory_date") or "").strip()
    exam_practice_date = (payload.get("exam_practice_date") or "").strip()
    practice_training_date = (payload.get("practice_training_date") or "").strip()

    if not name or not training_type:
        return jsonify({"ok": False, "error": "missing_name_or_training_type"}), 400

    session_id = uuid.uuid4().hex[:10]
    s = {
        "id": session_id,
        "name": name,
        "training_type": training_type,
        "date_start": date_start,
        "date_end": date_end,
        "exam_date": exam_date,
        "exam_theory_date": exam_theory_date,
        "exam_practice_date": exam_practice_date,
        "practice_training_date": practice_training_date,
        "created_at": _now_iso(),
        "trainees": [],
        "archived": False, 
    }
    data["sessions"].insert(0, s)
    save_data(data)
    return jsonify({"ok": True, "id": session_id})


@app.post("/api/sessions/<session_id>/update")
@admin_login_required
@admin_write_required
def api_update_session(session_id: str):
    data = load_data()
    s = find_session(data, session_id)
    if not s:
        return jsonify({"ok": False, "error": "session_not_found"}), 404

    payload = request.get_json(silent=True) or {}
    for key in (
        "date_start",
        "date_end",
        "exam_date",
        "exam_theory_date",
        "exam_practice_date",
        "practice_training_date",
    ):
        if key in payload:
            s[key] = (payload.get(key) or "").strip()

    s["updated_at"] = _now_iso()
    save_data(data)
    return jsonify({"ok": True})


@app.post("/api/sessions/<session_id>/delete")
@admin_login_required
@admin_write_required
def api_delete_session(session_id: str):
    data = load_data()
    _force_backup_snapshot(DATA_FILE)
    before = len(data.get("sessions", []))
    data["sessions"] = [s for s in data.get("sessions", []) if s.get("id") != session_id]
    save_data(data)
    return jsonify({"ok": True, "deleted": (len(data["sessions"]) != before)})

@app.post("/api/sessions/<session_id>/archive")
@admin_login_required
@admin_write_required
def api_archive_session(session_id: str):
    data = load_data()
    s = find_session(data, session_id)
    if not s:
        return jsonify({"ok": False, "error": "session_not_found"}), 404

    s["archived"] = True
    s["archived_at"] = _now_iso()
    save_data(data)
    return jsonify({"ok": True})


@app.post("/api/sessions/<session_id>/unarchive")
@admin_login_required
@admin_write_required
def api_unarchive_session(session_id: str):
    data = load_data()
    s = find_session(data, session_id)
    if not s:
        return jsonify({"ok": False, "error": "session_not_found"}), 404

    s["archived"] = False
    # optionnel: garder archived_at pour l'historique, ou le vider
    # s["archived_at"] = ""
    save_data(data)
    return jsonify({"ok": True})


# =========================
# API - Trainees (create + update for autosave)
# =========================

@app.post("/api/sessions/<session_id>/trainees/create")
@admin_login_required
@admin_write_required
def api_create_trainee(session_id: str):
    data = load_data()
    s = find_session(data, session_id)
    if not s:
        return jsonify({"ok": False, "error": "session_not_found"}), 404

    payload = request.get_json(silent=True) or {}
    last_name = normalize_last_name(payload.get("last_name") or "")
    first_name = normalize_first_name(payload.get("first_name") or "")
    birth_date = (payload.get("birth_date") or "").strip()
    birth_city = (payload.get("birth_city") or "").strip()
    email = (payload.get("email") or "").strip()
    phone = (payload.get("phone") or "").strip()
    address = (payload.get("address") or "").strip()
    zip_code = (payload.get("zip_code") or "").strip()
    city = (payload.get("city") or "").strip()
    cpf_amount = (payload.get("cpf_amount") or "").strip()
    personal_amount = (payload.get("personal_amount") or "").strip()
    other_amount = (payload.get("other_amount") or "").strip()
    carte_pro_ok = bool(payload.get("carte_pro_ok"))

    # ✅ nouveau : choisir si on envoie l'accès tout de suite
    send_access = payload.get("send_access", True)
    send_access = True if send_access in (True, "true", "1", 1, "yes", "on") else False

    if not last_name or not first_name:
        return jsonify({"ok": False, "error": "missing_name"}), 400

    trainee_id = "TRN-" + uuid.uuid4().hex[:8].upper()

    training_type = _session_get(s, "training_type", "")
    show_hosting = (training_type == "A3P")
    show_vae = (training_type == "DIRIGEANT VAE")
    default_price = default_training_price(training_type)

    public_token = uuid.uuid4().hex

    t = {
        "id": trainee_id,
        "personal_id": trainee_id,
        "last_name": last_name,
        "first_name": first_name,
        "birth_date": birth_date,
        "birth_city": birth_city,
        "email": email,
        "phone": phone,
        "vtc_cmar_id": "",
        "address": address,
        "zip_code": zip_code,
        "city": city,
        "comment": "",
        "cnaps": "CARTE PROFESSIONNELLE OK" if carte_pro_ok else "INCONNU",
        "convention_status": "soon",
        "test_fr_status": "soon",
        "dossier_status": "incomplete",
        "financement_status": "soon",
        "training_price": default_price if default_price is not None else "",
        "cpf_amount": cpf_amount,
        "personal_amount": personal_amount,
        "other_amount": other_amount,
        "vae_status": "soon" if show_vae else "",
        "hosting_status": "unknown" if show_hosting else "",
        "public_token": public_token,
        "no_permis": False,
        "force_dossier_complete": False,
        "vtc_cm_login": "",
        "vtc_cm_password": "",
        "vtc_cm_submitted_at": "",
        "vtc_cm_reminder_sent_at": "",
        "vtc_cm_reminder_email_ok": False,
        "vtc_cm_reminder_sms_ok": False,
        "vtc_cm_reminder_copy_email_ok": False,
        "exam_fees_paid": False,
        "exam_fees_paid_at": "",
        "elearning_link": "",
        "elearning_link_sent_at": "",
        "elearning_link_email_ok": False,
        "elearning_link_sms_ok": False,
        "vtc_book_sent_at": "",
        "documents": [],
        "created_at": _now_iso(),
        "phone_followups": [],
        "public_hide_popup": False,
    }

    ensure_documents_schema_for_trainee(t, training_type)
    t["dossier_status"] = "complete" if dossier_is_complete_total(t, training_type) else "incomplete"

    trainees = _session_trainees_list(s)
    trainees.insert(0, t)
    s["trainees"] = trainees
    s.pop("stagiaires", None)
    save_data(data)

    # ✅ ENVOI MAIL + SMS à la création (optionnel)
    link = f"{PUBLIC_STUDENT_PORTAL_BASE.rstrip('/')}/espace/{public_token}"

    if send_access:
        training_type = _session_get(s, "training_type", "")
        if "VTC" in (training_type or "").upper():
            subject, html = build_vtc_onboarding_email(first_name, link)
            sms = build_vtc_onboarding_sms(first_name, link)
            email_ok = brevo_send_email(email, subject, html) if email else False
            sms_ok = brevo_send_sms(phone, sms) if phone else False
        elif (training_type or "") == "DIRIGEANT VAE":
            subject = "Votre VAE Dirigeant d'entreprise de sécurité privée (DESP)"
            html = mail_layout(f"""
              <style>
                .vae-mail {{
                  font-size: 16px;
                  line-height: 1.65;
                  color: #0f172a;
                }}
                .vae-mail p {{
                  margin: 0 0 14px 0;
                }}
                .vae-mail .step {{
                  background: #f8fafc;
                  border: 1px solid #e2e8f0;
                  border-radius: 12px;
                  padding: 14px;
                  margin: 0 0 12px 0;
                }}
                .vae-mail .step-title {{
                  font-weight: 700;
                  margin-bottom: 6px;
                  display: block;
                }}
                @keyframes vaeCtaPulse {{
                  0% {{ box-shadow: 0 0 0 0 rgba(31, 143, 74, 0.45); transform: scale(1); }}
                  70% {{ box-shadow: 0 0 0 12px rgba(31, 143, 74, 0); transform: scale(1.02); }}
                  100% {{ box-shadow: 0 0 0 0 rgba(31, 143, 74, 0); transform: scale(1); }}
                }}
              </style>
              <div class="vae-mail">
              <h2 style="text-align:center;margin:0 0 16px 0">🚀 En route vers la VAE</h2>
              <p>Bonjour {first_name},</p>

              <p>
                Votre VAE (Validation des acquis de l'expérience) Dirigeant d'entreprise de sécurité privée (DESP)
                commence aujourd'hui.
              </p>

              <p><strong>Les étapes :</strong></p>

              <div class="step"><span class="step-title">1️⃣ Rédaction du Livret 1 (dossier de faisabilité)</span>
              Vous allez compléter en ligne votre dossier de faisabilité depuis votre Espace candidat.<br>
              Ce document permet de présenter votre parcours professionnel, vos fonctions exercées et vos responsabilités,
              afin de vérifier que votre expérience correspond bien aux compétences attendues pour le DESP.
              C’est en quelque sorte la « photographie » de votre expérience.<br>
              ⏳ Durée estimée : environ 30 minutes.</div>

              <div class="step"><span class="step-title">2️⃣ Étude du Livret 1 et attestation de recevabilité</span>
              Votre dossier est étudié par la commission.<br>
              Si les éléments fournis sont conformes et suffisants, une attestation de recevabilité vous est délivrée.<br>
              À partir de ce moment, nous prendrons contact avec vous pour mettre en place la convention de VAE
              et procéder au règlement de l’acompte (1 140 €).</div>

              <div class="step"><span class="step-title">3️⃣ Rédaction du Livret 2</span>
              Vous devrez ensuite compléter le Livret 2.<br>
              Ce document est le cœur de votre démarche : vous y détaillez précisément vos activités, vos missions,
              les situations professionnelles rencontrées, ainsi que les compétences mobilisées.<br>
              C’est ce dossier qui sera présenté au jury de certification.</div>

              <div class="step"><span class="step-title">4️⃣ Étude du Livret 2</span>
              La commission analyse votre dossier.<br>
              Si l’ensemble est conforme et complet, une date de passage devant le jury de certification est programmée.</div>

              <div class="step"><span class="step-title">5️⃣ Passage devant le jury de certification</span>
              Vous serez convoqué à un entretien professionnel d’environ une heure.<br>
              Lors de cet échange, le jury reviendra sur votre parcours et sur les éléments présentés dans le Livret 2.<br>
              L’objectif est de vérifier la maîtrise des compétences attendues, à travers des questions concrètes sur
              votre expérience et vos pratiques professionnelles.</div>

              <div class="step"><span class="step-title">6️⃣ Obtention de votre certification</span></div>

              <p style="text-align:center;margin:24px 0">
                <a href="{link}"
                   style="display:inline-block;background:#1f8f4a;color:white;padding:12px 18px;border-radius:10px;text-decoration:none;font-weight:bold;animation:vaeCtaPulse 1.8s ease-out infinite;box-shadow:0 0 0 0 rgba(31, 143, 74, 0.45)">
                  Démarrer ma VAE
                </a>
              </p>

              <p>
                Je reste à votre disposition pour tous renseignements complémentaires,<br>
                <strong>Clément VAILLANT</strong><br>
                Directeur Intégrale Academy
              </p>
              </div>
            """)

            sms = (
                f"Intégrale Academy Bonjour {first_name}, votre VAE Dirigeant d'entreprise de sécurité (DESP) commence aujourd'hui 🚀. "
                f"Pour démarrer votre VAE cliquez ici : {link} "
                "Je reste à votre disposition. Clément VAILLANT - Intégrale Academy"
            )

            email_ok = brevo_send_email(email, subject, html) if email else False
            sms_ok = brevo_send_sms(phone, sms) if phone else False
        else:
            formation_type = formation_label(training_type)
            dstart = fr_date(_session_get(s, "date_start", ""))
            dend = fr_date(_session_get(s, "date_end", ""))

            subject = "Votre inscription en formation – Intégrale Academy"

            html = mail_layout(f"""
              <h2 style="text-align:center">🎉 Confirmation d’inscription</h2>
              <p>Bonjour <strong>{first_name}</strong>,</p>
              <p>
                Je vous confirme que vous êtes inscrit(e) en formation
                <strong>{formation_type}</strong>, qui se déroulera
                du <strong>{dstart}</strong> au <strong>{dend}</strong>.
              </p>
              <p>Je vous remercie pour votre confiance !</p>
              <p>
                Vous recevrez prochainement par mail votre <strong>Contrat de formation</strong>
                que je vous invite à signer dès réception (signature électronique).
              </p>
              <p>
                📂 Je vous remercie de bien vouloir compléter dès que possible votre
                <strong>Dossier Formation</strong> depuis votre Espace Stagiaire en cliquant sur le bouton ci-dessous.
              </p>
              <p style="color:#b91c1c;font-weight:bold">
                ⚠️ Attention : votre dossier doit être complet au plus tard <u>10 jours avant le début de votre formation</u> !
              </p>

              <p style="text-align:center">
                <a href="{link}"
                   style="display:inline-block;background:#1f8f4a;color:white;padding:12px 18px;border-radius:10px;text-decoration:none;font-weight:bold">
                  👉 Accéder à mon espace stagiaire
                </a>
              </p>

              <p style="margin-top:25px">
                ☎️ Pour tous renseignements, vous pouvez nous contacter au <strong>04 22 47 07 68</strong>
                ou utiliser notre formulaire d’assistance :
              </p>

              <p style="text-align:center">
                <a href="https://assistance-alw9.onrender.com/"
                   style="display:inline-block;background:#2563eb;color:white;padding:10px 16px;border-radius:10px;text-decoration:none;font-weight:bold">
                  🛠️ Formulaire d’assistance
                </a>
              </p>

              <p style="margin-top:30px">
                Je reste à votre disposition pour tous renseignements complémentaires,<br>
                <strong>Clément VAILLANT</strong><br>
                Directeur Intégrale Academy
              </p>

              <hr style="margin:30px 0;border:none;border-top:1px solid #e5e7eb">

              <p style="font-size:12px;color:#6b7280;text-align:center;line-height:1.6">
                © Intégrale Academy — Merci de votre confiance 💛<br>
                54 chemin du Carreou 83480 PUGET SUR ARGENS / 142 rue de Rivoli 75001 PARIS<br>
                SIREN 840 899 884 - NDA 93830600283 - Certification Nationale QUALIOPI : n°03169 en date du 21/10/2024<br>
                UAI Côte d'Azur 0831774C - UAI Paris 0756548K<br>
                <a href="https://www.integraleacademy.com" style="color:#1f8f4a;text-decoration:none;font-weight:bold">
                  integraleacademy.com
                </a>
              </p>
            """)

            sms = (
                f"Intégrale Academy 🎓 Bonjour {first_name}, Votre inscription en formation {formation_type} est confirmée. "
                f"({dstart} au {dend}). Vous allez prochainement recevoir par mail votre Contrat de formation (signature électronique). "
                f"Vous devez à présent compléter votre Dossier Formation : {link} "
                f"(votre dossier doit être COMPLET au plus tard 10 jours avant votre entrée en formation). "
                f"Pour toute demande d'assistance vous pouvez nous contacter au 04 22 47 07 68."
            )

            email_ok = brevo_send_email(email, subject, html) if email else False
            sms_ok = brevo_send_sms(phone, sms) if phone else False

        t["access_sent_at"] = _now_iso()
        t["access_sent_email_ok"] = bool(email_ok)
        t["access_sent_sms_ok"] = bool(sms_ok)
    else:
        # pas d'envoi maintenant
        email_ok = False
        sms_ok = False
        t["access_sent_at"] = ""
        t["access_sent_email_ok"] = False
        t["access_sent_sms_ok"] = False

    save_data(data)

    return jsonify({
        "ok": True,
        "id": trainee_id,
        "access_email_ok": email_ok,
        "access_sms_ok": sms_ok,
        "public_link": link,
        "summary_url": url_for("admin_trainee_summary", session_id=session_id, trainee_id=trainee_id)
    })





@app.post("/api/sessions/<session_id>/stagiaires/<trainee_id>/update")
@admin_login_required
@admin_write_required
def api_update_trainee(session_id: str, trainee_id: str):
    data = load_data()
    s = find_session(data, session_id)
    if not s:
        return jsonify({"ok": False, "error": "session_not_found"}), 404

    trainees = _session_trainees_list(s)
    t = None
    for x in trainees:
        if x.get("id") == trainee_id:
            t = x
            break
    if not t:
        return jsonify({"ok": False, "error": "trainee_not_found"}), 404

    payload = request.get_json(silent=True) or {}
    was_exam_fees_paid = bool(t.get("exam_fees_paid"))
    previous_elearning_link = (t.get("elearning_link") or "").strip()
    previous_cnaps_status = (t.get("cnaps") or "").strip()

    # Your template uses:
    # - convention_status, test_fr_status, dossier_status, financement_status, vae_status, comment, cnaps
    allowed = {
        "convention_status",
        "convention_saisie_done",
        "convention_signed_done",
        "test_fr_status",
        "dossier_status",
        "force_dossier_complete",
        "financement_status",
        "vae_status",
        "comment",
        "financement_comment",
        "financement_new_date_seen",
        "vae_status_label",
        "vae_jury_date",
        "vae_action_dates",
        "cnaps",
        "no_permis", 
        "public_hide_infos",
        "public_hide_docs",
        "public_hide_suivi", 
        "public_hide_popup",
        "last_name",
        "first_name",
        "email",
        "phone",
        "training_price",
        "cpf_amount",
        "personal_amount",
        "other_amount",
        "birth_city",
        "address",
        "zip_code",
        "city",
        "vtc_cmar_id",
        "vtc_cm_login",
        "vtc_cm_password",
        "vtc_cm_submitted_at",
        "exam_fees_paid",
        "elearning_link",
        "vtc_book_sent_at",
        "vtc_theory_exam_sent_at",
        "vtc_practice_convocation_sent_at",

    }

    previous_vae_status = vae_status_view(t.get("vae_status"))["key"]

    send_vae_notification = True if payload.get("send_vae_notification", True) in (True, "true", "1", 1, "yes", "on") else False
    send_exam_fees_notification = True if payload.get("send_exam_fees_notification", True) in (True, "true", "1", 1, "yes", "on") else False
    send_elearning_notification = True if payload.get("send_elearning_notification", True) in (True, "true", "1", 1, "yes", "on") else False

    for k, v in payload.items():
        if k in ("send_vae_notification", "send_exam_fees_notification", "send_elearning_notification"):
            continue
        if k not in allowed:
            continue

        # bools
        if k in ("no_permis", "public_hide_infos", "public_hide_docs", "public_hide_suivi", "public_hide_popup", "force_dossier_complete", "financement_new_date_seen", "exam_fees_paid"):
            t[k] = True if v in (True, "true", "1", 1, "yes", "on") else False
            continue

        # strings (statuts / texte)
        if v is None:
            continue

        if isinstance(v, str):
            if k == "cnaps":
                new_val = v.strip()
                if _normalize_cnaps_status(new_val) != _normalize_cnaps_status(t.get(k)):
                    t[k] = new_val
                    record_cnaps_status_change(t, new_val)
                else:
                    t[k] = new_val
                continue
            if k == "last_name":
                t[k] = normalize_last_name(v)
            elif k == "first_name":
                t[k] = normalize_first_name(v)
            elif k == "elearning_link":
                t[k] = v.strip()
            else:
                t[k] = v.strip()
        else:
            t[k] = v

    if "vae_status" in payload or "vae_status_label" in payload:
        requested_vae = (payload.get("vae_status") or payload.get("vae_status_label") or "").strip()
        view = vae_status_view(requested_vae)
        t["vae_status"] = view["key"]
        t["vae_status_label"] = view["label"]
        if view["key"] != previous_vae_status and send_vae_notification:
            _notify_vae_status_change(t, view["key"])

    if (payload.get("financement_status") or "").strip() == "validated":
        t["financement_rejected_note"] = ""
        t["financement_new_date_seen"] = False
        t["comment"] = _remove_admin_comment_flag(t.get("comment", ""), "⚠️ Prélèvement rejeté")

    if "cnaps" in payload:
        new_cnaps_status = (t.get("cnaps") or "").strip()
        if _normalize_cnaps_status(new_cnaps_status) != _normalize_cnaps_status(previous_cnaps_status):
            trainee_display_name = _format_trainee_name(t.get("first_name", ""), t.get("last_name", ""))
            normalized_new = _normalize_cnaps_status(new_cnaps_status)
            icon = "🟠"
            if normalized_new == "ACCEPTE":
                icon = "🟢"
            add_admin_notification(
                data,
                f"{icon}{trainee_display_name} CNAPS {new_cnaps_status.lower()}",
                meta={
                    "type": "cnaps_status_change",
                    "session_id": s.get("id"),
                    "trainee_id": t.get("id"),
                    "old_status": previous_cnaps_status,
                    "new_status": new_cnaps_status,
                    "comment": (t.get("comment") or "").strip(),
                },
            )

    if "exam_fees_paid" in payload:
        now_paid = bool(t.get("exam_fees_paid"))
        if now_paid and not was_exam_fees_paid:
            t["exam_fees_paid_at"] = _now_iso()
            trainee_name = f"{t.get('first_name','').strip()} {t.get('last_name','').strip()}".strip()
            session_name = _session_get(s, "name", "")
            email = (t.get("email") or "").strip()
            phone = (t.get("phone") or "").strip()
            subject = "Frais d'examen réglés"
            message = "Nous venons de payer les frais d'examen, la Chambre des métiers vous enverra prochainement votre convocation."
            html = mail_layout(f"""
              <h2>Frais d'examen réglés</h2>
              <p>{message}</p>
              <p><strong>Stagiaire :</strong> {trainee_name or '—'}</p>
              <p><strong>Session :</strong> {session_name or '—'}</p>
            """)
            if send_exam_fees_notification:
                if email:
                    brevo_send_email(email, subject, html)
                if phone:
                    sms_prefix = f"Bonjour {t.get('first_name','').strip()}, " if (t.get("first_name") or "").strip() else "Bonjour, "
                    sms = f"{sms_prefix}{message}"
                    brevo_send_sms(phone, sms)
        elif not now_paid:
            t["exam_fees_paid_at"] = ""

    elearning_notifications = {"email_ok": False, "sms_ok": False}
    now_elearning_link = (t.get("elearning_link") or "").strip()
    if now_elearning_link and now_elearning_link != previous_elearning_link:
        if send_elearning_notification:
            elearning_notifications = notify_elearning_access_available(t, s, now_elearning_link)
            t["elearning_link_sent_at"] = _now_iso()
            t["elearning_link_email_ok"] = bool(elearning_notifications.get("email_ok"))
            t["elearning_link_sms_ok"] = bool(elearning_notifications.get("sms_ok"))
        else:
            t["elearning_link_sent_at"] = ""
            t["elearning_link_email_ok"] = False
            t["elearning_link_sms_ok"] = False
    elif not now_elearning_link:
        t["elearning_link_sent_at"] = ""
        t["elearning_link_email_ok"] = False
        t["elearning_link_sms_ok"] = False

    t["updated_at"] = _now_iso()
    s["trainees"] = trainees
    s.pop("stagiaires", None)
    training_type = _session_get(s, "training_type", "")
    t["dossier_status"] = "complete" if dossier_is_complete_total(t, training_type) else "incomplete"
    save_data(data)
    return jsonify({
        "ok": True,
        "dossier_status": t.get("dossier_status"),
        "force_dossier_complete": bool(t.get("force_dossier_complete")),
        "elearning_link_sent_at": t.get("elearning_link_sent_at") or "",
        "elearning_link_email_ok": bool(t.get("elearning_link_email_ok")),
        "elearning_link_sms_ok": bool(t.get("elearning_link_sms_ok")),
    })


@app.post("/api/sessions/<session_id>/stagiaires/<trainee_id>/pre-reception")
@admin_login_required
@admin_write_required
def api_admin_pre_reception(session_id: str, trainee_id: str):
    data = load_data()
    s = find_session(data, session_id)
    if not s:
        return jsonify({"ok": False, "error": "session_not_found"}), 404

    trainees = _session_trainees_list(s)
    t = next((x for x in trainees if x.get("id") == trainee_id), None)
    if not t:
        return jsonify({"ok": False, "error": "trainee_not_found"}), 404

    payload = request.get_json(silent=True) or {}
    pre_raw = (payload.get("pre_number") or "").strip()
    if not pre_raw:
        return jsonify({"ok": False, "error": "missing_pre"}), 400

    pre = pre_raw.upper().replace(" ", "")
    if not re.match(r"^(PRE|CAR)-(?:\d{3}-)?\d{4}-\d{2}-\d{2}-\d{11,}$", pre):
        return jsonify({"ok": False, "error": "invalid_pre"}), 400

    t["pre_number"] = pre
    training_type = _session_get(s, "training_type", "")
    t["dossier_status"] = "complete" if dossier_is_complete_total(t, training_type) else "incomplete"
    t["updated_at"] = _now_iso()

    s["trainees"] = trainees
    s.pop("stagiaires", None)
    save_data(data)

    email = (t.get("email") or "").strip()
    training_name = formation_label(_session_get(s, "training_type", "") or s.get("name") or "formation")
    dstart = fr_date(s.get("date_start") or "")
    dend = fr_date(s.get("date_end") or "")
    date_phrase = f"qui se déroulera du {dstart} au {dend}." if dstart and dend else "qui se déroulera prochainement."

    html = mail_layout(f"""
        <p>Bonjour,</p>
        <p>
          Nous revenons vers vous concernant votre formation <strong>{training_name}</strong>, {date_phrase}
        </p>
        <p>
          Vous nous avez indiqué ne pas avoir reçu par courrier postal votre autorisation préalable du CNAPS
          (Ministère de l'intérieur), alors que votre demande a bien été validée par leurs services.
        </p>
        <p>
          Afin de pouvoir finaliser votre entrée en formation, nous avons donc pris contact avec le CNAPS pour obtenir
          votre numéro d’autorisation.
          Après vérification, le CNAPS nous a confirmé que votre numéro d’autorisation est le suivant :
          <strong>{pre}</strong>
        </p>
        <p>Nous vous informons que votre dossier en ligne a été automatiquement complété avec ce numéro.</p>
        <p>Nous restons à votre disposition si besoin et vous souhaitons une excellente journée.</p>
        <p>Bien cordialement,<br>Intégrale Academy</p>
    """)

    email_ok = brevo_send_email(email, "Votre numéro d’autorisation CNAPS (PRE)", html) if email else False

    return jsonify({"ok": True, "pre_number": pre, "email_ok": bool(email_ok)})

@app.post("/api/sessions/<session_id>/stagiaires/<trainee_id>/cnaps-pre-relance/send")
@admin_login_required
@admin_write_required
def api_send_cnaps_pre_relance(session_id: str, trainee_id: str):
    data = load_data()
    s = find_session(data, session_id)
    if not s:
        return jsonify({"ok": False, "error": "session_not_found"}), 404

    trainees = _session_trainees_list(s)
    t = next((x for x in trainees if x.get("id") == trainee_id), None)
    if not t:
        return jsonify({"ok": False, "error": "trainee_not_found"}), 404

    training_name = formation_label(_session_get(s, "training_type", "") or s.get("name") or "formation")
    dstart = fr_date(s.get("date_start") or "")
    dend = fr_date(s.get("date_end") or "")
    if dstart and dend:
        date_phrase = f"qui se déroulera du <strong>{dstart}</strong> au <strong>{dend}</strong>."
        date_phrase_sms = f"du {dstart} au {dend}."
    else:
        date_phrase = "qui se déroulera prochainement."
        date_phrase_sms = "qui se déroulera prochainement."

    link = "https://cnapsv5-1.onrender.com/"

    html = mail_layout(f"""
        <p>Bonjour,</p>
        <p>
          Je me permets de revenir vers vous concernant votre formation
          <strong>{training_name}</strong> {date_phrase}
        </p>
        <p>
          À ce jour, nous n’avons toujours pas reçu les documents indispensables pour effectuer
          votre demande d’autorisation préalable d’entrée en formation auprès du CNAPS
          (Ministère de l’Intérieur).
        </p>
        <p>
          Nous vous rappelons que cette autorisation est obligatoire et doit impérativement
          être obtenue avant toute entrée en formation.
        </p>
        <p>
          Pour déposer votre demande, cliquez sur le bouton ci-dessous :
        </p>
        <p style="text-align:center;margin:18px 0;">
          <a href="{link}" style="display:inline-block;background:#1f8f4a;color:#ffffff;text-decoration:none;font-weight:700;padding:12px 18px;border-radius:8px;">
            Déposer ma demande CNAPS
          </a>
        </p>
        <p>
          Pour toute question vous pouvez nous contacter au 04 22 47 07 68.
        </p>
        <p>Je vous remercie par avance,</p>
        <p>
          Clément VAILLANT<br>
          Directeur Intégrale Academy
        </p>
    """)

    sms = (
        "Intégrale Academy : Bonjour, "
        f"je me permets de revenir vers vous concernant votre formation {training_name} {date_phrase_sms} "
        "Nous n'avons pas reçu les documents pour votre demande d'autorisation préalable CNAPS. "
        "Cette autorisation est obligatoire avant l'entrée en formation. "
        f"Déposez votre demande ici : {link}"
    )

    email = (t.get("email") or "").strip()
    phone = (t.get("phone") or "").strip()
    email_ok = brevo_send_email(email, "Relance documents CNAPS Ministère de l'intérieur", html) if email else False
    sms_ok = brevo_send_sms(phone, sms) if phone else False

    t["cnaps_pre_relance_last_sent_at"] = _now_iso()
    t["updated_at"] = _now_iso()

    first_name = (t.get("first_name") or "").strip()
    last_name = (t.get("last_name") or "").strip()
    phone_display = phone or ""
    email_display = email or ""
    label_name = _format_trainee_name(first_name, last_name)
    label_parts = [label_name]
    if training_name:
        label_parts.append(training_name)
    add_notification(
        data,
        "notifications_cnaps_pre_relances",
        " • ".join([part for part in label_parts if part]).strip(),
        meta={
            "first_name": first_name,
            "last_name": last_name,
            "training": training_name,
            "phone": phone_display,
            "email": email_display,
            "session_id": s.get("id"),
            "trainee_id": t.get("id"),
            "call_status": "À appeler",
            "no_answer_count": 0,
        },
    )

    s["trainees"] = trainees
    s.pop("stagiaires", None)
    save_data(data)

    return jsonify({"ok": True, "email_ok": bool(email_ok), "sms_ok": bool(sms_ok)})




@app.post("/api/sessions/<session_id>/stagiaires/<trainee_id>/vtc-theory-exam/send")
@admin_login_required
@admin_write_required
def api_send_vtc_theory_exam(session_id: str, trainee_id: str):
    data = load_data()
    s = find_session(data, session_id)
    if not s:
        return jsonify({"ok": False, "error": "session_not_found"}), 404

    training_type = (_session_get(s, "training_type", "") or "").upper()
    if "VTC" not in training_type:
        return jsonify({"ok": False, "error": "not_vtc_session"}), 400

    trainees = _session_trainees_list(s)
    t = next((x for x in trainees if x.get("id") == trainee_id), None)
    if not t:
        return jsonify({"ok": False, "error": "trainee_not_found"}), 404

    payload = request.get_json(silent=True) or {}
    send_email = payload.get("send_email", True) in (True, "true", "1", 1, "yes", "on")

    result = _send_vtc_theory_exam_notification(s, t, send_email=send_email)

    s["trainees"] = trainees
    s.pop("stagiaires", None)
    save_data(data)

    return jsonify({"ok": True, **result})


@app.post("/api/vtc/check/import")
@admin_login_required
@admin_write_required
def api_vtc_check_import():
    pdf = request.files.get("file")
    if not pdf:
        return jsonify({"ok": False, "error": "missing_file"}), 400

    filename = (pdf.filename or "").lower()
    if not filename.endswith(".pdf"):
        return jsonify({"ok": False, "error": "invalid_file_type"}), 400

    try:
        identifiers = _extract_cmar_identifiers_from_pdf(pdf.read())
    except Exception:
        return jsonify({"ok": False, "error": "pdf_parse_error"}), 400

    if not identifiers:
        return jsonify({"ok": True, "matches": [], "count": 0, "message": "aucun stagiaire VTC trouvé"})

    data = load_data()
    wanted = set(identifiers)
    matches = []
    seen = set()

    for sess in data.get("sessions", []):
        if "VTC" not in (_session_get(sess, "training_type", "") or "").upper():
            continue
        trainees = _session_trainees_list(sess)
        for trainee in trainees:
            cmar_id = _normalize_cmar_identifier(trainee.get("vtc_cmar_id") or "")
            if not cmar_id or cmar_id not in wanted:
                continue
            key = (sess.get("id"), trainee.get("id"))
            if key in seen:
                continue
            seen.add(key)
            matches.append({
                "session_id": sess.get("id") or "",
                "session_name": sess.get("name") or "",
                "trainee_id": trainee.get("id") or "",
                "first_name": (trainee.get("first_name") or "").strip(),
                "last_name": (trainee.get("last_name") or "").strip(),
                "cmar_id": trainee.get("vtc_cmar_id") or "",
            })

    count = len(matches)
    if count == 0:
        message = "aucun stagiaire VTC trouvé"
    elif count == 1:
        message = "1 stagiaire VTC trouvé"
    else:
        message = f"{count} stagiaires VTC trouvés"

    return jsonify({"ok": True, "matches": matches, "count": count, "message": message})


@app.post("/api/vtc/check/notify")
@admin_login_required
@admin_write_required
def api_vtc_check_notify():
    payload = request.get_json(silent=True) or {}
    items = payload.get("items") or []
    if not isinstance(items, list) or not items:
        return jsonify({"ok": False, "error": "missing_items"}), 400

    data = load_data()
    sent = 0
    failed = 0

    for item in items:
        session_id = str(item.get("session_id") or "").strip()
        trainee_id = str(item.get("trainee_id") or "").strip()
        if not session_id or not trainee_id:
            failed += 1
            continue

        sess = find_session(data, session_id)
        if not sess:
            failed += 1
            continue
        trainees = _session_trainees_list(sess)
        trainee = next((x for x in trainees if x.get("id") == trainee_id), None)
        if not trainee:
            failed += 1
            continue

        _send_vtc_theory_exam_notification(sess, trainee, send_email=True)
        sess["trainees"] = trainees
        sess.pop("stagiaires", None)
        sent += 1

    save_data(data)
    return jsonify({"ok": True, "sent": sent, "failed": failed})

@app.post("/api/sessions/<session_id>/trainees/<trainee_id>/delete")
@admin_login_required
@admin_write_required
def api_delete_trainee(session_id: str, trainee_id: str):
    data = load_data()
    _force_backup_snapshot(DATA_FILE)
    s = find_session(data, session_id)
    if not s:
        return jsonify({"ok": False, "error": "session_not_found"}), 404

    trainees = _session_trainees_list(s)
    before = len(trainees)
    trainees = [x for x in trainees if x.get("id") != trainee_id]
    s["trainees"] = trainees
    s.pop("stagiaires", None)
    save_data(data)
    return jsonify({"ok": True, "deleted": (len(trainees) != before)})


# =========================
# CNAPS lookup API (used by your refresh button)
# =========================

@app.get("/api/cnaps_lookup")
def api_cnaps_lookup():
    nom = (request.args.get("nom") or "").strip()
    prenom = (request.args.get("prenom") or "").strip()

    if not nom or not prenom:
        return jsonify({"ok": False, "error": "missing_nom_or_prenom"}), 400

    status = fetch_cnaps_status_by_name(nom, prenom) or "INCONNU"
    return jsonify({"ok": True, "nom": nom, "prenom": prenom, "statut_cnaps": str(status).upper()})


# =========================
# Health
# =========================

@app.get("/api/health")
def health():
    backup_files = []
    try:
        backup_files = [n for n in os.listdir(BACKUP_DIR) if n.endswith(".json")]
    except Exception:
        backup_files = []
    return jsonify({
        "ok": True,
        "data_file": DATA_FILE,
        "vae_data_file": VAE_DATA_FILE,
        "backup_dir": BACKUP_DIR,
        "backups_count": len(backup_files),
        "backup_retention": BACKUP_RETENTION,
    })


@app.post("/api/test-positionnement/submit")
def api_positioning_test_submit():
    payload = request.get_json(silent=True) or {}
    contact = payload.get("contact") or {}
    answers = payload.get("answers") or {}

    contact_data = {
        "last_name": (contact.get("last_name") or "").strip(),
        "first_name": (contact.get("first_name") or "").strip(),
        "email": (contact.get("email") or "").strip(),
        "phone": (contact.get("phone") or "").strip(),
    }
    if not all(contact_data.values()):
        return jsonify({"ok": False, "error": "missing_contact_fields"}), 400

    email_normalized = contact_data["email"].lower()

    score_data = positioning_test_score(answers)

    data = load_data()
    existing_entries = data.get("positioning_tests", [])
    already_submitted = any(
        (entry.get("contact") or {}).get("email", "").strip().lower() == email_normalized
        for entry in existing_entries
    )
    if already_submitted:
        return jsonify({"ok": False, "error": "email_already_submitted"}), 409

    entry = {
        "id": uuid.uuid4().hex,
        "created_at": _now_iso(),
        "contact": contact_data,
        "answers": answers,
        "score": score_data["score"],
        "total": score_data["total"],
        "score_over_20": score_data["score_over_20"],
    }

    data.setdefault("positioning_tests", []).append(entry)
    save_data(data)

    return jsonify(
        {
            "ok": True,
            "score": score_data["score"],
            "total": score_data["total"],
            "score_over_20": score_data["score_over_20"],
        }
    )


@app.post("/api/test-positionnement/<test_id>/delete")
@admin_login_required
@admin_write_required
def api_positioning_test_delete(test_id: str):
    data = load_data()
    _force_backup_snapshot(DATA_FILE)
    entries = list(data.get("positioning_tests", []))
    new_entries = [e for e in entries if e.get("id") != test_id]
    data["positioning_tests"] = new_entries
    save_data(data)
    return jsonify({"ok": True, "deleted": len(entries) - len(new_entries)})


@app.post("/api/test-positionnement/delete_all")
@admin_login_required
@admin_write_required
def api_positioning_test_delete_all():
    data = load_data()
    _force_backup_snapshot(DATA_FILE)
    deleted = len(data.get("positioning_tests", []))
    data["positioning_tests"] = []
    save_data(data)
    return jsonify({"ok": True, "deleted": deleted})

from werkzeug.utils import secure_filename



# =========================
# Upload helpers
# =========================
ALLOWED_EXT = {".pdf",".png",".jpg",".jpeg",".doc",".docx",".webp"}

def _safe_ext(filename: str) -> str:
    return os.path.splitext(filename)[1].lower()

def _store_file(session_id: str, trainee_id: str, folder: str, f) -> str:
    base = trainee_upload_dir(session_id, trainee_id)
    target_dir = os.path.join(base, folder)
    os.makedirs(target_dir, exist_ok=True)

    filename = secure_filename(f.filename or "file")
    ext = _safe_ext(filename)
    if ext and ext not in ALLOWED_EXT:
        raise ValueError("extension_not_allowed")

    name = uuid.uuid4().hex[:10] + (ext or "")
    path = os.path.join(target_dir, name)
    f.save(path)
    return path

def _tokenize_path(path: str) -> str:
    # on ne renvoie pas le chemin réel au template
    # token = path relatif à PERSIST_DIR
    rel = os.path.relpath(path, PERSIST_DIR).replace("\\","/")
    return rel

def _detokenize_path(token: str) -> str:
    token = (token or "").replace("..","").lstrip("/").replace("\\","/")
    return os.path.join(PERSIST_DIR, token)

@app.get("/admin/uploads/<path:path>")
@admin_login_required
def admin_view_upload(path: str):
    full = _detokenize_path(path)
    if not os.path.exists(full):
        abort(404)
    # simple serve
    return send_file(full, as_attachment=False)

def _token_belongs_to_trainee(t: dict, file_token: str) -> bool:
    file_token = (file_token or "").strip()
    if not file_token:
        return False

    # Deliverables (diplôme / SST / attestation)
    dv = t.get("deliverables") or {}
    for k in ("diplome", "carte_sst", "attestation_fin_formation", "attestation_recevabilite"):
        if (dv.get(k) or "").strip() == file_token:
            return True

    # Documents (mono + multi fichiers)
    for d in (t.get("documents") or []):
        if (d.get("file") or "").strip() == file_token:
            return True
        files = d.get("files")
        if isinstance(files, list) and file_token in [x.strip() for x in files if isinstance(x, str)]:
            return True

    # Photo identité (optionnel mais utile)
    if (t.get("identity_photo") or "").strip() == file_token:
        return True

    return False


@app.get("/espace/<token>/download/<path:file_token>")
def public_download_file(token: str, file_token: str):
    data = load_data()
    s, t = find_session_and_trainee_by_token(data, token)
    if not s or not t:
        abort(404)

    # Sécurité : le fichier doit appartenir à CE stagiaire
    if not _token_belongs_to_trainee(t, file_token):
        abort(403)

    full = _detokenize_path(file_token)
    if not os.path.exists(full):
        abort(404)

    return send_file(full, as_attachment=False)


@app.get("/espace/<token>/login")
def public_trainee_login(token: str):
    # ✅ si admin connecté, bypass
    if session.get("admin_logged_in"):
        return redirect(url_for("public_trainee_space", token=token))

    data = load_data()
    s, t = find_session_and_trainee_by_token(data, token)
    if not s or not t:
        abort(404)

    # si déjà auth, go direct
    if session.get(f"public_auth_{token}"):
        return redirect(url_for("public_trainee_space", token=token))

    error = (request.args.get("error") or "").strip()

    # mini page HTML (sans template) pour aller vite
    return f"""
    <!doctype html>
    <html lang="fr">
    <head>
      <meta charset="utf-8">
      <meta name="viewport" content="width=device-width,initial-scale=1">
      <title>Accès espace stagiaire</title>
    
<style>
  *, *::before, *::after {{
    box-sizing: border-box;
  }}

  body {{
    margin: 0;
    font-family: ui-sans-serif, system-ui, -apple-system, Segoe UI, Roboto, Arial;
    background: linear-gradient(180deg, #f6f8fb, #eef2f7);
    min-height: 100vh;
    display: flex;
    align-items: center;
    justify-content: center;
    padding: 16px;
  }}

  .card {{
    width: 100%;
    max-width: 420px;
    background: #fff;
    border-radius: 18px;
    padding: 28px 26px;
    box-shadow: 0 20px 50px rgba(2,6,23,0.12);
    border: 1px solid #e5e7eb;
    overflow: hidden;
  }}

  .logo {{
    display: block;
    margin: 0 auto 14px auto;
    max-height: 70px;
    max-width: 100%;
  }}

  h2 {{
    text-align: center;
    margin: 10px 0 6px 0;
    font-size: 22px;
    color: #0f172a;
  }}

  p {{
    text-align: center;
    margin: 0 0 22px 0;
    font-size: 14px;
    color: #64748b;
  }}

  label {{
    display: block;
    font-weight: 600;
    font-size: 14px;
    margin: 14px 0 6px 0;
    color: #0f172a;
  }}

  input {{
    width: 100%;
    max-width: 100%;
    padding: 12px 14px;
    border-radius: 12px;
    border: 1px solid #d1d5db;
    font-size: 15px;
    outline: none;
  }}

  input:focus {{
    border-color: #1f8f4a;
    box-shadow: 0 0 0 2px rgba(31,143,74,0.15);
  }}

  .btn {{
    margin-top: 22px;
    width: 100%;
    padding: 13px;
    border-radius: 14px;
    border: none;
    background: linear-gradient(135deg, #1f8f4a, #167a3e);
    color: #fff;
    font-size: 16px;
    font-weight: 700;
    cursor: pointer;
  }}

  .btn:hover {{
    filter: brightness(1.05);
  }}

  .hint {{
    margin-top: 12px;
    text-align: center;
    font-size: 13px;
    color: #6b7280;
  }}

  .err {{
    margin: 0 0 12px 0;
    padding: 10px 12px;
    border-radius: 12px;
    background: #fff1f2;
    border: 1px solid #fecdd3;
    color: #9f1239;
    font-size: 13px;
    text-align: center;
  }}
</style>

    </head>
    
    <body>
      <div class="card">
    
        <!-- 🔰 LOGO -->
        <img src="/static/logo-integrale.png" class="logo" alt="Intégrale Academy">
    
        <h2>Accès à votre espace stagiaire</h2>
        <p>Veuillez saisir votre nom de famille et votre date de naissance pour continuer.</p>
    
        <form method="post" action="/espace/{token}/login">
          <label>Nom de famille</label>
          <input name="last_name" autocomplete="family-name" required>
    
          <label>Date de naissance</label>
          <input name="birth" inputmode="numeric" placeholder="JJMMYYYY" required>
    
          <button class="btn">Se connecter</button>
        </form>
    
        <div class="hint">Format demandé : <strong>JJMMYYYY</strong> (ex : 16091993)</div>
      </div>
    </body>
    </html>
    """



@app.post("/espace/<token>/login")
def public_trainee_login_post(token: str):
    # ✅ si admin connecté, bypass
    if session.get("admin_logged_in"):
        return redirect(url_for("public_trainee_space", token=token))

    data = load_data()
    s, t = find_session_and_trainee_by_token(data, token)
    if not s or not t:
        abort(404)

    last_in = (request.form.get("last_name") or "").strip()
    birth_in = (request.form.get("birth") or "").strip()

    # normalisation saisies
    last_in_norm = _norm_lastname(last_in)
    birth_in_digits = re.sub(r"\D+", "", birth_in)  # doit donner 8 chiffres

    # valeurs attendues
    expected_last = _norm_lastname(t.get("last_name", ""))
    expected_birth = _birth_to_ddmmyyyy(t.get("birth_date", ""))

    print("[PUBLIC LOGIN] token =", token)
    print("[PUBLIC LOGIN] trainee keys =", list(t.keys()))
    print("[PUBLIC LOGIN] raw last_name =", t.get("last_name"))
    print("[PUBLIC LOGIN] raw birth_date =", t.get("birth_date"))


    # 🔒 contrôle strict
    if not expected_last or not expected_birth:
        # si les infos ne sont pas renseignées côté dossier, on refuse
        return redirect(url_for("public_trainee_login", token=token, error="1"))

    if last_in_norm == expected_last and birth_in_digits == expected_birth:
        session[f"public_auth_{token}"] = True
        session.permanent = True  # cookie persistant (comme admin)
        _mark_public_login(data, s, t)
        return redirect(url_for("public_trainee_space", token=token))

    return redirect(url_for("public_trainee_login", token=token, error="1"))



@app.post("/admin/sessions/<session_id>/stagiaires/<trainee_id>/documents/<doc_key>/upload")
@admin_login_required
@admin_write_required
def admin_upload_doc_file(session_id: str, trainee_id: str, doc_key: str):
    data = load_data()
    s = find_session(data, session_id)
    if not s:
        abort(404)

    trainees = _session_trainees_list(s)
    t = next((x for x in trainees if x.get("id") == trainee_id), None)
    if not t:
        abort(404)

    training_type = _session_get(s, "training_type", "")

    # ✅ s'assure que la liste de documents correspond à la formation (et supprime dom)
    ensure_documents_schema_for_trainee(t, training_type)

    # ✅ refuse les doc_key inconnus pour cette formation
    if doc_key not in allowed_doc_keys_for_training(training_type):
        return redirect(url_for("admin_trainee_page", session_id=session_id, trainee_id=trainee_id))

    f = request.files.get("file")
    if not f or not f.filename:
        return redirect(url_for("admin_trainee_page", session_id=session_id, trainee_id=trainee_id))

    try:
        stored = _store_file(session_id, trainee_id, "documents", f)
    except Exception:
        return redirect(url_for("admin_trainee_page", session_id=session_id, trainee_id=trainee_id))

    token = _tokenize_path(stored)

    docs = t.get("documents") or []
    for d in docs:
        if d.get("key") == doc_key:
            cur_files = d.get("files")
            if not isinstance(cur_files, list):
                cur_files = []

            old = (d.get("file") or "").strip()
            if old and old not in cur_files:
                cur_files.append(old)

            cur_files.append(token)

            d["files"] = cur_files
            d["file"] = cur_files[0] if cur_files else ""

            cur = (d.get("status") or "").strip().upper()
            if cur in ("", "NON DÉPOSÉ", "NON DEPOSE", "NON_DEPOSE"):
                d["status"] = "A CONTRÔLER"
            if d.get("status") == "A CONTROLER":
                d["status"] = "A CONTRÔLER"
            break

    t["updated_at"] = _now_iso()

    # ✅ recalcul dossier_status
    t["dossier_status"] = "complete" if dossier_is_complete_total(t, training_type) else "incomplete"

    s["trainees"] = trainees
    s.pop("stagiaires", None)
    save_data(data)

    return redirect(url_for("admin_trainee_page", session_id=session_id, trainee_id=trainee_id))

@app.post("/admin/sessions/<session_id>/stagiaires/<trainee_id>/documents/<doc_key>/delete")
@admin_login_required
@admin_write_required
def admin_delete_doc_file(session_id: str, trainee_id: str, doc_key: str):
    data = load_data()
    _force_backup_snapshot(DATA_FILE)
    s = find_session(data, session_id)
    if not s:
        abort(404)

    trainees = _session_trainees_list(s)
    t = next((x for x in trainees if x.get("id") == trainee_id), None)
    if not t:
        abort(404)

    training_type = _session_get(s, "training_type", "")
    ensure_documents_schema_for_trainee(t, training_type)

    # sécurité: n'accepte que les doc_key requis
    if doc_key not in allowed_doc_keys_for_training(training_type):
        return redirect(url_for("admin_trainee_page", session_id=session_id, trainee_id=trainee_id))

    docs = t.get("documents") or []
    target = next((d for d in docs if d.get("key") == doc_key), None)
    if not target:
        return redirect(url_for("admin_trainee_page", session_id=session_id, trainee_id=trainee_id))

    file_token = (request.args.get("file_token") or "").strip()
    existing_files = [x for x in (target.get("files") or []) if x]
    if not existing_files:
        tok = (target.get("file") or "").strip()
        existing_files = [tok] if tok else []

    if file_token:
        tokens = [file_token] if file_token in existing_files else []
    else:
        tokens = list(existing_files)

    # suppression fichiers sur disque
    for tok in tokens:
        try:
            fp = _detokenize_path(tok)
            if os.path.exists(fp):
                os.remove(fp)
        except Exception:
            pass

    # reset du doc
    remaining = [x for x in existing_files if x not in tokens]
    target["files"] = remaining
    target["file"] = remaining[0] if remaining else ""
    if not remaining:
        target["status"] = "NON DÉPOSÉ"
    # on garde le commentaire (pratique), ou tu peux le vider si tu préfères

    t["updated_at"] = _now_iso()

    # recalcul dossier_status
    t["dossier_status"] = "complete" if dossier_is_complete_total(t, training_type) else "incomplete"

    s["trainees"] = trainees
    s.pop("stagiaires", None)
    save_data(data)

    return redirect(url_for("admin_trainee_page", session_id=session_id, trainee_id=trainee_id))




# =========================
# Documents logic
# =========================


def _normalize_status(value: str) -> str:
    return (
        value.strip()
        .upper()
        .replace("_", " ")
        .replace("É", "E")
        .replace("È", "E")
        .replace("Ê", "E")
        .replace("À", "A")
        .replace("Â", "A")
        .replace("Ô", "O")
        .replace("Û", "U")
        .replace("Ï", "I")
        .replace("Î", "I")
    )


def docs_summary_text(
    trainee: Dict[str, Any],
    allowed_statuses: Optional[Iterable[str]] = None,
) -> str:
    lines = []
    allowed_norms = (
        {_normalize_status(s) for s in allowed_statuses} if allowed_statuses else None
    )
    for d in (trainee.get("documents") or []):
        st_raw = (d.get("status") or "A CONTRÔLER").strip().upper()
        st_norm = _normalize_status(st_raw)
        if allowed_norms and st_norm not in allowed_norms:
            continue
        com = (d.get("comment") or "").strip()
        if com:
            lines.append(f"- {d.get('label','document')} : {st_raw} — {com}")
        else:
            lines.append(f"- {d.get('label','document')} : {st_raw}")
    return "\n".join(lines)


import re

def infos_missing_text(trainee: dict) -> str:
    """
    Retourne une liste texte des infos à compléter (ou invalides),
    exactement comme dans l'espace stagiaire (Infos à compléter).
    """
    missing = []

    # --- champs simples obligatoires ---
    simple_required = [
        ("birth_date", "Date de naissance"),
        ("birth_city", "Ville de naissance"),
        ("birth_country", "Pays de naissance"),
        ("nationality", "Nationalité"),
        ("address", "Adresse postale"),
        ("zip_code", "Code postal"),
        ("city", "Ville"),
    ]
    for key, label in simple_required:
        if not (trainee.get(key) or "").strip():
            missing.append(f"- {label}")

    # --- Numéro de sécu : 15 chiffres ---
    secu_raw = (trainee.get("carte_vitale") or "").strip()
    secu_digits = re.sub(r"\D+", "", secu_raw)
    if not secu_raw:
        missing.append("- Numéro de sécurité sociale")
    elif len(secu_digits) != 15:
        missing.append("- Numéro de sécurité sociale (15 chiffres)")

    # --- PRE/CAR ---
    pre_raw = (trainee.get("pre_number") or "").strip()
    pre = pre_raw.upper().replace(" ", "")
    if not pre_raw:
        missing.append("- Numéro PRE / CAR")
    elif not re.match(r"^(PRE|CAR)-(?:\d{3}-)?\d{4}-\d{2}-\d{2}-\d{11,}$", pre):
        missing.append("- Numéro PRE / CAR (format invalide)")

    return "\n".join(missing)


# =========================
# Admin actions — trainee
# =========================
@app.post("/admin/sessions/<session_id>/stagiaires/<trainee_id>/delete")
@admin_login_required
@admin_write_required
def admin_delete_trainee(session_id: str, trainee_id: str):
    data = load_data()
    _force_backup_snapshot(DATA_FILE)
    s = find_session(data, session_id)
    if not s:
        abort(404)
    trainees = _session_trainees_list(s)
    trainees = [x for x in trainees if x.get("id") != trainee_id]
    s["trainees"] = trainees
    s.pop("stagiaires", None)
    save_data(data)
    return redirect(url_for("admin_trainees", session_id=session_id))

def _replace_in_docx(doc: Document, replacements: dict) -> None:
    def replace_in_paragraph(p):
        # Remplace dans les runs pour garder le style
        for run in p.runs:
            for k, v in replacements.items():
                if k in run.text:
                    run.text = run.text.replace(k, v)

    def replace_in_table(table):
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)
                for t2 in cell.tables:
                    replace_in_table(t2)

    # Corps du document
    for p in doc.paragraphs:
        replace_in_paragraph(p)

    for table in doc.tables:
        replace_in_table(table)

    # En-têtes / pieds de page
    for section in doc.sections:
        for p in section.header.paragraphs:
            replace_in_paragraph(p)
        for table in section.header.tables:
            replace_in_table(table)

        for p in section.footer.paragraphs:
            replace_in_paragraph(p)
        for table in section.footer.tables:
            replace_in_table(table)


@app.get("/admin/sessions/<session_id>/stagiaires/<trainee_id>/etiquette.docx")
def admin_etiquette_docx(session_id: str, trainee_id: str):
    data = load_data()
    s = find_session(data, session_id)
    if not s:
        abort(404)

    trainees = _session_trainees_list(s)
    t = next((x for x in trainees if x.get("id") == trainee_id), None)
    if not t:
        abort(404)

    # 1) Choix du modèle Word selon le type de formation
    training_type = (_session_get(s, "training_type", "") or "").strip().upper()

    TEMPLATE_MAP = {
        "A3P": "etiquette_a3p.docx",
        "APS": "etiquette_aps.docx",
        "CHAUFFEUR VTC": "etiquette_vtc.docx",
        "VTC": "etiquette_vtc.docx",
        "DIRIGEANT": "etiquette_dirigeant.docx",
        "DIRIGEANT INITIAL": "etiquette_dirigeant_initial.docx",
        "DIRIGEANT VAE": "etiquette_dirigeant.docx",
    }

    template_name = TEMPLATE_MAP.get(training_type)
    if not template_name:
        abort(400, f"Aucun modèle Word prévu pour la formation : {training_type}")

    template_path = os.path.join("templates_word", template_name)
    if not os.path.exists(template_path):
        abort(500, f"Fichier Word manquant : {template_name} (dans /templates_word)")

    # 2) Ouvrir le modèle
    doc = Document(template_path)

    # 3) Remplacements
    replacements = {
        "{{NOM}}": (t.get("last_name", "") or "").upper(),
        "{{PRENOM}}": (t.get("first_name", "") or "").upper(),
        "{{FORMATION}}": _session_get(s, "name", ""),
        "{{TYPE_FORMATION}}": training_type,
        "{{DATES}}": f"{fr_date(_session_get(s,'date_start',''))} → {fr_date(_session_get(s,'date_end',''))}",
    }

    _replace_in_docx(doc, replacements)

    # ✅ Photo identité dans l'étiquette (même taille, sans déformation)
    photo_token = (t.get("identity_photo") or "").strip()
    if photo_token:
        photo_path = _detokenize_path(photo_token)
        _insert_label_photo(doc, "{{PHOTO}}", photo_path, width_cm=5.41, height_cm=6.41)
    else:
        # si pas de photo, on enlève le placeholder
        _replace_in_docx(doc, {"{{PHOTO}}": ""})

    # 4) Télécharger
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    t["etiquette_word_downloaded_at"] = _now_iso()
    s["trainees"] = trainees
    save_data(data)

    filename = f"etiquette_{t.get('last_name','')}_{t.get('first_name','')}.docx".replace(" ", "_")
    return send_file(
        buf,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

def _prepare_photo_for_label(src_path: str, target_ratio: float) -> str:
    """
    Recadre la photo au centre au bon ratio (sans déformation),
    et retourne un chemin vers un JPG temporaire compatible Word.
    """
    im = Image.open(src_path).convert("RGB")
    w, h = im.size
    src_ratio = w / h

    if src_ratio > target_ratio:
        # image trop large → on coupe sur les côtés
        new_w = int(h * target_ratio)
        left = (w - new_w) // 2
        im = im.crop((left, 0, left + new_w, h))
    else:
        # image trop haute → on coupe en haut/bas
        new_h = int(w / target_ratio)
        top = (h - new_h) // 2
        im = im.crop((0, top, w, top + new_h))

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".jpg")
    im.save(tmp.name, "JPEG", quality=90)
    return tmp.name


def _insert_label_photo(doc: Document, placeholder: str, photo_path: str, width_cm: float, height_cm: float) -> bool:
    if not photo_path or not os.path.exists(photo_path):
        return False

    target_ratio = width_cm / height_cm
    prepared = _prepare_photo_for_label(photo_path, target_ratio)

    width = Inches(width_cm / 2.54)
    height = Inches(height_cm / 2.54)

    def process_paragraph(p) -> bool:
        full = "".join(run.text for run in p.runs)
        if placeholder not in full:
            return False

        # vide le paragraphe
        for run in p.runs:
            run.text = ""

        # insère l'image recadrée au bon ratio, donc pas de déformation
        r = p.add_run()
        r.add_picture(prepared, width=width, height=height)
        return True

    def process_table(table) -> bool:
        ok = False
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    ok = process_paragraph(p) or ok
                for t2 in cell.tables:
                    ok = process_table(t2) or ok
        return ok

    inserted = False
    for p in doc.paragraphs:
        inserted = process_paragraph(p) or inserted
    for table in doc.tables:
        inserted = process_table(table) or inserted

    return inserted


@app.post("/admin/sessions/<session_id>/stagiaires/<trainee_id>/send-access")
@admin_login_required
def admin_send_access(session_id: str, trainee_id: str):
    # Autorisé aussi pour le profil consultation (lecture seule)
    # afin de permettre l'envoi de l'accès à l'espace stagiaire.
    data = load_data()
    s = find_session(data, session_id)
    if not s:
        abort(404)
    trainees = _session_trainees_list(s)
    t = next((x for x in trainees if x.get("id") == trainee_id), None)
    if not t:
        abort(404)

    link = f"{PUBLIC_STUDENT_PORTAL_BASE.rstrip('/')}/espace/{t.get('public_token','')}"
    training_type = _session_get(s, "training_type", "")
    if "VTC" in (training_type or "").upper():
        first_name = t.get("first_name", "")
        subject, html = build_vtc_onboarding_email(first_name, link)
        sms = build_vtc_onboarding_sms(first_name, link)
        brevo_send_email(t.get("email", ""), subject, html)
        brevo_send_sms(t.get("phone", ""), sms)
    else:
        subject = "Accès à votre espace stagiaire – Intégrale Academy"
        html = mail_layout(f"""
          <h2>Votre espace stagiaire est disponible</h2>
          <p>Formation : <strong>{_session_get(s,'name','')}</strong></p>
          <p>
            <a href="{link}" style="display:inline-block;background:#1f8f4a;color:white;padding:10px 14px;border-radius:10px;text-decoration:none">
              Accéder à mon espace stagiaire
            </a>
          </p>
        """)
        sms = f"Intégrale Academy : votre espace stagiaire est disponible : {link}"
        brevo_send_email(t.get("email", ""), subject, html)
        brevo_send_sms(t.get("phone", ""), sms)

    t["access_sent_at"] = _now_iso()
    s["trainees"] = trainees
    save_data(data)
    return redirect(url_for("admin_trainee_page", session_id=session_id, trainee_id=trainee_id))


@app.post("/admin/sessions/<session_id>/stagiaires/<trainee_id>/vtc-cmar-relance")
@admin_login_required
@admin_write_required
def admin_vtc_cmar_relance(session_id: str, trainee_id: str):
    data = load_data()
    s = find_session(data, session_id)
    if not s:
        abort(404)

    training_type = (_session_get(s, "training_type", "") or "").upper()
    if "VTC" not in training_type:
        return redirect(url_for("admin_trainee_page", session_id=session_id, trainee_id=trainee_id))

    trainees = _session_trainees_list(s)
    t = next((x for x in trainees if x.get("id") == trainee_id), None)
    if not t:
        abort(404)

    _send_vtc_credentials_reminder(data, s, t, "Relance manuelle CMAR (admin)")

    s["trainees"] = trainees
    s.pop("stagiaires", None)
    save_data(data)

    return redirect(url_for("admin_trainee_page", session_id=session_id, trainee_id=trainee_id))

# =========================
# Convention — non signée
# =========================
@app.post("/admin/sessions/<session_id>/stagiaires/<trainee_id>/convention/unsigned-notify")
@admin_login_required
@admin_write_required
def admin_convention_unsigned_notify(session_id: str, trainee_id: str):
    data = load_data()
    s = find_session(data, session_id)
    if not s:
        abort(404)

    trainees = _session_trainees_list(s)
    t = next((x for x in trainees if x.get("id") == trainee_id), None)
    if not t:
        abort(404)

    formation_type = formation_label(_session_get(s, "training_type", ""))
    dstart = fr_date(_session_get(s, "date_start", ""))
    dend = fr_date(_session_get(s, "date_end", ""))

    subject = "Relance contrat de formation à signer – Intégrale Academy"
    html = mail_layout(f"""
      <h2 style="text-align:center;color:#b91c1c">Contrat de formation à signer</h2>

      <p>Bonjour <strong>{t.get("first_name","").strip() or "Madame, Monsieur"}</strong>,</p>

      <p>
        Vous n'avez pas encore signé votre <strong>Contrat</strong>
        concernant votre formation <strong>{formation_type}</strong> qui se déroulera (du <strong>{dstart}</strong> au <strong>{dend}</strong>).
      </p>

      <p>
        Nous vous remercions de bien vouloir finaliser la signature électronique dès que possible. Si vous n’avez pas reçu le lien de signature, nous vous remercions de bien vouloir nous contacter au 04 22 47 07 68.
      </p>

      <p style="margin-top:22px">
        Merci par avance,<br>
        <strong>Clément VAILLANT</strong><br>
        Directeur Intégrale Academy
      </p>
    """)

    sms_prefix = f"Bonjour {t.get('first_name','').strip()}, " if (t.get("first_name") or "").strip() else "Bonjour, "
    sms = (
        f"Intégrale Academy {sms_prefix}"
        f"Vous n'avez pas encore signé votre Contrat"
        f"concernant votre formation {formation_type} ({dstart} au {dend}). "
        "Nous vous remercions de bien vouloir procéder à la signature de ce document. Besoin d'aide ? 04 22 47 07 68."
    )

    brevo_send_email(t.get("email", ""), subject, html)
    brevo_send_sms(t.get("phone", ""), sms)

    full_name = _format_trainee_name(t.get("first_name", ""), t.get("last_name", ""))
    period = f"{dstart} au {dend}" if (dstart and dend) else "Dates à confirmer"
    add_notification(
        data,
        "notifications_convention_unsigned",
        f"{full_name} • {formation_type} • {period}",
        meta={
            "first_name": t.get("first_name", ""),
            "last_name": t.get("last_name", ""),
            "phone": t.get("phone", ""),
            "email": t.get("email", ""),
            "training": formation_type,
            "training_start_date": dstart,
            "training_end_date": dend,
            "session_id": s.get("id"),
            "trainee_id": t.get("id"),
        },
    )

    t["convention_unsigned_notified_at"] = _now_iso()
    t["updated_at"] = _now_iso()

    s["trainees"] = trainees
    save_data(data)
    return redirect(url_for("admin_trainee_page", session_id=session_id, trainee_id=trainee_id))

# =========================
# Test de français — notify/relance
# =========================
def _build_test_fr_payload(t, s, code, deadline, mode):
    link = "https://testb1.lapreventionsecurite.org/Public/"
    formation_type = formation_label(_session_get(s, "training_type", ""))
    dstart = fr_date(_session_get(s, "date_start", ""))
    dend = fr_date(_session_get(s, "date_end", ""))
    deadline_fr = fr_date(deadline)

    if mode == "notify":
        subject = "Test de français à réaliser – Intégrale Academy"
        html = mail_layout(f"""
      <h2 style="text-align:center">📝 Test de français obligatoire</h2>

      <p>Bonjour <strong>{t.get("first_name","").strip() or "Madame, Monsieur"}</strong>,</p>

      <p>
        Je me permets de revenir vers vous concernant votre inscription en formation
        <strong>{formation_type}</strong>, qui se déroulera du <strong>{dstart}</strong> au <strong>{dend}</strong>.
      </p>

      <p>
        Conformément à la réglementation, nous vous demandons de bien vouloir procéder au
        <strong>Test de français obligatoire</strong> avant votre entrée en formation.
      </p>

      <div style="background:#f3f4f6;border:1px solid #e5e7eb;border-radius:12px;padding:14px;margin:16px 0">
        <p style="margin:0 0 10px 0"><strong>🔗 Lien du test :</strong>
          <a href="{link}" style="color:#1f8f4a;text-decoration:none;font-weight:bold">{link}</a>
        </p>

        <p style="margin:0 0 10px 0"><strong>🔑 Code d’activation :</strong>
          <span style="font-size:16px;letter-spacing:1px">{code}</span>
        </p>

        <p style="margin:0;color:#b91c1c;font-weight:bold">
          ⚠️ Attention : le test doit être réalisé le <u>{deadline_fr}</u>.
        </p>
      </div>

      <p>Je vous remercie par avance et je vous souhaite une excellente journée,</p>

      <p style="margin-top:22px">
        <strong>Clément VAILLANT</strong><br>
        Directeur Intégrale Academy
      </p>

      <p style="text-align:center;margin-top:18px">
        <a href="{link}"
           style="display:inline-block;background:#1f8f4a;color:white;padding:12px 18px;border-radius:10px;text-decoration:none;font-weight:bold">
          👉 Accéder au test de français
        </a>
      </p>
    """)

        sms = (
            f"Intégrale Academy 📝 Bonjour {t.get('first_name','')}, "
            f"Vous devez réaliser le Test de français obligatoire pour votre formation {formation_type}. "
            f"Lien : {link} | Code : {code} | À faire le {deadline_fr}. "
            f"Besoin d’aide ? 04 22 47 07 68"
        )
        status = "in_progress"
        stamp_field = "test_fr_last_notified_at"
    elif mode == "relance":
        subject = "Relance – Test de français à réaliser"
        html = mail_layout(f"""
      <h2 style="text-align:center;color:#b91c1c">⏰ Relance – Test de français obligatoire</h2>

      <p>Bonjour <strong>{t.get("first_name","").strip() or "Madame, Monsieur"}</strong>,</p>

      <p>
        Nous revenons vers vous concernant votre inscription en formation
        <strong>{formation_type}</strong> (du <strong>{dstart}</strong> au <strong>{dend}</strong>).
      </p>

      <p>
        À ce jour, nous n’avons pas encore reçu la validation de votre <strong>Test de français obligatoire</strong>.
        Merci de le réaliser dès que possible.
      </p>

      <div style="background:#fef2f2;border:1px solid #fecaca;border-radius:12px;padding:14px;margin:16px 0">
        <p style="margin:0 0 10px 0"><strong>🔗 Lien du test :</strong>
          <a href="{link}" style="color:#1f8f4a;text-decoration:none;font-weight:bold">{link}</a>
        </p>

        <p style="margin:0 0 10px 0"><strong>🔑 Code d’activation :</strong>
          <span style="font-size:16px;letter-spacing:1px">{code}</span>
        </p>

        <p style="margin:0;color:#b91c1c;font-weight:bold">
          ⚠️ Date limite : <u>{deadline_fr}</u>
        </p>
      </div>

      <p style="margin-top:22px">
        Si vous avez la moindre difficulté, contactez-nous au <strong>04 22 47 07 68</strong>.
      </p>

      <p style="margin-top:22px">
        Merci par avance,<br>
        <strong>Clément VAILLANT</strong><br>
        Directeur Intégrale Academy
      </p>

      <p style="text-align:center;margin-top:18px">
        <a href="{link}"
           style="display:inline-block;background:#1f8f4a;color:white;padding:12px 18px;border-radius:10px;text-decoration:none;font-weight:bold">
          👉 Accéder au test de français
        </a>
      </p>
    """)

        sms = (
            f"Intégrale Academy ⏰ Relance : Bonjour {t.get('first_name','')}, "
            f"Vous n'avez pas encore réalisé votre Test de français obligatoire avant votre entrée en formation {formation_type}. "
            f"Lien : {link} | Code : {code} | Date limite : {deadline_fr}. "
            f"Besoin d’aide ? 04 22 47 07 68"
        )
        status = "relance"
        stamp_field = "test_fr_last_relance_at"
    else:
        subject = "Échec au test de français – Test à refaire"
        html = mail_layout(f"""
      <h2 style="text-align:center;color:#b91c1c">❌ Échec au test de français</h2>

      <p>Bonjour <strong>{t.get("first_name","").strip() or "Madame, Monsieur"}</strong>,</p>

      <p>
        Nous revenons vers vous concernant votre inscription en formation
        <strong>{formation_type}</strong> (du <strong>{dstart}</strong> au <strong>{dend}</strong>).
      </p>

      <p>
        Suite à un <strong>échec au test de français</strong>, vous devez <strong>refaire le test de français obligatoire</strong>.
      </p>

      <div style="background:#fef2f2;border:1px solid #fecaca;border-radius:12px;padding:14px;margin:16px 0">
        <p style="margin:0 0 10px 0"><strong>🔗 Lien du test :</strong>
          <a href="{link}" style="color:#1f8f4a;text-decoration:none;font-weight:bold">{link}</a>
        </p>

        <p style="margin:0 0 10px 0"><strong>🔑 Code d’activation :</strong>
          <span style="font-size:16px;letter-spacing:1px">{code}</span>
        </p>

        <p style="margin:0;color:#b91c1c;font-weight:bold">
          ⚠️ Date limite : <u>{deadline_fr}</u>
        </p>
      </div>

      <p style="margin-top:22px">
        Si vous avez la moindre difficulté, contactez-nous au <strong>04 22 47 07 68</strong>.
      </p>

      <p style="margin-top:22px">
        Merci par avance,<br>
        <strong>Clément VAILLANT</strong><br>
        Directeur Intégrale Academy
      </p>

      <p style="text-align:center;margin-top:18px">
        <a href="{link}"
           style="display:inline-block;background:#1f8f4a;color:white;padding:12px 18px;border-radius:10px;text-decoration:none;font-weight:bold">
          👉 Refaire le test de français
        </a>
      </p>
    """)

        sms = (
            f"Intégrale Academy ❌ Bonjour {t.get('first_name','')}, "
            f"Suite à un échec, vous devez refaire votre Test de français obligatoire pour la formation {formation_type}. "
            f"Lien : {link} | Code : {code} | Date limite : {deadline_fr}. "
            f"Besoin d’aide ? 04 22 47 07 68"
        )
        status = "relance"
        stamp_field = "test_fr_last_relance_at"

    return {
        "subject": subject,
        "html": html,
        "sms": sms,
        "status": status,
        "stamp_field": stamp_field,
    }


@app.post("/admin/sessions/<session_id>/stagiaires/<trainee_id>/test-fr/notify")
@admin_login_required
@admin_write_required
def admin_test_fr_notify(session_id: str, trainee_id: str):
    code = (request.form.get("code") or "").strip()
    deadline = (request.form.get("deadline") or "").strip()
    if not code or not deadline:
        return redirect(url_for("admin_trainee_page", session_id=session_id, trainee_id=trainee_id))

    data = load_data()
    s = find_session(data, session_id)
    if not s:
        abort(404)
    trainees = _session_trainees_list(s)
    t = next((x for x in trainees if x.get("id")==trainee_id), None)
    if not t:
        abort(404)

    payload = _build_test_fr_payload(t, s, code, deadline, "notify")
    brevo_send_email(t.get("email",""), payload["subject"], payload["html"])
    brevo_send_sms(t.get("phone",""), payload["sms"])

    now = _now_iso()
    t["test_fr_status"] = payload["status"]
    t["test_fr_code"] = code
    t["test_fr_deadline"] = deadline
    t[payload["stamp_field"]] = now
    t["updated_at"] = now


    s["trainees"] = trainees
    save_data(data)
    return redirect(url_for("admin_trainee_page", session_id=session_id, trainee_id=trainee_id))

@app.post("/admin/sessions/<session_id>/stagiaires/<trainee_id>/test-fr/relance")
@admin_login_required
@admin_write_required
def admin_test_fr_relance(session_id: str, trainee_id: str):
    code = (request.form.get("code") or "").strip()
    deadline = (request.form.get("deadline") or "").strip()
    if not code or not deadline:
        return redirect(url_for("admin_trainee_page", session_id=session_id, trainee_id=trainee_id))

    data = load_data()
    s = find_session(data, session_id)
    if not s:
        abort(404)

    trainees = _session_trainees_list(s)
    t = next((x for x in trainees if x.get("id") == trainee_id), None)
    if not t:
        abort(404)

    payload = _build_test_fr_payload(t, s, code, deadline, "relance")
    brevo_send_email(t.get("email", ""), payload["subject"], payload["html"])
    brevo_send_sms(t.get("phone", ""), payload["sms"])

    now = _now_iso()
    t["test_fr_status"] = payload["status"]
    t["test_fr_code"] = code
    t["test_fr_deadline"] = deadline
    t[payload["stamp_field"]] = now
    t["updated_at"] = now

    add_notification(
        data,
        "notifications_test_fr",
        f"Cette personne n'a pas réalisé son test de français, un nouveau lien a été envoyé, le test doit être réalisé le {fr_date(deadline) or deadline}",
        meta={
            "type": "test_fr_relance",
            "session_id": s.get("id"),
            "trainee_id": t.get("id"),
            "first_name": t.get("first_name", ""),
            "last_name": t.get("last_name", ""),
            "phone": t.get("phone", ""),
            "deadline": deadline,
        },
    )

    s["trainees"] = trainees
    save_data(data)
    return redirect(url_for("admin_trainee_page", session_id=session_id, trainee_id=trainee_id))


@app.post("/admin/sessions/<session_id>/stagiaires/<trainee_id>/test-fr/echec")
@admin_login_required
@admin_write_required
def admin_test_fr_echec(session_id: str, trainee_id: str):
    code = (request.form.get("code") or "").strip()
    deadline = (request.form.get("deadline") or "").strip()
    if not code or not deadline:
        return redirect(url_for("admin_trainee_page", session_id=session_id, trainee_id=trainee_id))

    data = load_data()
    s = find_session(data, session_id)
    if not s:
        abort(404)

    trainees = _session_trainees_list(s)
    t = next((x for x in trainees if x.get("id") == trainee_id), None)
    if not t:
        abort(404)

    payload = _build_test_fr_payload(t, s, code, deadline, "failure")
    brevo_send_email(t.get("email", ""), payload["subject"], payload["html"])
    brevo_send_sms(t.get("phone", ""), payload["sms"])

    now = _now_iso()
    t["test_fr_status"] = payload["status"]
    t["test_fr_code"] = code
    t["test_fr_deadline"] = deadline
    t[payload["stamp_field"]] = now
    t["test_fr_last_failed_at"] = now
    t["updated_at"] = now

    add_notification(
        data,
        "notifications_test_fr",
        f"Cette personne a échoué son test de français, un nouveau lien a été envoyé, le test doit être réalisé le {fr_date(deadline) or deadline}",
        meta={
            "type": "test_fr_echec",
            "session_id": s.get("id"),
            "trainee_id": t.get("id"),
            "first_name": t.get("first_name", ""),
            "last_name": t.get("last_name", ""),
            "phone": t.get("phone", ""),
            "deadline": deadline,
        },
    )

    s["trainees"] = trainees
    save_data(data)
    return redirect(url_for("admin_trainee_page", session_id=session_id, trainee_id=trainee_id))


@app.post("/admin/sessions/<session_id>/stagiaires/test-fr/notify-bulk")
@admin_login_required
@admin_write_required
def admin_test_fr_notify_bulk(session_id: str):
    code = (request.form.get("code") or "").strip()
    deadline = (request.form.get("deadline") or "").strip()
    if not code or not deadline:
        if request.headers.get("X-Requested-With") == "XMLHttpRequest":
            return jsonify({"ok": False, "error": "code_or_deadline_missing"}), 400
        return redirect(url_for("admin_trainees", session_id=session_id))

    data = load_data()
    s = find_session(data, session_id)
    if not s:
        abort(404)

    trainees = _session_trainees_list(s)
    total = len(trainees)
    sent = 0
    for t in trainees:
        payload = _build_test_fr_payload(t, s, code, deadline, "notify")
        email = (t.get("email") or "").strip()
        phone = (t.get("phone") or "").strip()
        can_send = bool(email or phone)
        if can_send:
            sent += 1
        if email:
            brevo_send_email(email, payload["subject"], payload["html"])
        if phone:
            brevo_send_sms(phone, payload["sms"])

        t["test_fr_status"] = payload["status"]
        t["test_fr_code"] = code
        t["test_fr_deadline"] = deadline
        t[payload["stamp_field"]] = _now_iso()
        t["updated_at"] = _now_iso()

    s["trainees"] = trainees
    save_data(data)
    if request.headers.get("X-Requested-With") == "XMLHttpRequest":
        return jsonify(
            {
                "ok": True,
                "total": total,
                "sent": sent,
                "missing": total - sent,
                "all_ok": total > 0 and sent == total,
            }
        )
    return redirect(url_for("admin_trainees", session_id=session_id))

# =========================
# Documents — notify / nonconform / relance / zip
# =========================
@app.post("/admin/sessions/<session_id>/stagiaires/<trainee_id>/docs/notify")
@admin_login_required
@admin_write_required
def admin_docs_notify(session_id: str, trainee_id: str):
    data = load_data()
    s = find_session(data, session_id)
    if not s:
        abort(404)

    trainees = _session_trainees_list(s)
    t = next((x for x in trainees if x.get("id") == trainee_id), None)
    if not t:
        abort(404)

    link = f"{PUBLIC_STUDENT_PORTAL_BASE.rstrip('/')}/espace/{t.get('public_token','')}"
    subject = "Envoi de documents – Action requise (Intégrale Academy)"

    formation_type = formation_label(_session_get(s, "training_type", ""))
    dstart = fr_date(_session_get(s, "date_start", ""))
    dend = fr_date(_session_get(s, "date_end", ""))

    first_name = (t.get("first_name") or "").strip() or "Madame, Monsieur"

    html = mail_layout(f"""
      <h2 style="text-align:center">📄 Envoi de documents – Dossier formation</h2>

      <p>Bonjour <strong>{first_name}</strong>,</p>

      <p>
        Dans le cadre de votre inscription en formation
        <strong>{formation_type}</strong> (du <strong>{dstart}</strong> au <strong>{dend}</strong>),
        nous vous invitons à compléter votre Dossier Formation via votre espace stagiaire.
      </p>

      <div style="background:#f3f4f6;border:1px solid #e5e7eb;border-radius:12px;padding:14px;margin:16px 0">
        <p style="margin:0 0 10px 0">
          <strong>📍 Accès à votre espace stagiaire :</strong><br>
          <a href="{link}" style="color:#1f8f4a;text-decoration:none;font-weight:bold">{link}</a>
        </p>

        <p style="margin:0;color:#b91c1c;font-weight:bold">
          ⚠️ Pour un meilleur traitement de votre inscription, nous vous invitons à compléter votre dossier dès que possible. Attention, votre dossier doit être complet au plus tard 10 jours avant votre entrée en formation.
        </p>
      </div>

      <p style="margin-top:22px">
        Si vous avez la moindre difficulté, vous pouvez nous contacter au <strong>04 22 47 07 68</strong>.
      </p>

      <p style="margin-top:22px">
        Merci par avance,<br>
        <strong>Clément VAILLANT</strong><br>
        Directeur Intégrale Academy
      </p>

      <p style="text-align:center;margin-top:18px">
        <a href="{link}"
           style="display:inline-block;background:#1f8f4a;color:white;padding:12px 18px;border-radius:10px;text-decoration:none;font-weight:bold">
          👉 Accéder à mon espace stagiaire
        </a>
      </p>
    """)

    sms = (
        f"Intégrale Academy 📄 Bonjour {t.get('first_name','')}, "
        f"Nous vous remercions de bien vouloir compléter votre Dossier Formation concernant votre formation {formation_type} "
        f"({dstart} au {dend}) via votre espace : {link} "
        f"Besoin d’aide ? 04 22 47 07 68"
    )

    brevo_send_email(t.get("email", ""), subject, html)
    brevo_send_sms(t.get("phone", ""), sms)

    t["docs_notified_at"] = _now_iso()
    t["updated_at"] = _now_iso()

    s["trainees"] = trainees
    s.pop("stagiaires", None)
    save_data(data)

    return redirect(url_for("admin_trainee_page", session_id=session_id, trainee_id=trainee_id))
    
@app.post("/admin/sessions/<session_id>/stagiaires/<trainee_id>/docs/nonconform/notify")
@admin_login_required
@admin_write_required
def admin_docs_nonconform_notify(session_id: str, trainee_id: str):
    data = load_data()
    s = find_session(data, session_id)
    if not s:
        abort(404)

    trainees = _session_trainees_list(s)
    t = next((x for x in trainees if x.get("id") == trainee_id), None)
    if not t:
        abort(404)

    link = f"{PUBLIC_STUDENT_PORTAL_BASE.rstrip('/')}/espace/{t.get('public_token','')}"
    training_type = _session_get(s, "training_type", "")
    ensure_documents_schema_for_trainee(t, training_type)

    details = docs_summary_text(
        t,
        allowed_statuses={
            "NON CONFORME",
            "NON_CONFORME",
            "NON DÉPOSÉ",
            "NON DEPOSE",
        },
    )

    subject = "Documents non conformes – Action requise (Intégrale Academy)"

    html = mail_layout(f"""
      <h2 style="text-align:center;color:#b91c1c">❌ Documents non conformes / à corriger</h2>

      <p>Bonjour <strong>{(t.get("first_name") or "").strip() or "Madame, Monsieur"}</strong>,</p>

      <p>
        Certains documents déposés dans votre dossier ne sont pas conformes (ou doivent être corrigés).
        Merci de consulter le détail ci-dessous et de déposer les documents corrigés depuis votre espace stagiaire.
      </p>

      <div style="background:#fef2f2;border:1px solid #fecaca;border-radius:12px;padding:14px;margin:16px 0">
        <p style="margin:0 0 10px 0"><strong>📌 Détail de vos documents :</strong></p>
        <pre style="white-space:pre-wrap;background:#fff;border:1px solid #fee2e2;padding:10px;border-radius:10px;margin:0">{details or "Aucun détail disponible."}</pre>

        <p style="margin:14px 0 0 0">
          <strong>📍 Déposer les documents corrigés :</strong><br>
          <a href="{link}" style="color:#1f8f4a;text-decoration:none;font-weight:bold">{link}</a>
        </p>

        <p style="margin:10px 0 0 0;color:#b91c1c;font-weight:bold">
          ⚠️ Merci de corriger et renvoyer dès que possible pour valider votre inscription.
        </p>
      </div>

      <p style="margin-top:22px">
        Besoin d’aide ? Contactez-nous au <strong>04 22 47 07 68</strong>.
      </p>

      <p style="margin-top:22px">
        Merci par avance,<br>
        <strong>Clément VAILLANT</strong><br>
        Directeur Intégrale Academy
      </p>

      <p style="text-align:center;margin-top:18px">
        <a href="{link}"
           style="display:inline-block;background:#1f8f4a;color:white;padding:12px 18px;border-radius:10px;text-decoration:none;font-weight:bold">
          👉 Accéder à mon espace stagiaire
        </a>
      </p>
    """)

    sms = (
        f"Intégrale Academy ❌ Bonjour {t.get('first_name','')}, "
        f"Certains documents déposés sont NON CONFORMES. Nous vous invitons à corriger votre dépôt. La liste détaillée des non conformités vous a été adressée par mail. "
        f"Merci de déposer les documents corrigés sur votre espace : {link} "
        f"Aide : 04 22 47 07 68"
    )

    brevo_send_email(t.get("email",""), subject, html)
    brevo_send_sms(t.get("phone",""), sms)

    t["docs_last_nonconform_notified_at"] = _now_iso()
    t["updated_at"] = _now_iso()
    s["trainees"] = trainees
    s.pop("stagiaires", None)
    save_data(data)

    return redirect(url_for("admin_trainee_page", session_id=session_id, trainee_id=trainee_id))

@app.post("/admin/sessions/<session_id>/stagiaires/<trainee_id>/docs/relance")
@admin_login_required
@admin_write_required
def admin_docs_relance(session_id: str, trainee_id: str):
    data = load_data()
    s = find_session(data, session_id)
    if not s:
        abort(404)

    trainees = _session_trainees_list(s)
    t = next((x for x in trainees if x.get("id") == trainee_id), None)
    if not t:
        abort(404)

    link = f"{PUBLIC_STUDENT_PORTAL_BASE.rstrip('/')}/espace/{t.get('public_token','')}"
    training_type = _session_get(s, "training_type", "")
    ensure_documents_schema_for_trainee(t, training_type)
    
    docs_details = docs_summary_text(
        t,
        allowed_statuses={
            "A CONTRÔLER",
            "A CONTROLER",
            "NON CONFORME",
            "NON_CONFORME",
            "NON DÉPOSÉ",
            "NON DEPOSE",
            "NON_DEPOSE",
        },
    )
    infos_details = infos_missing_text(t)

    formation_type = formation_label(_session_get(s, "training_type", ""))
    dstart = fr_date(_session_get(s, "date_start", ""))
    dend = fr_date(_session_get(s, "date_end", ""))

    first_name = (t.get("first_name") or "").strip() or "Madame, Monsieur"

    subject = "Relance : Dossier Formation incomplet"

    html = mail_layout(f"""
      <h2 style="text-align:center;color:#b91c1c">⏰ Relance – Votre Dossier Formation est incomplet</h2>

      <p>Bonjour <strong>{first_name}</strong>,</p>

      <p>
        Nous revenons vers vous concernant votre inscription en formation
        <strong>{formation_type}</strong> (du <strong>{dstart}</strong> au <strong>{dend}</strong>).
      </p>

      <p>
        À ce jour, votre dossier est INCOMPLET (éléments manquants et/ou à corriger).
        Merci de déposer les éléments nécessaires dès que possible via votre espace stagiaire.
      </p>

      <div style="background:#fef2f2;border:1px solid #fecaca;border-radius:12px;padding:14px;margin:16px 0">
        <p style="margin:0 0 10px 0"><strong>📌 Votre dossier détaillé :</strong></p>
       <pre style="white-space:pre-wrap;background:#fff;border:1px solid #fee2e2;padding:10px;border-radius:10px;margin:0">{docs_details or "Aucun document en attente."}</pre>

    <p style="margin:14px 0 10px 0"><strong>🧾 Informations à compléter :</strong></p>
    <pre style="white-space:pre-wrap;background:#fff;border:1px solid #fee2e2;padding:10px;border-radius:10px;margin:0">{infos_details or "Aucune information manquante."}</pre>

        <p style="margin:12px 0 0 0">
          <strong>📍 Informations à compléter et Dépôt des documents :</strong><br>
          <a href="{link}" style="color:#1f8f4a;text-decoration:none;font-weight:bold">{link}</a>
        </p>

        <p style="margin:10px 0 0 0;color:#b91c1c;font-weight:bold">
          ⚠️ Nous vous remercions de bien vouloir compléter votre dossier dès que possible !
        </p>
      </div>

      <p style="margin-top:22px">
        Si vous avez la moindre difficulté, contactez-nous au <strong>04 22 47 07 68</strong>.
      </p>

      <p style="margin-top:22px">
        Merci par avance,<br>
        <strong>Clément VAILLANT</strong><br>
        Directeur Intégrale Academy
      </p>

      <p style="text-align:center;margin-top:18px">
        <a href="{link}"
           style="display:inline-block;background:#1f8f4a;color:white;padding:12px 18px;border-radius:10px;text-decoration:none;font-weight:bold">
          👉 Accéder à mon espace stagiaire
        </a>
      </p>
    """)

    sms = (
        f"Intégrale Academy ⏰ Relance : Bonjour {t.get('first_name','')}, "
        f"Nous revenons vers vous au sujet de votre formation {formation_type}. A ce jour votre Dossier Formation est INCOMPLET. Votre formation approche, et pour un meilleur suivi de votre inscription, nous vous remercions de bien vouloir compléter votre dossier dès que possible. "
        f"Pour rappel, votre dossier doit être COMPLET au plus tard 10 jours avant votre entrée en formation. Vous pouvez compléter votre dossier en cliquant ici : {link} "
        f"Besoin d’aide ? 04 22 47 07 68"
    )

    brevo_send_email(t.get("email", ""), subject, html)
    brevo_send_sms(t.get("phone", ""), sms)

    t["docs_last_relance_at"] = _now_iso()
    t["updated_at"] = _now_iso()

    s["trainees"] = trainees
    s.pop("stagiaires", None)
    save_data(data)

    return redirect(url_for("admin_trainee_page", session_id=session_id, trainee_id=trainee_id))

@app.get("/admin/sessions/<session_id>/stagiaires/<trainee_id>/documents.zip")
@admin_login_required
def admin_docs_zip(session_id: str, trainee_id: str):
    data = load_data()
    s = find_session(data, session_id)
    if not s:
        abort(404)

    trainees = _session_trainees_list(s)
    t = next((x for x in trainees if x.get("id") == trainee_id), None)
    if not t:
        abort(404)

    docs = t.get("documents") or []
    buf = BytesIO()

    with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as z:
        for d in docs:
            tokens = []

            # ✅ multi-fichiers en priorité
            if isinstance(d.get("files"), list) and d["files"]:
                tokens = [x for x in d["files"] if x]
            else:
                # compat: 1 fichier
                tok = (d.get("file") or "")
                if tok:
                    tokens = [tok]

            if not tokens:
                continue

            label = (d.get("label") or "document").replace("/", "-")
            prenom = (t.get("first_name") or "").strip()
            nom = (t.get("last_name") or "").strip()

            for i, token in enumerate(tokens, start=1):
                fp = _detokenize_path(token)
                if not os.path.exists(fp):
                    continue

                ext = os.path.splitext(fp)[1] or ""
                base = f"{label} {prenom} {nom}".strip().replace("  ", " ")

                # ✅ si plusieurs fichiers: suffixe _1, _2...
                arc = (base + ext) if len(tokens) == 1 else (f"{base}_{i}{ext}")
                z.write(fp, arcname=arc)

    buf.seek(0)
    zipname = f"Documents_{t.get('first_name','')}_{t.get('last_name','')}.zip".replace(" ", "_")
    return send_file(buf, as_attachment=True, download_name=zipname, mimetype="application/zip")

# =========================
# API docs autosave (status/comment)
# =========================
@app.post("/api/sessions/<session_id>/stagiaires/<trainee_id>/documents/update")
@admin_login_required
@admin_write_required
def api_docs_update(session_id: str, trainee_id: str):
    data = load_data()
    s = find_session(data, session_id)
    if not s:
        return jsonify({"ok": False, "error": "session_not_found"}), 404

    trainees = _session_trainees_list(s)
    t = next((x for x in trainees if x.get("id") == trainee_id), None)
    if not t:
        return jsonify({"ok": False, "error": "trainee_not_found"}), 404

    payload = request.get_json(silent=True) or {}
    doc_key = payload.get("key")
    field = payload.get("field")
    value = payload.get("value")

    if field not in ("status", "comment"):
        return jsonify({"ok": False, "error": "invalid_field"}), 400

    docs = t.get("documents") or []
    for d in docs:
        if d.get("key") == doc_key:
            d[field] = value
            break

    t["updated_at"] = _now_iso()

    # ✅ Synchronisation automatique du statut dossier
    training_type = _session_get(s, "training_type", "")
    t["dossier_status"] = "complete" if dossier_is_complete_total(t, training_type) else "incomplete"

    # ✅ PERSISTENCE (sinon ça se perd au refresh)
    s["trainees"] = trainees
    s.pop("stagiaires", None)
    save_data(data)

    return jsonify({
        "ok": True,
        "dossier_is_complete": dossier_is_complete_total(t, training_type),
        "dossier_status": t["dossier_status"]
    })

# =========================
# Deliverables required (fin de formation)
# =========================
DELIVERABLE_LABELS = {
    "carte_sst": "Carte SST",
    "diplome": "Diplôme",
    "attestation_fin_formation": "Attestation fin de formation",
    "attestation_recevabilite": "Attestation de recevabilité VAE",
}

DELIVERABLE_REQUIRED_KEYS = ["diplome", "carte_sst", "attestation_fin_formation"]

VAE_STATUS_STEPS = {
    "livret_1_todo": {"label": "Livret 1 à compléter", "pill": "orange"},
    "livret_1_analysis": {"label": "Livret 1 en cours d'analyse", "pill": "gray"},
    "livret_1_validated": {"label": "Livret 1 validé", "pill": "green"},
    "financement_validated": {"label": "Financement validé", "pill": "green"},
    "livret_2_todo": {"label": "Livret 2 à compléter", "pill": "orange"},
    "livret_2_analysis": {"label": "Réception livret 2", "pill": "gray"},
    "livret_2_validated": {"label": "Livret 2 validé", "pill": "green"},
    "jury": {"label": "Passage devant le jury", "pill": "yellow"},
    "certified": {"label": "Diplôme obtenu", "pill": "green"},
}


def vae_status_view(status_key: Optional[str]) -> Dict[str, str]:
    key = (status_key or "").strip()
    if key not in VAE_STATUS_STEPS:
        key = "livret_1_todo"
    step = VAE_STATUS_STEPS[key]
    return {"key": key, "label": step["label"], "pill": step["pill"]}


def _notify_vae_status_change(t: Dict[str, Any], status_key: str) -> None:
    status_key = (status_key or "").strip()
    email = (t.get("email") or "").strip()
    first_name = (t.get("first_name") or "").strip() or "Madame, Monsieur"
    trainee_link = (t.get("public_link") or "").strip()

    subject = ""
    html = ""
    primary_btn = "display:inline-block;background:#1f8f4a;color:#ffffff;padding:12px 20px;border-radius:10px;text-decoration:none;font-weight:700"
    secondary_btn = "display:inline-block;background:#ffffff;color:#1f8f4a;padding:12px 20px;border:1px solid #1f8f4a;border-radius:10px;text-decoration:none;font-weight:700"
    public_token = (t.get("public_token") or "").strip()
    space_url = trainee_link or (f"{PUBLIC_STUDENT_PORTAL_BASE.rstrip('/')}/espace/{public_token}" if public_token else f"{PUBLIC_STUDENT_PORTAL_BASE.rstrip('/')}")
    booking_url = "https://calendly.com/integraleacademy/dirigeant"

    if status_key == "livret_1_analysis":
        subject = "Réception de votre Livret 1✅"
        html = mail_layout(f"""
        <h2 style="margin:0 0 12px 0;color:#0f172a;text-align:center;">✅ Réception de votre Livret 1</h2>
        <p>Bonjour <strong>{first_name}</strong>,</p>
        <p>Nous vous informons que nous avons bien reçu votre Livret 1 dans le cadre de votre VAE Dirigeant d'entreprise de sécurité privée (DESP).</p>
        <p>Votre dossier de faisabilité est désormais transmis à la commission pour étude. Cette étape permet de vérifier la conformité de votre demande et l'adéquation de votre expérience avec le référentiel de la certification visée.</p>
        <p>Notre équipe reviendra vers vous dès que l'analyse sera finalisée. Si votre dossier est recevable, vous recevrez :</p>
        <p>1️⃣ votre attestation de recevabilité,<br>
        2️⃣ les consignes pour démarrer la rédaction du Livret 2,<br>
        3️⃣ la suite de votre parcours VAE étape par étape.</p>
        <p>En attendant, vous pouvez continuer à consulter votre espace candidat pour suivre l'avancement de votre dossier.</p>
        <p style="margin-top:18px;text-align:center;"><a href="{space_url}" style="{secondary_btn}">Accéder à mon espace candidat</a></p>
        <p>Nous vous souhaitons une bonne journée et nous restons à votre disposition, la Team Intégrale Academy</p>
        """)
    elif status_key == "livret_1_validated":
        subject = "Livret 1 validé par la commission 🎉"
        html = mail_layout(f"""
        <h2 style=\"margin:0 0 12px 0;color:#0f172a;text-align:center;\">Votre Livret 1 est validé 🥳</h2>
        <p>Bonjour <strong>{first_name}</strong>,</p>
        <p>Nous avons le plaisir de vous informer que la commission a rendu un <strong>avis favorable</strong> à votre demande de VAE.</p>
        <p>Pour passer à l'étape suivante, nous devons organiser un rendez-vous téléphonique afin de :</p>
        <ul style=\"list-style:none;padding-left:0;margin:0;\">
          <li>1️⃣ Mettre en place le financement de votre VAE (versement de l'acompte),</li>
          <li>2️⃣ Finaliser et signer votre convention de VAE,</li>
          <li>3️⃣ Vous transmettre le cadre de travail pour la rédaction de votre Livret 2.</li>
        </ul>
        <p style=\"margin:18px 0;text-align:center;\"><a href=\"{booking_url}\" style=\"{primary_btn}\">Réserver un RDV téléphonique</a></p>
        <p>Vous pouvez également récupérer votre <strong>attestation de recevabilité</strong> directement dans votre espace candidat :</p>
        <p style=\"margin-top:10px;text-align:center;\"><a href=\"{space_url}\" style=\"{secondary_btn}\">Aller vers mon espace candidat</a></p>
        <p style=\"margin-top:18px;\">Notre équipe reste disponible si vous souhaitez être accompagné(e) dans la préparation de votre Livret 2.</p>
        <p>Nous vous souhaitons une excellente journée.<br>L'équipe Intégrale Academy</p>
        """)
    elif status_key == "livret_2_analysis":
        subject = "Transmission du Livret 2"
        html = mail_layout(f"""
        <h2 style=\"margin:0 0 12px 0;color:#0f172a;text-align:center;\">✅ Réception de votre Livret 2</h2>
        <p>Bonjour <strong>{first_name}</strong>,</p>
        <p>Nous avons bien reçu votre Livret 2 dans le cadre de votre VAE Dirigeant d'entreprise de sécurité privée (DESP).</p>
        <p>Votre dossier est à présent en cours d'étude par la commission, qui va en vérifier la conformité et la cohérence au regard du référentiel de certification.</p>
        <p>Dès que l'analyse sera terminée, nous reviendrons vers vous. Si tout est conforme, nous pourrons passer à la dernière étape de votre parcours : votre passage devant le <strong>jury de certification</strong>.</p>
        <p>Vous pouvez suivre votre progression à tout moment depuis votre espace candidat.</p>
        <p style=\"margin-top:18px;text-align:center;\"><a href=\"{space_url}\" style=\"{secondary_btn}\">Suivre mon dossier VAE</a></p>
        <p>Nous vous souhaitons une excellente journée.<br>L'équipe Intégrale Academy</p>
        """)
    elif status_key == "livret_2_validated":
        subject = "Livret 2 VAE dirigeant validé 🎉"
        html = mail_layout(f"""
        <h2 style=\"margin:0 0 12px 0;color:#0f172a;text-align:center;\">Votre Livret 2 VAE est validé 🎉</h2>
        <p>Bonjour <strong>{first_name}</strong>,</p>
        <p>Nous avons le plaisir de vous informer que votre <strong>Livret 2 VAE Dirigeant d'entreprise de sécurité (DESP)</strong> est conforme.</p>
        <p>Nous allons désormais programmer votre date de passage devant le jury de certification.</p>
        <p>Afin d'organiser votre passage et finaliser le financement de votre VAE, nous vous remercions de bien vouloir réserver un rendez-vous téléphonique avec notre équipe.</p>
        <p style=\"margin-top:18px;text-align:center;\"><a href=\"{booking_url}\" style=\"{primary_btn}\">Réserver un RDV téléphonique</a></p>
        <p style=\"margin-top:18px;\">Lors de cet échange, nous vous préciserons les prochaines étapes administratives et pratiques jusqu'à l'obtention de votre certification dirigeant.</p>
        <p>Nous restons à votre disposition et nous vous souhaitons une agréable journée,</p>
        <p><strong>La Team Intégrale Academy</strong></p>
        """)
    elif status_key == "jury":
        subject = "VAE : date de passage devant le jury"
        jury_date_iso = (t.get("vae_jury_date") or "").strip()
        jury_date = jury_date_iso
        if jury_date_iso and re.match(r"^\d{4}-\d{2}-\d{2}$", jury_date_iso):
            y, m, d = jury_date_iso.split("-")
            jury_date = f"{d}/{m}/{y}"
        html = mail_layout(f"""
        <h2 style=\"margin:0 0 12px 0;color:#0f172a;text-align:center;\">📅 Votre date d'examen VAE Dirigeant (DESP)</h2>
        <p>Bonjour <strong>{first_name}</strong>,</p>
        <p>Nous revenons vers vous concernant votre passage devant le jury de certification.</p>
        <p>Votre examen est planifié le <strong>{jury_date or 'DD/MM/YYYY'}</strong>.</p>
        <p>Nous vous communiquerons prochainement toutes les informations utiles : horaires, modalités de passage, documents à prévoir et consignes pratiques.</p>
        <p>En attendant, n'hésitez pas à consulter votre espace candidat pour suivre votre dossier.</p>
        <p style=\"margin-top:18px;text-align:center;\"><a href=\"{space_url}\" style=\"{secondary_btn}\">Ouvrir mon espace candidat</a></p>
        """)

    if not subject:
        print(f"[VAE][EMAIL] status inconnu, aucun envoi déclenché: status={status_key!r}")
        return

    if not email:
        trainee_id = str(t.get("id") or "")
        print(f"[VAE][EMAIL] aucun email stagiaire, envoi ignoré: trainee_id={trainee_id!r} status={status_key!r}")
        return

    email_ok = brevo_send_email(email, subject, html)
    sent_at = _now_iso()

    phone_followups = t.get("phone_followups")
    if not isinstance(phone_followups, list):
        phone_followups = []
    status_label = vae_status_view(status_key)["label"]
    phone_followups.insert(0, {
        "type": "Suivi VAE",
        "details": f"Mail VAE - {status_label}",
        "comment": f"Objet : {subject} · Envoi {'confirmé' if email_ok else 'tenté'}",
        "at": sent_at,
    })
    t["phone_followups"] = phone_followups

    trainee_id = str(t.get("id") or "")
    print(
        f"[VAE][EMAIL] envoi statut VAE: trainee_id={trainee_id!r} status={status_key!r} "
        f"to={email!r} ok={bool(email_ok)}"
    )

def deliverables_progress(t: Dict[str, Any]):
    """
    Retourne (done, total, is_complete) pour les 3 deliverables.
    done = nb de fichiers présents dans t['deliverables'] pour les clés attendues.
    """
    dv = t.get("deliverables") or {}
    done = 0
    for k in DELIVERABLE_REQUIRED_KEYS:
        tok = (dv.get(k) or "").strip()
        if tok:
            done += 1
    total = len(DELIVERABLE_REQUIRED_KEYS)
    return done, total, (done == total)


@app.post("/admin/sessions/<session_id>/stagiaires/<trainee_id>/deliverables/<kind>/upload")
@admin_login_required
@admin_write_required
def admin_upload_deliverable(session_id: str, trainee_id: str, kind: str):
    if kind not in DELIVERABLE_LABELS:
        abort(404)

    data = load_data()
    s = find_session(data, session_id)
    if not s:
        abort(404)

    trainees = _session_trainees_list(s)
    t = next((x for x in trainees if x.get("id") == trainee_id), None)
    if not t:
        abort(404)

    f = request.files.get("file")
    if not f or not f.filename:
        return redirect(url_for("admin_trainee_page", session_id=session_id, trainee_id=trainee_id))

    try:
        stored = _store_file(session_id, trainee_id, "deliverables", f)
    except Exception:
        return redirect(url_for("admin_trainee_page", session_id=session_id, trainee_id=trainee_id))

    token = _tokenize_path(stored)

    t.setdefault("deliverables", {})
    t["deliverables"][kind] = token
    t["updated_at"] = _now_iso()

    link = f"{PUBLIC_STUDENT_PORTAL_BASE.rstrip('/')}/espace/{t.get('public_token','')}"
    label = DELIVERABLE_LABELS[kind]

    # =========================
    # ✅ Jolis mails + SMS
    # =========================
    first_name = (t.get("first_name") or "").strip() or "Madame, Monsieur"
    formation_type = formation_label(_session_get(s, "training_type", ""))
    dstart = fr_date(_session_get(s, "date_start", ""))
    dend = fr_date(_session_get(s, "date_end", ""))

    # ✅ type formation brut pour la logique CNAPS
    tt_raw = (_session_get(s, "training_type", "") or "").strip()
    tt = tt_raw.upper()

    extra_line = ""
    cnaps_block = ""

    if kind == "diplome":
        extra_line = "🎉 Félicitations ! Votre diplôme est maintenant disponible."

        # --- CNAPS (différent selon formation) ---
        if tt == "APS":
            cnaps_block = f"""
            <div style="background:#fff7ed;border:1px solid #fed7aa;border-radius:12px;padding:14px;margin:16px 0">
              <p style="margin:0 0 8px 0;font-weight:900;color:#9a3412">🛡️ Carte professionnelle – Information importante</p>
              <p style="margin:0;color:#7c2d12;line-height:1.55">
                <strong>Vous n'avez aucune démarche à effectuer pour votre carte professionnelle.</strong>
                Votre diplôme a été automatiquement transmis au CNAPS qui procède actuellement à une enquête administrative.
                Dès que l'enquête sera terminée, vous recevrez votre carte professionnelle directement chez vous par courrier postal.
                <br><br>
                <strong>Pour rappel, vous ne pouvez pas exercer la profession tant que vous n'avez pas reçu votre carte professionnelle.</strong>
              </p>
            </div>
            """
        elif tt == "A3P":
            cnaps_block = f"""
            <div style="background:#eff6ff;border:1px solid #bfdbfe;border-radius:12px;padding:14px;margin:16px 0">
              <p style="margin:0 0 8px 0;font-weight:900;color:#1d4ed8">🛡️ Demande de carte professionnelle (CNAPS)</p>
              <p style="margin:0;color:#1e3a8a;line-height:1.55">
                Vous pouvez à présent procéder à la demande de carte professionnelle depuis l'espace Téléservices du CNAPS.
                <br><br>
                Si vous êtes déjà agent de sécurité, cliquez sur <strong>"Ma demande concerne une extension de carte professionnelle"</strong>.<br>
                Si vous n'êtes pas agent de sécurité, cliquez sur <strong>"Ma demande concerne une carte professionnelle"</strong>.
                <br><br>
                Dans les deux cas, complétez la rubrique <strong>"J'ai un NUB"</strong> en indiquant <strong>votre NOM</strong>
                (uniquement votre nom, pas votre prénom) et votre <strong>NUB</strong>
                (7 derniers chiffres de votre numéro d'autorisation préalable ou de votre carte professionnelle).
                Suivez les étapes et téléchargez les pièces justificatives : votre pièce d'identité, votre justificatif de domicile de moins de 3 mois et votre diplôme.
              </p>

              <p style="margin:12px 0 0 0;text-align:center">
                <a href="https://depot-teleservices-cnaps.interieur.gouv.fr/"
                   style="display:inline-block;background:#1d4ed8;color:white;padding:12px 16px;border-radius:10px;text-decoration:none;font-weight:900">
                  👉 Demander ma carte professionnelle CNAPS
                </a>
              </p>
            </div>
            """
        elif "DIRIGEANT" in tt:
            cnaps_block = f"""
            <div style="background:#f0fdf4;border:1px solid #bbf7d0;border-radius:12px;padding:14px;margin:16px 0">
              <p style="margin:0 0 8px 0;font-weight:900;color:#166534">🏛️ Agrément dirigeant (CNAPS)</p>
              <p style="margin:0;color:#14532d;line-height:1.55">
                Vous pouvez à présent procéder à votre demande d'agrément dirigeant directement depuis le site internet du CNAPS
                en complétant le formulaire en cliquant ci-dessous.
              </p>

              <p style="margin:12px 0 0 0;text-align:center">
                <a href="https://www.cnaps.interieur.gouv.fr/Demarches-en-ligne/Vous-etes-un-particulier/Diriger-une-entreprise-de-securite-privee-un-organisme-de-formation-un-service-interne-de-securite/Diriger-un-organisme-de-formation-une-entreprise-de-securite-privee-un-service-interne-de-securite"
                   style="display:inline-block;background:#16a34a;color:white;padding:12px 16px;border-radius:10px;text-decoration:none;font-weight:900">
                  👉 Faire ma demande d’agrément dirigeant
                </a>
              </p>
            </div>
            """

    elif kind == "attestation_fin_formation":
        extra_line = "📄 Votre attestation de fin de formation est disponible et peut être téléchargée à tout moment."
    elif kind == "carte_sst":
        extra_line = (
            "🩺 Votre carte SST est disponible sur votre espace en ligne. "
            "Nous vous remettrons également un exemplaire papier en main propre "
            "(attention : aucun duplicata ne sera délivré). "
            "Conservez-la précieusement, elle peut être demandée par un employeur."
        )


    subject = f"{label} disponible – Intégrale Academy"

    # ✅ Bloc avis Google (tous les cas)
    google_block = """
      <div style="background:#fff;border:1px solid #e5e7eb;border-radius:12px;padding:14px;margin:16px 0">
        <p style="margin:0 0 8px 0;font-weight:900">⭐ Un petit service (1 minute)</p>
        <p style="margin:0;color:#374151;line-height:1.55">
          Si la formation vous a été utile, votre avis aide énormément les futurs stagiaires à choisir une école sérieuse
          et nous permet d’améliorer encore notre accompagnement.
        </p>
        <p style="margin:12px 0 0 0;text-align:center">
          <a href="https://g.page/r/CZ0Ug-feyXjHEAE"
             style="display:inline-block;background:#f59e0b;color:#111827;padding:12px 16px;border-radius:10px;text-decoration:none;font-weight:900">
            👉 Laisser un avis Google
          </a>
        </p>
      </div>
    """

    html = mail_layout(f"""
      <h2 style="text-align:center">✅ {label} disponible</h2>

      <p>Bonjour <strong>{first_name}</strong>,</p>

      <p>
        Nous avons le plaisir de vous informer que votre <strong>{label}</strong>
        est désormais disponible dans votre espace stagiaire.
      </p>

      {"<p style='margin-top:10px;font-weight:700'>" + extra_line + "</p>" if extra_line else ""}

      {cnaps_block}

      <div style="background:#f3f4f6;border:1px solid #e5e7eb;border-radius:12px;padding:14px;margin:16px 0">
        <p style="margin:0 0 10px 0">
          <strong>📌 Formation :</strong> {formation_type}
          {" — <strong>Dates :</strong> " + dstart + " au " + dend if (dstart or dend) else ""}
        </p>

        <p style="margin:0">
          <strong>📍 Accéder à votre espace stagiaire :</strong><br>
          <a href="{link}" style="color:#1f8f4a;text-decoration:none;font-weight:bold">{link}</a>
        </p>
      </div>

      <p style="text-align:center;margin-top:18px">
        <a href="{link}"
           style="display:inline-block;background:#1f8f4a;color:white;padding:12px 18px;border-radius:10px;
                  text-decoration:none;font-weight:bold">
          👉 Accéder à mon espace stagiaire
        </a>
      </p>

      {google_block}

      <p style="margin-top:22px">
        Pour toute question, vous pouvez nous contacter au <strong>04 22 47 07 68</strong>.
      </p>

      <p style="margin-top:22px">
        Bien cordialement,<br>
        <strong>Clément VAILLANT</strong><br>
        Directeur Intégrale Academy
      </p>

      <hr style="margin:26px 0;border:none;border-top:1px solid #e5e7eb">

      <p style="font-size:12px;color:#6b7280;text-align:center;line-height:1.6">
        © Intégrale Academy — Merci de votre confiance 💛<br>
        54 chemin du Carreou 83480 PUGET SUR ARGENS / 142 rue de Rivoli 75001 PARIS<br>
        <a href="https://www.integraleacademy.com"
           style="color:#1f8f4a;text-decoration:none;font-weight:bold">
          integraleacademy.com
        </a>
      </p>
    """)

    sms_name = (t.get("first_name") or "").strip()
    sms = (
        f"Intégrale Academy ✅ {sms_name + ', ' if sms_name else ''}\n"
        f"Votre {label} est disponible sur votre espace :\n"
        f"{link}\n"
        f"A bientôt, la Team Intégrale Academy"
)

    if kind != "attestation_recevabilite":
        brevo_send_email(t.get("email", ""), subject, html)
        brevo_send_sms(t.get("phone", ""), sms)

    # ✅ persistance
    s["trainees"] = trainees
    s.pop("stagiaires", None)
    save_data(data)

    return redirect(url_for("admin_trainee_page", session_id=session_id, trainee_id=trainee_id))

def find_session_and_trainee_by_token(data: Dict[str, Any], token: str):
    token = (token or "").strip()
    if not token:
        return None, None

    sessions = data.get("sessions", []) or []
    for s in sessions:
        trainees = s.get("trainees") or s.get("stagiaires") or []
        for t in trainees:
            public_token = (t.get("public_token") or "").strip()
            legacy_token = (t.get("token") or "").strip()
            if public_token == token or legacy_token == token:
                return s, t
    return None, None


def _vae_extract_trainee_token_from_referer(referer: str) -> str:
    raw = str(referer or "").strip()
    if not raw:
        return ""

    marker = "/espace/"
    idx = raw.find(marker)
    if idx < 0:
        return ""

    tail = raw[idx + len(marker):]
    token = tail.split("?", 1)[0].split("#", 1)[0].split("/", 1)[0].strip()
    return token


@app.get("/espace/<token>")
def public_trainee_space(token):
    data = load_data()
    s, t = find_session_and_trainee_by_token(data, token)

    if not s or not t:
        abort(404)

    # 🔒 Verrou public : si pas auth → login
    if not _public_is_authed(token):
        return redirect(url_for("public_trainee_login", token=token))

    if not session.get("admin_logged_in") and not t.get("public_has_logged_in"):
        _mark_public_login(data, s, t)

    training_type = _session_get(s, "training_type", "")

    # ✅ aligne la liste des docs requis
    ensure_documents_schema_for_trainee(t, training_type)
    if (training_type or "").strip().upper() == "DIRIGEANT VAE":
        _ensure_livret2_document_entry(t)

    for d in (t.get("documents") or []):
        file_token = d.get("file") or ""
        d["file_token"] = file_token
        files = d.get("files")
        file_tokens = [x for x in files if x] if isinstance(files, list) else []
        if file_token and file_token not in file_tokens:
            file_tokens.insert(0, file_token)
        d["file_tokens"] = file_tokens

    show_hosting = ((training_type or "").strip().upper() == "A3P")
    show_vae = ("VAE" in (training_type or "").upper())
    show_vtc = ("VTC" in (training_type or "").upper())

    # ✅ persistance
    s["trainees"] = _session_trainees_list(s)
    s.pop("stagiaires", None)
    save_data(data)

    return render_template(
        "public_trainee.html",
        session=s,
        trainee=t,
        token=token,
        show_hosting=show_hosting,
        show_vae=show_vae,
        show_vtc=show_vtc,
        dossier_ok=dossier_is_complete_total(t, training_type),
        vae_required_docs_deposited=required_docs_are_deposited(t, training_type),
    )
    

@app.post("/admin/sessions/<session_id>/stagiaires/<trainee_id>/identity-photo/upload")
@admin_login_required
@admin_write_required
def admin_upload_identity_photo(session_id: str, trainee_id: str):
    data = load_data()
    s = find_session(data, session_id)
    if not s:
        abort(404)

    trainees = _session_trainees_list(s)
    t = next((x for x in trainees if x.get("id") == trainee_id), None)
    if not t:
        abort(404)

    f = request.files.get("file")
    if not f or not f.filename:
        return redirect(url_for("admin_trainee_page", session_id=session_id, trainee_id=trainee_id))

    # ✅ on limite aux images
    ext = _safe_ext(f.filename)
    if ext not in (".jpg", ".jpeg", ".png", ".webp"):
        return redirect(url_for("admin_trainee_page", session_id=session_id, trainee_id=trainee_id))

    try:
        stored = _store_file(session_id, trainee_id, "identity_photo", f)
    except Exception:
        return redirect(url_for("admin_trainee_page", session_id=session_id, trainee_id=trainee_id))

    token = _tokenize_path(stored)

    # ✅ on sauvegarde le token dans le stagiaire
    t["identity_photo"] = token
    t["updated_at"] = _now_iso()

    s["trainees"] = trainees
    s.pop("stagiaires", None)
    save_data(data)

    return redirect(url_for("admin_trainee_page", session_id=session_id, trainee_id=trainee_id))


@app.post("/espace/<token>/infos/update")
def public_infos_update(token: str):
    data = load_data()
    s, t = find_session_and_trainee_by_token(data, token)
    if not s or not t:
        return jsonify({"ok": False}), 404

    payload = request.get_json(silent=True) or {}

    # champs autorisés (sécurité)
    allowed = {
        "carte_vitale",
        "pre_number",
        "birth_date",
        "birth_city",
        "birth_country",
        "nationality",
        "address",
        "zip_code",
        "city",
        "no_permis",
    }

    previous_vae_status = vae_status_view(t.get("vae_status"))["key"]

    for k, v in payload.items():
        if k not in allowed:
            continue

        # no_permis = bool
        if k == "no_permis":
            t["no_permis"] = bool(v)
            continue

        # strings : on n'écrase PAS avec vide
        if v is None:
            continue
        if isinstance(v, str):
            vv = v.strip()
            if vv == "":
                continue
            t[k] = vv
        else:
            # si jamais tu envoies autre chose
            t[k] = v

    training_type = _session_get(s, "training_type", "")
    t["dossier_status"] = "complete" if dossier_is_complete_total(t, training_type) else "incomplete"
    t["updated_at"] = _now_iso()

    # ✅ IMPORTANT : persister la session normalisée comme ailleurs
    s["trainees"] = _session_trainees_list(s)
    s.pop("stagiaires", None)
    save_data(data)

    return jsonify({"ok": True})


@app.post("/espace/<token>/vtc-credentials")
def public_vtc_credentials(token: str):
    data = load_data()
    s, t = find_session_and_trainee_by_token(data, token)
    if not s or not t:
        return jsonify({"ok": False}), 404

    if not _public_is_authed(token):
        return jsonify({"ok": False, "error": "not_authenticated"}), 403

    payload = request.get_json(silent=True) or {}
    login = (payload.get("login") or "").strip()
    password = (payload.get("password") or "").strip()

    if not login or not password:
        return jsonify({"ok": False, "error": "missing_credentials"}), 400

    t["vtc_cm_login"] = login
    t["vtc_cm_password"] = password
    t["vtc_cm_submitted_at"] = _now_iso()
    t["updated_at"] = _now_iso()

    trainee_display_name = _format_trainee_name(t.get("first_name", ""), t.get("last_name", ""))
    add_admin_notification(
        data,
        f"🟢{trainee_display_name} - Identifiants chambre des métiers VTC envoyés",
        meta={
            "type": "vtc_credentials",
            "session_id": s.get("id"),
            "trainee_id": t.get("id"),
            "comment": (t.get("comment") or "").strip(),
        },
    )

    s["trainees"] = _session_trainees_list(s)
    s.pop("stagiaires", None)
    save_data(data)

    trainee_name = f"{t.get('first_name','').strip()} {t.get('last_name','').strip()}".strip()
    session_name = _session_get(s, "name", "")
    subject = "Identifiants examen VTC transmis"
    html = mail_layout(f"""
      <h2>Identifiants VTC transmis</h2>
      <p><strong>Stagiaire :</strong> {trainee_name or '—'}</p>
      <p><strong>Session :</strong> {session_name or '—'}</p>
      <p><strong>Email :</strong> {t.get("email") or "—"}</p>
      <p><strong>Téléphone :</strong> {t.get("phone") or "—"}</p>
      <hr style="margin:16px 0;border:none;border-top:1px solid #e5e7eb">
      <p><strong>Login :</strong> {login}</p>
      <p><strong>Mot de passe :</strong> {password}</p>
    """)
    brevo_send_email("clement@integraleacademy.com", subject, html)

    return jsonify({"ok": True})



@app.post("/espace/<token>/documents/<doc_key>/upload")
def public_doc_upload(token: str, doc_key: str):
    data = load_data()
    s, t = find_session_and_trainee_by_token(data, token)
    if not s or not t:
        abort(404)

    training_type = _session_get(s, "training_type", "")
    ensure_documents_schema_for_trainee(t, training_type)
    if (training_type or "").strip().upper() == "DIRIGEANT VAE":
        _ensure_livret2_document_entry(t)

    # ✅ doc_key doit être dans la liste requise
    if doc_key not in allowed_doc_keys_for_training(training_type):
        return redirect(url_for("public_trainee_space", token=token))

    # ✅ retrouver la config du doc (accept)
    docs = t.get("documents") or []
    target = next((d for d in docs if d.get("key") == doc_key), None)
    if not target:
        return redirect(url_for("public_trainee_space", token=token))

    # ✅ 1 fichier par envoi (mais l'ID peut en contenir 2)
    incoming_files = request.files.getlist("files") or request.files.getlist("file")
    incoming_files = [f for f in incoming_files if f and f.filename]
    if not incoming_files:
        if doc_key == "livret_2":
            existing_files = target.get("files") if isinstance(target.get("files"), list) else []
            legacy_file = (target.get("file") or "").strip()
            has_existing_livret_2 = bool(existing_files or legacy_file)
            if has_existing_livret_2:
                if not isinstance(t.get("vae_action_dates"), dict):
                    t["vae_action_dates"] = {}
                t["vae_action_dates"]["livret_2_received"] = datetime.date.today().strftime("%d/%m/%Y")
                view = vae_status_view("livret_2_analysis")
                previous_status = (t.get("vae_status") or "").strip()
                t["vae_status"] = view["key"]
                t["vae_status_label"] = view["label"]
                if previous_status != view["key"]:
                    _notify_vae_status_change(t, "livret_2_analysis")
                    trainee_display_name = _format_trainee_name(t.get("first_name", ""), t.get("last_name", ""))
                    add_admin_notification(
                        data,
                        f"VAE Livret 2️⃣ Déposé par {trainee_display_name}",
                        meta={
                            "type": "vae_livret_2_upload",
                            "session_id": s.get("id"),
                            "trainee_id": t.get("id"),
                        },
                    )
                t["updated_at"] = _now_iso()
                t["dossier_status"] = "complete" if dossier_is_complete_total(t, training_type) else "incomplete"
                s["trainees"] = _session_trainees_list(s)
                s.pop("stagiaires", None)
                save_data(data)
        return redirect(url_for("public_trainee_space", token=token))

    accept = (target.get("accept") or "").lower()

    def _accepts_file(ext: str) -> bool:
        acc = [a.strip().lower() for a in accept.split(",") if a.strip()]
        allowed_exts = set()
        if "application/pdf" in acc:
            allowed_exts.add(".pdf")
        if any(a.startswith("image/") for a in acc) or ("image/jpeg" in acc) or ("image/png" in acc):
            allowed_exts.update({".jpg", ".jpeg", ".png", ".webp"})
        if allowed_exts:
            return ext in allowed_exts
        return ext in ALLOWED_EXT

    # ✅ stockage du fichier
    session_id = s.get("id")
    trainee_id = t.get("id")

    # ✅ MAJ du doc: on APPEND dans files (sans écraser)
    original_name = ""
    for d in docs:
        if d.get("key") == doc_key:
            cur_files = d.get("files")
            if not isinstance(cur_files, list):
                cur_files = []

            # compat: si un ancien "file" existe mais pas dans files, on le garde
            old = (d.get("file") or "").strip()
            if old and old not in cur_files:
                cur_files.append(old)

            cur_status = (d.get("status") or "").strip().upper()
            if cur_status in ("NON CONFORME", "NON_CONFORME"):
                cur_files = []
            max_files = 2 if doc_key == "id" else (-1 if doc_key == "livret_2" else 1)
            if max_files < 0:
                files_to_store = incoming_files
            else:
                remaining_slots = max(max_files - len(cur_files), 0)
                files_to_store = incoming_files[:remaining_slots] if remaining_slots else []
            if not files_to_store:
                return redirect(url_for("public_trainee_space", token=token))

            for f in files_to_store:
                ext = _safe_ext(f.filename)
                if not _accepts_file(ext):
                    return redirect(url_for("public_trainee_space", token=token))

            for f in files_to_store:
                if not original_name:
                    original_name = secure_filename(f.filename or "document")
                stored = _store_file(session_id, trainee_id, "public_documents", f)
                new_token = _tokenize_path(stored)
                cur_files.append(new_token)

            # on garde le premier fichier dans "file" (pour compat template/admin)
            d["files"] = cur_files
            d["file"] = cur_files[0] if cur_files else ""

            cur = (d.get("status") or "").strip().upper()
            if cur in ("", "NON DÉPOSÉ", "NON DEPOSE", "NON_DEPOSE") or cur_status in ("NON CONFORME", "NON_CONFORME"):
                d["status"] = "A CONTRÔLER"
            if d.get("status") == "A CONTROLER":
                d["status"] = "A CONTRÔLER"
            break

    t["updated_at"] = _now_iso()
    t["dossier_status"] = "complete" if dossier_is_complete_total(t, training_type) else "incomplete"

    # ✅ persistance
    s["trainees"] = _session_trainees_list(s)
    s.pop("stagiaires", None)
    save_data(data)

    if doc_key == "livret_2":
        if not isinstance(t.get("vae_action_dates"), dict):
            t["vae_action_dates"] = {}
        t["vae_action_dates"]["livret_2_received"] = datetime.date.today().strftime("%d/%m/%Y")
        view = vae_status_view("livret_2_analysis")
        t["vae_status"] = view["key"]
        t["vae_status_label"] = view["label"]
        _notify_vae_status_change(t, "livret_2_analysis")
        trainee_display_name = _format_trainee_name(t.get("first_name", ""), t.get("last_name", ""))
        add_admin_notification(
            data,
            f"VAE Livret 2️⃣ Déposé par {trainee_display_name}",
            meta={
                "type": "vae_livret_2_upload",
                "session_id": s.get("id"),
                "trainee_id": t.get("id"),
            },
        )
        save_data(data)

    # ✅ IMPORTANT: on renvoie l’info au GET (pour popup + scroll ensuite)
    return redirect(url_for(
        "public_trainee_space",
        token=token,
        uploaded=doc_key,
        fname=original_name
    ))


@app.post("/espace/<token>/documents/candidate_info_sheet/validate")
def public_candidate_sheet_validate(token: str):
    data = load_data()
    s, t = find_session_and_trainee_by_token(data, token)
    if not s or not t:
        abort(404)

    if not _public_is_authed(token):
        return redirect(url_for("public_trainee_login", token=token))

    training_type = _session_get(s, "training_type", "")
    if (training_type or "").strip().upper() != "DIRIGEANT VAE":
        return redirect(url_for("public_trainee_space", token=token))

    ensure_documents_schema_for_trainee(t, training_type)

    docs = t.get("documents") or []
    target = next((d for d in docs if d.get("key") == "candidate_info_sheet"), None)
    if target:
        target["status"] = "A CONTRÔLER"
        if target.get("status") == "A CONTROLER":
            target["status"] = "A CONTRÔLER"

    t["updated_at"] = _now_iso()
    t["dossier_status"] = "complete" if dossier_is_complete_total(t, training_type) else "incomplete"

    s["trainees"] = _session_trainees_list(s)
    s.pop("stagiaires", None)
    save_data(data)

    return redirect(url_for("public_trainee_space", token=token))




# =========================
# ✅ Remplace ta page JSON par une vraie page HTML
# =========================
@app.get("/admin/sessions/<session_id>/stagiaires/<trainee_id>")
@admin_login_required
def admin_trainee_page(session_id: str, trainee_id: str):
    data = load_data()
    s = find_session(data, session_id)
    if not s:
        abort(404)

    session_view = {
        "id": s.get("id"),
        "name": _session_get(s, "name", ""),
        "training_type": _session_get(s, "training_type", ""),
        "date_start": _session_get(s, "date_start", ""),
        "date_end": _session_get(s, "date_end", ""),
        "exam_date": _session_get(s, "exam_date", ""),
        "exam_theory_date": _session_get(s, "exam_theory_date", ""),
        "exam_practice_date": _session_get(s, "exam_practice_date", ""),
    }

    trainees = _session_trainees_list(s)
    t = next((x for x in trainees if x.get("id") == trainee_id), None)
    if not t:
        abort(404)

    training_type = session_view["training_type"]
    default_price = default_training_price(training_type)

    # ✅ IMPORTANT : on impose la liste de documents selon la formation (et supprime dom)
    ensure_documents_schema_for_trainee(t, training_type)

    # ✅ deliverables
    t.setdefault("deliverables", {})

    # file tokens for template links (documents)
    for d in (t.get("documents") or []):
        # compat: 1 fichier
        token = d.get("file") or ""
        d["file_token"] = token

        # ✅ multi-fichiers
        files = d.get("files")
        file_tokens = [x for x in files if x] if isinstance(files, list) else []
        if token and token not in file_tokens:
            file_tokens.insert(0, token)
        d["file_tokens"] = file_tokens

    # deliverables view
    deliverables_view = []
    for k, label in DELIVERABLE_LABELS.items():
        token = (t.get("deliverables", {}) or {}).get(k, "")
        deliverables_view.append({
            "key": k,
            "label": label,
            "file": token,
            "file_token": token,
        })

    show_vae = (training_type == "DIRIGEANT VAE")
    vae_steps = [{"key": k, "label": v["label"], "pill": v["pill"]} for k, v in VAE_STATUS_STEPS.items()]
    vae_dossier = _vae_find_latest_for_trainee(str(t.get('id') or '')) if show_vae else None
    t["vae_status"] = vae_status_view(t.get("vae_status") or t.get("vae_status_label"))["key"]
    t["vae_status_label"] = vae_status_view(t.get("vae_status"))["label"]
    if not isinstance(t.get("vae_action_dates"), dict):
        t["vae_action_dates"] = {}

    # ✅ s'assure que no_permis est bien un bool
    t["no_permis"] = bool(t.get("no_permis"))
    t["force_dossier_complete"] = bool(t.get("force_dossier_complete"))
    if not (str(t.get("training_price") or "").strip()) and default_price is not None:
        t["training_price"] = default_price

    # ✅ dossier_status cohérent avec les docs requis
    dossier_complete = dossier_is_complete_total(t, training_type)
    t["dossier_status"] = "complete" if dossier_complete else "incomplete"
    t["updated_at"] = _now_iso()
    ensure_cnaps_history(t)

    # ✅ persistance
    s["trainees"] = trainees
    s.pop("stagiaires", None)
    save_data(data)

    trainee_history = build_trainee_history_entries(t)

    return render_template(
        "admin_trainee.html",
        session=session_view,
        trainee=t,
        show_vae=show_vae,
        vae_steps=vae_steps,
        vae_dossier=vae_dossier,
        dossier_is_complete=dossier_complete,
        deliverables_view=deliverables_view,
        default_training_price=default_price,
        trainee_history=trainee_history,
        PUBLIC_STUDENT_PORTAL_BASE=PUBLIC_STUDENT_PORTAL_BASE,
        fr_date=fr_date,
    )




@app.get("/admin/sessions/<session_id>/stagiaires/<trainee_id>/vtc-cm-autologin")
@admin_login_required
def admin_vtc_cm_autologin(session_id: str, trainee_id: str):
    data = load_data()
    s = find_session(data, session_id)
    if not s:
        abort(404)

    trainees = _session_trainees_list(s)
    t = next((x for x in trainees if x.get("id") == trainee_id), None)
    if not t:
        abort(404)

    login = (t.get("vtc_cm_login") or "").strip()
    password = (t.get("vtc_cm_password") or "").strip()
    trainee_name = f"{(t.get('first_name') or '').strip()} {(t.get('last_name') or '').strip()}".strip() or "stagiaire"

    if not login or not password:
        return f"""
<!doctype html>
<html lang="fr">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width,initial-scale=1">
  <title>Connexion auto Exament3P</title>
  <style>
    body {{ font-family: Arial, sans-serif; max-width: 760px; margin: 40px auto; padding: 0 16px; line-height: 1.45; }}
    .card {{ border:1px solid #e5e7eb; border-radius:14px; padding:18px; background:#fff; }}
    .warn {{ color:#b45309; font-weight:700; margin-bottom:10px; }}
    a {{ color:#1d4ed8; }}
  </style>
</head>
<body>
  <div class="card">
    <div class="warn">⚠️ Connexion automatique impossible</div>
    <p>Les identifiants Exament3P de <strong>{html.escape(trainee_name)}</strong> sont incomplets.</p>
    <p>Renseignez le login et le mot de passe dans l’espace stagiaire, puis réessayez.</p>
    <p><a href="{url_for('admin_trainee_page', session_id=session_id, trainee_id=trainee_id)}">← Retour à la fiche stagiaire</a></p>
  </div>
</body>
</html>
""", 400

    login_esc = html.escape(login)
    password_esc = html.escape(password)
    target_url = "https://www.exament3p.fr/id/14"

    return f"""
<!doctype html>
<html lang="fr">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width,initial-scale=1">
  <title>Connexion auto Exament3P</title>
  <style>
    body {{ font-family: Arial, sans-serif; max-width: 820px; margin: 40px auto; padding: 0 16px; line-height: 1.45; }}
    .card {{ border:1px solid #e5e7eb; border-radius:14px; padding:18px; background:#fff; }}
    .ok {{ color:#166534; font-weight:700; }}
    .hint {{ color:#6b7280; }}
    .btn {{ display:inline-block; margin-top:12px; border:1px solid #d1d5db; border-radius:8px; padding:8px 12px; text-decoration:none; color:#111827; background:#fff; cursor:pointer; }}
    .btn + .btn {{ margin-left:8px; }}
  </style>
</head>
<body>
  <div class="card">
    <div class="ok">✅ Tentative de connexion automatique à Exament3P en cours…</div>
    <p>Nous envoyons automatiquement les identifiants enregistrés pour <strong>{html.escape(trainee_name)}</strong>.</p>
    <p class="hint">Le site Exament3P utilise une modale JavaScript : si la connexion n’aboutit pas, cliquez sur « Relancer » puis « Se connecter » sur leur fenêtre.</p>

    <button class="btn" type="button" onclick="runAutoLogin()">Relancer la connexion auto</button>
    <a class="btn" href="{target_url}" target="_blank" rel="noopener">Ouvrir Exament3P manuellement</a>
  </div>

  <form id="autoLogin" method="post" action="{target_url}" style="display:none;">
    <input name="email" value="{login_esc}">
    <input name="loginEmail" value="{login_esc}">
    <input name="uac_email" value="{login_esc}">
    <input name="login" value="{login_esc}">
    <input name="username" value="{login_esc}">
    <input name="identifiant" value="{login_esc}">

    <input name="password" type="password" value="{password_esc}">
    <input name="loginPassword" type="password" value="{password_esc}">
    <input name="uac_password" type="password" value="{password_esc}">
    <input name="passwd" type="password" value="{password_esc}">
    <input name="mot_de_passe" type="password" value="{password_esc}">

    <input name="idpage" value="14">
    <input name="pageid" value="14">
    <input name="_remember_me" value="1">
    <input name="remember" value="1">
  </form>

  <script>
    function runAutoLogin() {{
      const form = document.getElementById('autoLogin');
      if (!form) {{
        window.location.href = '{target_url}';
        return;
      }}
      form.submit();
    }}

    setTimeout(runAutoLogin, 120);
  </script>
</body>
</html>
"""



@app.get("/admin/sessions/<session_id>/stagiaires/<trainee_id>/summary")
@admin_login_required
def admin_trainee_summary(session_id: str, trainee_id: str):
    data = load_data()
    s = find_session(data, session_id)
    if not s:
        abort(404)

    trainees = _session_trainees_list(s)
    t = next((x for x in trainees if x.get("id") == trainee_id), None)
    if not t:
        abort(404)

    training_name = (s.get("name") or "").strip() or formation_label(_session_get(s, "training_type", ""))
    dstart = fr_date(_session_get(s, "date_start", ""))
    dend = fr_date(_session_get(s, "date_end", ""))
    formation_dates = f"Du {dstart} au {dend}" if dstart and dend else "Dates à confirmer"

    session_view = {
        "id": s.get("id"),
        "name": _session_get(s, "name", ""),
        "training_type": _session_get(s, "training_type", ""),
        "date_start": _session_get(s, "date_start", ""),
        "date_end": _session_get(s, "date_end", ""),
    }

    return render_template(
        "admin_trainee_summary.html",
        session=session_view,
        trainee=t,
        training_name=training_name or "Formation",
        formation_dates=formation_dates,
    )


def _build_candidate_sheet_data(session_data: Dict[str, Any], trainee_data: Dict[str, Any]) -> Dict[str, Any]:
    dossier = _vae_find_latest_for_trainee(str(trainee_data.get("id") or ""))
    candidat = (dossier or {}).get("candidat") or {}

    def pick(*values: Any) -> str:
        for value in values:
            txt = str(value or "").strip()
            if txt:
                return txt
        return ""

    def clean_phone(value: Any) -> str:
        raw = re.sub(r"\D", "", str(value or ""))
        if len(raw) == 10:
            return " ".join(raw[i:i + 2] for i in range(0, len(raw), 2))
        return str(value or "").strip()

    payload = {
        "formation_type": "Validation des acquis de l'expérience (VAE)",
        "date_entree_stage": fr_date(_session_get(session_data, "date_start", "")) or "",
        "situation": pick(candidat.get("statut")),
        "company": pick(trainee_data.get("company_name"), trainee_data.get("employer"), trainee_data.get("entreprise")),
        "financing": "",
        "interviewer": "Clément VAILLANT",
        "last_name": pick(trainee_data.get("last_name"), candidat.get("nom_naissance")),
        "usage_name": pick(candidat.get("nom_usage")),
        "first_names": pick(trainee_data.get("first_name"), candidat.get("prenoms")),
        "address": pick(trainee_data.get("address"), candidat.get("adresse")),
        "postal_code": pick(trainee_data.get("zip_code"), trainee_data.get("postal_code")),
        "city": pick(trainee_data.get("city"), trainee_data.get("birth_city")),
        "phone": clean_phone(pick(trainee_data.get("phone"), candidat.get("telephone"))),
        "email": pick(trainee_data.get("email"), candidat.get("email")),
        "birth_date": fr_date(pick(trainee_data.get("birth_date"), candidat.get("date_naissance"))) or pick(trainee_data.get("birth_date"), candidat.get("date_naissance")),
        "birth_city": pick(trainee_data.get("birth_city"), trainee_data.get("birth_place")),
        "department": pick(trainee_data.get("department")),
        "country": pick(trainee_data.get("country"), "France"),
        "nationality": pick(trainee_data.get("nationality"), candidat.get("nationalite")),
        "emergency_contact": pick(trainee_data.get("emergency_contact"), trainee_data.get("emergency_phone")),
        "cnaps_number": "",
        "study_level": pick(candidat.get("niveau_formation")),
        "study_domain": pick(trainee_data.get("study_domain")),
        "last_certification_level": pick(candidat.get("niveau_certification")),
        "last_certification_domain": pick(trainee_data.get("last_certification_domain")),
        "last_job": pick(trainee_data.get("last_job")),
        "years_experience": pick(trainee_data.get("years_experience")),
        "company_name": pick(trainee_data.get("company_name"), trainee_data.get("employer")),
        "gross_salary": pick(trainee_data.get("gross_salary")),
    }

    saved_sheet = trainee_data.get("candidate_sheet")
    if isinstance(saved_sheet, dict):
        for key in payload.keys():
            if key in saved_sheet:
                payload[key] = str(saved_sheet.get(key) or "").strip()

    return payload


@app.post("/espace/<token>/fiche-candidat/enregistrer")
def public_trainee_candidate_sheet_save(token: str):
    data = load_data()
    s, t = find_session_and_trainee_by_token(data, token)
    if not s or not t:
        abort(404)

    if not _public_is_authed(token):
        return redirect(url_for("public_trainee_login", token=token))

    training_type = (_session_get(s, "training_type", "") or "").strip().upper()
    if training_type != "DIRIGEANT VAE":
        abort(404)

    base_sheet = _build_candidate_sheet_data(s, t)
    saved_sheet: Dict[str, str] = {}
    for key in base_sheet.keys():
        saved_sheet[key] = str(request.form.get(key, "") or "").strip()
    t["candidate_sheet"] = saved_sheet
    t["candidate_sheet_saved_at"] = _now_iso()

    ensure_documents_schema_for_trainee(t, training_type)

    t["updated_at"] = _now_iso()
    t["dossier_status"] = "complete" if dossier_is_complete_total(t, training_type) else "incomplete"

    s["trainees"] = _session_trainees_list(s)
    s.pop("stagiaires", None)
    save_data(data)

    return redirect(url_for("public_trainee_space", token=token) + "#doc_candidate_info_sheet")


@app.get("/espace/<token>/fiche-candidat")
def public_trainee_candidate_sheet(token: str):
    data = load_data()
    s, t = find_session_and_trainee_by_token(data, token)
    if not s or not t:
        abort(404)

    if not _public_is_authed(token):
        return redirect(url_for("public_trainee_login", token=token))

    training_type = (_session_get(s, "training_type", "") or "").strip().upper()
    if training_type != "DIRIGEANT VAE":
        abort(404)

    photo_url = ""
    photo_token = (t.get("identity_photo") or "").strip()
    if photo_token:
        photo_url = url_for("public_download_file", token=token, file_token=photo_token)

    return render_template(
        "public_candidate_sheet_form.html",
        candidate=_build_candidate_sheet_data(s, t),
        photo_url=photo_url,
        save_url=url_for("public_trainee_candidate_sheet_save", token=token),
    )


@app.get("/admin/sessions/<session_id>/stagiaires/<trainee_id>/fiche-candidat-completee")
@admin_login_required
def admin_trainee_candidate_sheet(session_id: str, trainee_id: str):
    data = load_data()
    s = find_session(data, session_id)
    if not s:
        abort(404)

    trainees = _session_trainees_list(s)
    t = next((x for x in trainees if x.get("id") == trainee_id), None)
    if not t:
        abort(404)

    training_type = (_session_get(s, "training_type", "") or "").strip().upper()
    if training_type != "DIRIGEANT VAE":
        abort(404)

    photo_url = ""
    photo_token = (t.get("identity_photo") or "").strip()
    if photo_token:
        photo_url = url_for("admin_view_upload", path=photo_token)

    return render_template(
        "admin_trainee_candidate_sheet.html",
        candidate=_build_candidate_sheet_data(s, t),
        photo_url=photo_url,
    )

@app.get("/api/docs_to_control")
@admin_login_required
def api_docs_to_control():
    data = load_data()
    out = []

    for s in data.get("sessions", []):
        session_id = s.get("id")
        session_name = _session_get(s, "name", "")
        training_type = _session_get(s, "training_type", "")

        trainees = _session_trainees_list(s)

        for t in trainees:
            # s'assure que les docs requis existent (sinon liste vide => pas détecté)
            ensure_documents_schema_for_trainee(t, training_type)

            docs = t.get("documents") or []
            pending = 0
            for d in docs:
                st = (d.get("status") or "").strip().upper()
                if st in ("A CONTRÔLER", "A CONTROLER"):
                    pending += 1

            if pending > 0:
                out.append({
                    "session_id": session_id,
                    "session_name": session_name,
                    "training_type": training_type,
                    "trainee_id": t.get("id"),
                    "last_name": t.get("last_name", ""),
                    "first_name": t.get("first_name", ""),
                    "pending_count": pending,
                    "admin_url": f"/admin/sessions/{session_id}/stagiaires/{t.get('id')}",
                })

    # tri: plus urgent d'abord (plus de docs à contrôler)
    out.sort(key=lambda x: x.get("pending_count", 0), reverse=True)

    return jsonify({"ok": True, "items": out, "count": len(out)})


from flask import make_response

@app.get("/docs_to_control.json")
def public_docs_to_control():
    data = load_data()
    out = []

    for s in data.get("sessions", []):
        session_id = s.get("id")
        session_name = _session_get(s, "name", "")
        training_type = _session_get(s, "training_type", "")

        trainees = _session_trainees_list(s)

        for t in trainees:
            ensure_documents_schema_for_trainee(t, training_type)

            docs = t.get("documents") or []
            pending = 0
            for d in docs:
                st = (d.get("status") or "").strip().upper()
                if st in ("A CONTRÔLER", "A CONTROLER"):
                    pending += 1

            if pending > 0:
                out.append({
                    "session_id": session_id,
                    "session_name": session_name,
                    "training_type": training_type,
                    "trainee_id": t.get("id"),
                    "last_name": t.get("last_name", ""),
                    "first_name": t.get("first_name", ""),
                    "pending_count": pending,
                    "admin_url": f"/admin/sessions/{session_id}/stagiaires/{t.get('id')}",
                })

    out.sort(key=lambda x: x.get("pending_count", 0), reverse=True)

    resp = make_response(jsonify({"ok": True, "items": out, "count": len(out)}))

    # ✅ autorise le fetch depuis ton dashboard (autre domaine)
    resp.headers["Access-Control-Allow-Origin"] = "*"
    resp.headers["Access-Control-Allow-Methods"] = "GET, OPTIONS"
    resp.headers["Access-Control-Allow-Headers"] = "Content-Type"

    return resp


@app.get("/api/trainees_search")
@admin_login_required
def api_trainees_search():
    q = (request.args.get("q") or "").strip().lower()
    if not q or len(q) < 2:
        return jsonify({"ok": True, "items": []})

    data = load_data()
    out = []

    for s in data.get("sessions", []):
        session_id = s.get("id")
        session_name = _session_get(s, "name", "")
        training_type = _session_get(s, "training_type", "")

        trainees = _session_trainees_list(s)

        for t in trainees:
            fn = (t.get("first_name") or "").strip()
            ln = (t.get("last_name") or "").strip()
            full = f"{fn} {ln}".strip().lower()

            # match prénom/nom (contient)
            if q in full or q in fn.lower() or q in ln.lower():
                out.append({
                    "session_id": session_id,
                    "session_name": session_name,
                    "training_type": training_type,
                    "trainee_id": t.get("id"),
                    "first_name": fn,
                    "last_name": ln,
                    "convention_status": t.get("convention_status") or "soon",
                    "convention_saisie_done": bool(t.get("convention_saisie_done")),
                    "convention_signed_done": bool(t.get("convention_signed_done")),
                    "test_fr_status": t.get("test_fr_status") or "soon",
                    "admin_url": f"/admin/sessions/{session_id}/stagiaires/{t.get('id')}",
                })

    # tri: nom puis prénom
    out.sort(key=lambda x: ((x.get("last_name") or "").lower(), (x.get("first_name") or "").lower()))

    # limite pour éviter des réponses énormes
    out = out[:30]

    return jsonify({"ok": True, "items": out, "count": len(out)})


@app.get("/api/cnaps/trainees")
@admin_login_required
def api_cnaps_trainees():
    data = load_data()
    sessions_out = []
    for s in data.get("sessions", []):
        if bool(s.get("archived")):
            continue
        trainees = _session_trainees_list(s)
        sessions_out.append({
            "id": s.get("id"),
            "name": _session_get(s, "name", ""),
            "training_type": _session_get(s, "training_type", ""),
            "trainees": [
                {
                    "id": t.get("id"),
                    "first_name": t.get("first_name", ""),
                    "last_name": t.get("last_name", ""),
                    "birth_date": t.get("birth_date", ""),
                }
                for t in trainees
            ],
        })

    return jsonify({"ok": True, "sessions": sessions_out})


@app.post("/api/cnaps/pre_request")
@admin_login_required
@admin_write_required
def api_cnaps_pre_request():
    payload = request.get_json(silent=True) or {}
    items = payload.get("items") if isinstance(payload.get("items"), list) else []
    if not items:
        return jsonify({"ok": False, "error": "no_items"}), 400

    data = load_data()
    updated = 0

    for item in items:
        session_id = str(item.get("session_id") or "").strip()
        trainee_id = str(item.get("trainee_id") or "").strip()
        if not session_id or not trainee_id:
            continue
        s = find_session(data, session_id)
        if not s:
            continue
        trainees = _session_trainees_list(s)
        t = next((x for x in trainees if x.get("id") == trainee_id), None)
        if not t:
            continue
        record_cnaps_pre_request(t)
        updated += 1
        s["trainees"] = trainees
        s.pop("stagiaires", None)

    if updated:
        save_data(data)

    return jsonify({"ok": True, "updated": updated})


@app.get("/admin/sessions/archived")
@admin_login_required
def admin_sessions_archived():
    data = load_data()
    out_sessions = []

    for s in data.get("sessions", []):
        if not bool(s.get("archived")):
            continue

        st = compute_stats(s)
        trainees = _session_trainees_list(s)
        dossier_complete_total = sum(
            1 for t in trainees if dossier_is_complete_total(t, _session_get(s, "training_type", ""))
        )
        session_dossier_complete = (len(trainees) > 0 and dossier_complete_total == len(trainees))
        out_sessions.append({
            "id": s.get("id"),
            "name": _session_get(s, "name", ""),
            "training_type": _session_get(s, "training_type", ""),
            "date_start": _session_get(s, "date_start", ""),
            "date_end": _session_get(s, "date_end", ""),
            "exam_date": _session_get(s, "exam_date", ""),
            "exam_theory_date": _session_get(s, "exam_theory_date", ""),
            "exam_practice_date": _session_get(s, "exam_practice_date", ""),
            "total": st["total"],
            "session_is_conform": st["session_is_conform"],
            "session_dossier_complete": session_dossier_complete,
        })

    return render_template(
        "admin_sessions_archived.html",
        sessions=out_sessions,
        formation_types=FORMATION_TYPES,
    )

# =========================
# RELANCE TELEPHONIQUE
# =========================

def phone_missing_details_text(t: Dict[str, Any], training_type: str) -> str:
    # docs requis alignés
    ensure_documents_schema_for_trainee(t, training_type)

    # documents incomplets = tout ce qui n'est pas CONFORME (et le permis si A3P + no_permis)
    docs_lines = []
    tt = (training_type or "").strip().upper()
    no_permis = bool(t.get("no_permis"))

    for d in (t.get("documents") or []):
        key = (d.get("key") or "").strip()
        label = (d.get("label") or "Document").strip()
        st = (d.get("status") or "").strip().upper()

        # permis optionnel si no_permis
        if tt == "A3P" and key == "permis" and no_permis:
            continue

        if st != "CONFORME":
            if not st:
                st = "NON DÉPOSÉ"
            docs_lines.append(f"- {label} : {st}")

    docs_txt = "\n".join(docs_lines) if docs_lines else "- Aucun (selon statuts actuels)"

    infos_txt = infos_missing_text(t) or "- Aucune"

    return (
        "📄 Documents incomplets :\n"
        f"{docs_txt}\n\n"
        "🧾 Informations à compléter :\n"
        f"{infos_txt}\n"
    )


def _find_session_and_trainee(data: Dict[str, Any], session_id: str, trainee_id: str):
    s = find_session(data, session_id)
    if not s:
        return None, None
    trainees = _session_trainees_list(s)
    t = next((x for x in trainees if x.get("id") == trainee_id), None)
    return s, t


@app.post("/api/sessions/<session_id>/stagiaires/<trainee_id>/phone-relance/send")
@admin_login_required
@admin_write_required
def api_phone_relance_send(session_id: str, trainee_id: str):
    payload = request.get_json(silent=True) or {}
    admin_comment = (payload.get("comment") or "").strip()

    data = load_data()
    s, t = _find_session_and_trainee(data, session_id, trainee_id)
    if not s or not t:
        return jsonify({"ok": False, "error": "not_found"}), 404

    training_type = _session_get(s, "training_type", "")
    t.setdefault("phone_followups", [])

    # Détails incomplets
    missing_details = phone_missing_details_text(t, training_type)

    # Token unique pour les actions secrétaire
    followup_token = uuid.uuid4().hex
    followup_id = "PHN-" + followup_token[:10].upper()

    # Enregistre la demande
    entry = {
        "id": followup_id,
        "token": followup_token,
        "type": "DEMANDE RELANCE",
        "at": _now_iso(),
        "details": missing_details,
        "comment": admin_comment,
        "status": "PENDING",
    }
    t["phone_followups"].insert(0, entry)

    # Infos mail
    first_name = (t.get("first_name") or "").strip()
    last_name = (t.get("last_name") or "").strip()
    email = (t.get("email") or "").strip()
    phone = (t.get("phone") or "").strip()

    formation_type = formation_label(_session_get(s, "training_type", ""))
    dstart = fr_date(_session_get(s, "date_start", ""))
    dend = fr_date(_session_get(s, "date_end", ""))

    # Liens actions secrétaire (page qui ouvre une modale)
    base = PUBLIC_BASE_URL.rstrip("/")
    action_url = f"{base}/phone-followup/{followup_token}"

    url_called = action_url + "?action=called"
    url_noanswer = action_url + "?action=no_answer"

    subject = f"📞 Relance téléphonique – Dossier incomplet – {first_name} {last_name}".strip()

    html = mail_layout(f"""
      <h2 style="text-align:center">📞 Relance téléphonique – Dossier incomplet</h2>

      <div style="background:#f3f4f6;border:1px solid #e5e7eb;border-radius:12px;padding:14px;margin:14px 0">
        <p style="margin:0 0 8px 0"><strong>Stagiaire :</strong> {first_name} {last_name}</p>
        <p style="margin:0 0 8px 0"><strong>Formation :</strong> {formation_type}</p>
        <p style="margin:0 0 8px 0"><strong>Dates :</strong> {dstart} → {dend}</p>
        <p style="margin:0 0 8px 0"><strong>Téléphone :</strong> {phone or "—"}</p>
        <p style="margin:0"><strong>Email :</strong> {email or "—"}</p>
      </div>

      <p style="margin:12px 0 8px 0"><strong>Éléments incomplets :</strong></p>
      <pre style="white-space:pre-wrap;background:#fff;border:1px solid #e5e7eb;padding:12px;border-radius:12px;margin:0">{missing_details}</pre>

      {"<p style='margin-top:12px'><strong>Commentaire admin :</strong><br>" + admin_comment + "</p>" if admin_comment else ""}

      <div style="text-align:center;margin-top:18px;display:flex;gap:10px;justify-content:center;flex-wrap:wrap;">
        <a href="{url_called}"
           style="display:inline-block;background:#16a34a;color:white;padding:12px 16px;border-radius:10px;text-decoration:none;font-weight:800">
          ✅ J’ai appelé la personne
        </a>

        <a href="{url_noanswer}"
           style="display:inline-block;background:#dc2626;color:white;padding:12px 16px;border-radius:10px;text-decoration:none;font-weight:800">
          ❌ Je n’ai pas pu joindre la personne
        </a>
      </div>

      <p class="hint" style="margin-top:14px;color:#6b7280;font-size:13px;text-align:center">
        Ces boutons ouvrent une page avec une modale pour saisir le commentaire.
      </p>
    """)

    # envoi au mail secrétaire
    ok = brevo_send_email(
        "znaw83@gmail.com",
        subject,
        html,
        cc_emails=["clement@integraleacademy.com"],
    )

    add_notification(
        data,
        "notifications_phone_relances",
        f"{first_name} {last_name} • {formation_type}",
        meta={
            "first_name": first_name,
            "last_name": last_name,
            "training": formation_type,
            "phone": phone,
            "email": email,
            "session_id": s.get("id"),
            "trainee_id": t.get("id"),
            "followup_id": followup_id,
            "missing_details": missing_details,
            "admin_comment": admin_comment,
            "call_status": "À appeler",
            "no_answer_count": 0,
        },
    )

    # persistance
    s["trainees"] = _session_trainees_list(s)
    s.pop("stagiaires", None)
    save_data(data)

    return jsonify({"ok": True, "email_ok": bool(ok), "followup_id": followup_id})


# =========================
# PRELEVEMENT REJETE
# =========================

def _append_admin_comment_flag(current: str, flag_text: str) -> str:
    current = (current or "").strip()
    if not current:
        return flag_text
    if flag_text in current:
        return current
    return current + "\n" + flag_text

def _remove_admin_comment_flag(current: str, flag_text: str) -> str:
    current = (current or "").strip()
    if not current:
        return ""
    kept = [line for line in current.splitlines() if line.strip() != flag_text]
    return "\n".join(kept).strip()

def _send_prelevement_new_date_email(
    trainee: dict,
    session: dict,
    rejected_request: dict,
    new_date: str,
    comment: str = "",
) -> None:
    first_name = (trainee.get("first_name") or "").strip()
    last_name = (trainee.get("last_name") or "").strip()
    formation_type = formation_label(_session_get(session, "training_type", ""))

    amount = rejected_request.get("amount", "")
    scheduled_date = fr_date(rejected_request.get("scheduled_date", ""))
    new_date_fr = fr_date(new_date) or new_date

    subject = f"📩 Nouveau prélèvement proposé – {first_name} {last_name}".strip()
    html = mail_layout(f"""
      <h2 style="text-align:center">📩 Nouveau prélèvement proposé</h2>

      <div style="background:#f3f4f6;border:1px solid #e5e7eb;border-radius:12px;padding:14px;margin:14px 0">
        <p style="margin:0 0 8px 0"><strong>Stagiaire :</strong> {first_name} {last_name}</p>
        <p style="margin:0 0 8px 0"><strong>Formation :</strong> {formation_type}</p>
        <p style="margin:0 0 8px 0"><strong>Montant :</strong> {amount}</p>
        <p style="margin:0 0 8px 0"><strong>Date initiale :</strong> {scheduled_date}</p>
        <p style="margin:0"><strong>Nouvelle date proposée :</strong> {new_date_fr}</p>
      </div>

      {"<p><strong>Commentaire :</strong><br>" + comment + "</p>" if comment else ""}
    """)

    brevo_send_email("clement@integraleacademy.com", subject, html)


@app.post("/api/sessions/<session_id>/stagiaires/<trainee_id>/financement-rejet/send")
@admin_login_required
@admin_write_required
def api_financement_rejet_send(session_id: str, trainee_id: str):
    payload = request.get_json(silent=True) or {}
    amount = (payload.get("amount") or "").strip()
    scheduled_date = (payload.get("scheduled_date") or "").strip()

    if not amount or not scheduled_date:
        return jsonify({"ok": False, "error": "missing_fields"}), 400

    data = load_data()
    s, t = _find_session_and_trainee(data, session_id, trainee_id)
    if not s or not t:
        return jsonify({"ok": False, "error": "not_found"}), 404

    token = uuid.uuid4().hex
    secretariat_token = uuid.uuid4().hex
    entry_id = "PAY-" + token[:10].upper()

    t.setdefault("financement_rejected_requests", [])
    t["financement_rejected_requests"].insert(0, {
        "id": entry_id,
        "token": token,
        "secretariat_token": secretariat_token,
        "amount": amount,
        "scheduled_date": scheduled_date,
        "at": _now_iso(),
        "status": "PENDING",
    })

    t["financement_status"] = "in_review"
    t["financement_rejected_note"] = "⚠️ Prélèvement rejeté"
    t["comment"] = _append_admin_comment_flag(t.get("comment", ""), "⚠️ Prélèvement rejeté")

    first_name = (t.get("first_name") or "").strip()
    last_name = (t.get("last_name") or "").strip()
    email = (t.get("email") or "").strip()
    phone = (t.get("phone") or "").strip()
    formation_type = formation_label(_session_get(s, "training_type", ""))
    training_name = formation_label(_session_get(s, "training_type", "") or s.get("name") or "formation")
    scheduled_fr = fr_date(scheduled_date) or scheduled_date

    base = PUBLIC_BASE_URL.rstrip("/")
    reply_url = f"{base}/prelevement-rejete/{token}"
    secretariat_url = f"{base}/prelevement-rejete-secretaire/{secretariat_token}"

    subject = "⚠️ Prélèvement rejeté – action requise"
    html = mail_layout(f"""
      <p>Bonjour {first_name} {last_name},</p>

      <p>Je me permets de revenir vers vous concernant votre formation
      <strong>{formation_type}</strong>.</p>

      <p>Nous avons pu constater que votre prélèvement d'un montant de
      <strong>{amount}</strong> euros initialement prévu le <strong>{scheduled_fr}</strong> a été rejeté.</p>

      <p>Pourriez-vous svp nous indiquer à quelle date nous pouvons prévoir un nouveau prélèvement
      en cliquant ici ?</p>

      <p style="text-align:center;margin:18px 0">
        <a href="{reply_url}"
           style="display:inline-block;background:#2563eb;color:#fff;padding:12px 16px;border-radius:10px;text-decoration:none;font-weight:800">
          Indiquer une nouvelle date
        </a>
      </p>

      <p>En cas de difficulté vous pouvez nous contacter au 04 22 47 07 68.</p>

      <p>Je vous remercie par avance,</p>

      <p>Clément VAILLANT<br>Directeur Intégrale Academy</p>
    """)

    email_ok = brevo_send_email(email, subject, html) if email else False

    dstart = fr_date(_session_get(s, "date_start", ""))
    dend = fr_date(_session_get(s, "date_end", ""))
    exam_date = fr_date(_session_get(s, "exam_date", ""))
    formation_dates = "Du <strong>{}</strong> au <strong>{}</strong>".format(dstart, dend) if dstart and dend else "Dates à confirmer"

    secretariat_subject = f"⚠️ Prélèvement rejeté – rappel à prévoir ({first_name} {last_name})".strip()
    secretariat_html = mail_layout(f"""
      <h2 style="text-align:center">⚠️ Prélèvement rejeté</h2>
      <p>Merci de rappeler le stagiaire pour convenir d’une nouvelle date de prélèvement.</p>

      <div style="background:#f3f4f6;border:1px solid #e5e7eb;border-radius:12px;padding:14px;margin:14px 0">
        <p style="margin:0 0 8px 0"><strong>Stagiaire :</strong> {first_name} {last_name}</p>
        <p style="margin:0 0 8px 0"><strong>Téléphone :</strong> {phone or "—"}</p>
        <p style="margin:0 0 8px 0"><strong>Email :</strong> {email or "—"}</p>
        <p style="margin:0 0 8px 0"><strong>Formation :</strong> {training_name}</p>
        <p style="margin:0 0 8px 0"><strong>Dates de formation :</strong> {formation_dates}</p>
        <p style="margin:0 0 8px 0"><strong>Date d’examen :</strong> {exam_date or "—"}</p>
        <p style="margin:0 0 8px 0"><strong>Montant :</strong> {amount}</p>
        <p style="margin:0"><strong>Date initiale :</strong> {scheduled_fr}</p>
      </div>

      <p style="text-align:center;margin:18px 0">
        <a href="{secretariat_url}"
           style="display:inline-block;background:#2563eb;color:#fff;padding:12px 16px;border-radius:10px;text-decoration:none;font-weight:800">
          Proposer une nouvelle date
        </a>
      </p>
    """)
    brevo_send_email("znaw83@gmail.com", secretariat_subject, secretariat_html)

    sms_name = first_name.strip()
    sms_prefix = f"Bonjour {sms_name}, " if sms_name else "Bonjour, "
    sms = (
        f"{sms_prefix}Je reviens vers vous concernant votre formation {training_name}. "
        f"Votre prélèvement d'un montant de {amount} euros prévu le {scheduled_fr} a été rejeté. "
        "Nous vous remercions de bien vouloir nous indiquer une nouvelle date de prélèvement "
        f"en cliquant ici : {reply_url} "
        "En cas de difficultés, vous pouvez nous contacter au 04 22 47 07 68. "
        "Je vous remercie par avance, "
        "Clément VAILLANT - Intégrale Academy"
    ).strip()
    sms_ok = brevo_send_sms(phone, sms) if phone else False

    add_notification(
        data,
        "notifications_prelevements",
        f"{first_name} {last_name} • {training_name} • {amount} • {scheduled_fr}",
        meta={
            "first_name": first_name,
            "last_name": last_name,
            "training": training_name,
            "amount": amount,
            "scheduled_date": scheduled_date,
            "session_id": s.get("id"),
            "trainee_id": t.get("id"),
            "entry_id": entry_id,
            "secretariat_token": secretariat_token,
        },
    )

    s["trainees"] = _session_trainees_list(s)
    s.pop("stagiaires", None)
    save_data(data)

    return jsonify({
        "ok": True,
        "email_ok": bool(email_ok),
        "sms_ok": bool(sms_ok),
        "reply_url": reply_url,
        "note": t.get("financement_rejected_note"),
        "comment": t.get("comment", ""),
    })


@app.get("/prelevement-rejete/<token>")
def prelevement_rejete_page(token: str):
    data = load_data()
    found = None
    found_trainee = None
    found_session = None

    for s in data.get("sessions", []) or []:
        for t in (s.get("trainees") or []):
            for it in (t.get("financement_rejected_requests") or []):
                if (it.get("token") or "").strip() == token:
                    found = it
                    found_trainee = t
                    found_session = s
                    break
            if found:
                break
        if found:
            break

    if not found:
        return "<h3>Lien invalide ou expiré.</h3>", 404

    new_date = found.get("new_date")

    return render_template(
        "prelevement_rejete.html",
        token=token,
        trainee=found_trainee,
        session=found_session,
        formation_label=formation_label(_session_get(found_session, "training_type", "")),
        amount=found.get("amount", ""),
        scheduled_date=fr_date(found.get("scheduled_date", "")),
        ref_id=found.get("id", ""),
        new_date=new_date,
    )


@app.get("/prelevement-rejete-secretaire/<token>")
def prelevement_rejete_secretaire_page(token: str):
    data = load_data()
    found = None
    found_trainee = None
    found_session = None

    for s in data.get("sessions", []) or []:
        for t in (s.get("trainees") or []):
            for it in (t.get("financement_rejected_requests") or []):
                if (it.get("secretariat_token") or "").strip() == token:
                    found = it
                    found_trainee = t
                    found_session = s
                    break
            if found:
                break
        if found:
            break

    if not found:
        return "<h3>Lien invalide ou expiré.</h3>", 404

    new_date = found.get("new_date")
    new_date_fr = fr_date(new_date) if new_date else ""

    return render_template(
        "prelevement_rejete_secretaire.html",
        token=token,
        trainee=found_trainee,
        session=found_session,
        formation_label=formation_label(_session_get(found_session, "training_type", "")),
        amount=found.get("amount", ""),
        scheduled_date=fr_date(found.get("scheduled_date", "")),
        ref_id=found.get("id", ""),
        new_date=new_date_fr or new_date or "",
    )


@app.post("/prelevement-rejete-secretaire/<token>/reply")
def prelevement_rejete_secretaire_reply(token: str):
    new_date = (request.form.get("new_date") or "").strip()

    data = load_data()
    found = None
    found_trainee = None
    found_session = None

    for s in data.get("sessions", []) or []:
        for t in (s.get("trainees") or []):
            for it in (t.get("financement_rejected_requests") or []):
                if (it.get("secretariat_token") or "").strip() == token:
                    found = it
                    found_trainee = t
                    found_session = s
                    break
            if found:
                break
        if found:
            break

    if not found:
        return "<h3>Lien invalide ou expiré.</h3>", 404

    if found.get("new_date"):
        return render_template(
            "prelevement_rejete_secretaire.html",
            token=token,
            trainee=found_trainee,
            session=found_session,
            formation_label=formation_label(_session_get(found_session, "training_type", "")),
            amount=found.get("amount", ""),
            scheduled_date=fr_date(found.get("scheduled_date", "")),
            ref_id=found.get("id", ""),
            new_date=fr_date(found.get("new_date")) or found.get("new_date"),
        )

    if not new_date:
        return "<h3>Veuillez indiquer une date.</h3>", 400

    found["status"] = "DONE"
    found["responded_at"] = _now_iso()
    found["new_date"] = new_date
    found["new_date_source"] = "SECRETARIAT"

    _send_prelevement_new_date_email(found_trainee, found_session, found, new_date)

    trainee_display_name = _format_trainee_name(found_trainee.get("first_name", ""), found_trainee.get("last_name", ""))
    add_admin_notification(
        data,
        f"🟢{trainee_display_name} - Nouveau prélèvement proposé le {fr_date(new_date) or new_date}",
        meta={
            "type": "prelevement_new_date",
            "source": "secretariat_public_page",
            "session_id": found_session.get("id"),
            "trainee_id": found_trainee.get("id"),
            "entry_id": found.get("id"),
            "comment": (found.get("comment") or "").strip(),
        },
    )

    found_session["trainees"] = _session_trainees_list(found_session)
    found_session.pop("stagiaires", None)
    save_data(data)

    return """
    <div style="font-family:Arial,sans-serif;max-width:520px;margin:60px auto;padding:18px;border:1px solid #e5e7eb;border-radius:14px">
      <h2 style="margin:0 0 10px 0">✅ Merci !</h2>
      <p style="margin:0;color:#374151">La nouvelle date a bien été enregistrée.</p>
    </div>
    """


@app.post("/prelevement-rejete/<token>/reply")
def prelevement_rejete_reply(token: str):
    new_date = (request.form.get("new_date") or "").strip()
    comment = (request.form.get("comment") or "").strip()

    data = load_data()
    found = None
    found_trainee = None
    found_session = None

    for s in data.get("sessions", []) or []:
        for t in (s.get("trainees") or []):
            for it in (t.get("financement_rejected_requests") or []):
                if (it.get("token") or "").strip() == token:
                    found = it
                    found_trainee = t
                    found_session = s
                    break
            if found:
                break
        if found:
            break

    if not found:
        return "<h3>Lien invalide ou expiré.</h3>", 404

    if found.get("new_date"):
        return render_template(
            "prelevement_rejete.html",
            token=token,
            trainee=found_trainee,
            session=found_session,
            formation_label=formation_label(_session_get(found_session, "training_type", "")),
            amount=found.get("amount", ""),
            scheduled_date=fr_date(found.get("scheduled_date", "")),
            ref_id=found.get("id", ""),
            new_date=found.get("new_date"),
        )

    found["status"] = "DONE"
    found["responded_at"] = _now_iso()
    found["new_date"] = new_date
    found["comment"] = comment
    found["new_date_source"] = "TRAINEE"

    _send_prelevement_new_date_email(found_trainee, found_session, found, new_date, comment)

    trainee_display_name = _format_trainee_name(found_trainee.get("first_name", ""), found_trainee.get("last_name", ""))
    add_admin_notification(
        data,
        f"🟢{trainee_display_name} - Nouveau prélèvement proposé le {fr_date(new_date) or new_date}",
        meta={
            "type": "prelevement_new_date",
            "source": "trainee_public_page",
            "session_id": found_session.get("id"),
            "trainee_id": found_trainee.get("id"),
            "entry_id": found.get("id"),
            "comment": comment,
        },
    )

    found_session["trainees"] = _session_trainees_list(found_session)
    found_session.pop("stagiaires", None)
    save_data(data)

    return """
    <div style="font-family:Arial,sans-serif;max-width:520px;margin:60px auto;padding:18px;border:1px solid #e5e7eb;border-radius:14px">
      <h2 style="margin:0 0 10px 0">✅ Merci !</h2>
      <p style="margin:0;color:#374151">Votre réponse a bien été transmise. Nous revenons vers vous rapidement.</p>
    </div>
    """


@app.get("/phone-followup/<token>")
def phone_followup_page(token: str):
    # page publique "action secrétaire" (sans login), basée sur un token unique
    action = (request.args.get("action") or "").strip()  # called / no_answer

    data = load_data()
    found = None
    found_session_id = None
    found_trainee_id = None

    for s in data.get("sessions", []) or []:
        for t in (s.get("trainees") or []):
            for it in (t.get("phone_followups") or []):
                if (it.get("token") or "").strip() == token:
                    found = it
                    found_session_id = s.get("id")
                    found_trainee_id = t.get("id")
                    break
            if found:
                break
        if found:
            break

    if not found:
        return "<h3>Lien invalide ou expiré.</h3>", 404

    # petite page qui ouvre automatiquement une modale (comme demandé)
    # et envoie le commentaire via POST
    return render_template(
    "phone_followup.html",
    token=token,
    action=action,
    ref_id=found.get("id", ""),
)


@app.post("/phone-followup/<token>/reply")
def phone_followup_reply(token: str):
    outcome = (request.form.get("outcome") or "").strip().upper()
    comment = (request.form.get("comment") or "").strip()

    if outcome not in ("CALLED", "NO_ANSWER"):
        return "<h3>Action invalide.</h3>", 400

    data = load_data()

    s_found = None
    t_found = None
    entry_found = None

    for s in data.get("sessions", []) or []:
        for t in (s.get("trainees") or []):
            for it in (t.get("phone_followups") or []):
                if (it.get("token") or "").strip() == token:
                    s_found = s
                    t_found = t
                    entry_found = it
                    break
            if entry_found:
                break
        if entry_found:
            break

    if not entry_found:
        return "<h3>Lien invalide ou expiré.</h3>", 404

    # on enregistre la réponse comme un nouvel événement (historique)
    t_found.setdefault("phone_followups", [])
    t_found["phone_followups"].insert(0, {
        "id": "PHN-REP-" + uuid.uuid4().hex[:8].upper(),
        "type": "RÉPONSE SECRÉTAIRE",
        "at": _now_iso(),
        "details": ("✅ Appelé" if outcome == "CALLED" else "❌ Pas pu joindre"),
        "comment": comment,
        "ref": entry_found.get("id", ""),
    })

    # marque la demande comme traitée (optionnel)
    entry_found["status"] = "DONE"
    entry_found["done_at"] = _now_iso()
    entry_found["done_outcome"] = outcome

    trainee_display_name = _format_trainee_name(t_found.get("first_name", ""), t_found.get("last_name", ""))
    if outcome == "CALLED":
        add_admin_notification(
            data,
            f"🟢Relance téléphonique {trainee_display_name} a été appelé",
            meta={
                "type": "relance_call_result",
                "source": "phone_followup_public_page",
                "outcome": outcome,
                "session_id": s_found.get("id"),
                "trainee_id": t_found.get("id"),
                "comment": comment,
                "call_status": "Personne jointe",
            },
        )
    else:
        current_no_answer = _parse_no_answer_count(entry_found.get("no_answer_count"))
        no_answer_count = min(3, current_no_answer + 1)
        entry_found["no_answer_count"] = no_answer_count
        display = {
            1: "1er appel pas de réponse",
            2: "2ème appel pas de réponse",
            3: "3ème appel pas de réponse",
        }[no_answer_count]
        icon = {1: '🟡', 2: '🟠', 3: '🔴'}[no_answer_count]
        add_admin_notification(
            data,
            f"{icon}Relance téléphonique {trainee_display_name} {display}",
            meta={
                "type": "relance_call_result",
                "source": "phone_followup_public_page",
                "outcome": outcome,
                "no_answer_count": no_answer_count,
                "session_id": s_found.get("id"),
                "trainee_id": t_found.get("id"),
                "comment": comment,
                "call_status": display,
            },
        )

    # persist
    s_found["trainees"] = _session_trainees_list(s_found)
    s_found.pop("stagiaires", None)
    save_data(data)

    return """
    <div style="font-family:Arial,sans-serif;max-width:520px;margin:60px auto;padding:18px;border:1px solid #e5e7eb;border-radius:14px">
      <h2 style="margin:0 0 10px 0">✅ Réponse enregistrée</h2>
      <p style="margin:0;color:#374151">Merci, le retour a bien été ajouté à l’historique.</p>
    </div>
    """

# =========================
# ALIAS rétrocompatibilité (stagiaires <-> trainees)
# =========================

@app.post("/api/sessions/<session_id>/stagiaires/create")
@admin_login_required
@admin_write_required
def api_create_trainee_alias(session_id: str):
    # redirige vers la vraie fonction
    return api_create_trainee(session_id)

@app.post("/api/sessions/<session_id>/stagiaires/<trainee_id>/delete")
@admin_login_required
@admin_write_required
def api_delete_trainee_alias(session_id: str, trainee_id: str):
    return api_delete_trainee(session_id, trainee_id)

@app.post("/api/sessions/<session_id>/trainees/<trainee_id>/update")
@admin_login_required
@admin_write_required
def api_update_trainee_alias(session_id: str, trainee_id: str):
    # ton update actuel est en /stagiaires/.../update
    return api_update_trainee(session_id, trainee_id)

import re
import unicodedata
import threading
from flask import request, jsonify

def _norm_name(s: str) -> str:
    s = (s or "").strip().lower()
    s = unicodedata.normalize("NFD", s)
    s = "".join(ch for ch in s if unicodedata.category(ch) != "Mn")  # enlève accents
    s = re.sub(r"[^a-z0-9]+", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s

def _match_trainee_from_filename(trainees: list, filename: str):
    """
    Match si NOM et PRÉNOM apparaissent dans le nom de fichier (normalisé).
    Renvoie (trainee, None) si OK, sinon (None, reason).
    """
    fn = _norm_name(filename)

    hits = []
    for t in trainees:
        ln = _norm_name(t.get("last_name", ""))
        fnm = _norm_name(t.get("first_name", ""))
        if not ln or not fnm:
            continue
        if ln in fn and fnm in fn:
            hits.append(t)

    if len(hits) == 1:
        return hits[0], None
    if len(hits) == 0:
        return None, "nom/prénom non trouvés dans le fichier"
    return None, "plusieurs stagiaires correspondent (homonyme)"

@app.post("/api/sessions/<session_id>/sst/bulk_upload")
@admin_login_required
@admin_write_required
def api_sst_bulk_upload(session_id: str):
    import traceback

    data = load_data()
    s = find_session(data, session_id)
    if not s:
        return jsonify({"ok": False, "error": "session_not_found"}), 404

    files = request.files.getlist("files")
    if not files:
        return jsonify({"ok": False, "error": "no_files"}), 400

    trainees = _session_trainees_list(s)

    received = 0
    added = []
    failed = []

    for f in files:
        if not f or not f.filename:
            continue

        received += 1
        original_name = f.filename
        ext = _safe_ext(original_name)

        if ext not in (".pdf", ".jpg", ".jpeg", ".png"):
            failed.append({"filename": original_name, "reason": "extension non autorisée"})
            continue

        trainee, reason = _match_trainee_from_filename(trainees, original_name)
        if not trainee:
            failed.append({"filename": original_name, "reason": reason or "non rattaché"})
            continue

        # ✅ sécurise l'id (selon ton schéma)
        trainee_id = trainee.get("id") or trainee.get("trainee_id") or trainee.get("personal_id")
        if not trainee_id:
            failed.append({"filename": original_name, "reason": "trainee_id introuvable (id manquant dans data.json)"})
            continue

        # ✅ si déjà un SST, on n'écrase pas + on n'envoie pas
        existing = ((trainee.get("deliverables") or {}).get("carte_sst") or "").strip()
        if existing:
            failed.append({"filename": original_name, "reason": "déjà un SST existant (non remplacé)"})
            continue
 

        try:
            # ✅ sécurité : remet le curseur au début (selon navigateur / proxy ça évite des fichiers vides)
            try:
                f.stream.seek(0)
            except Exception:
                pass

            stored = _store_file(session_id, trainee_id, "deliverables", f)
            token = _tokenize_path(stored)

        except Exception as e:
            # ✅ on log l'erreur complète dans Render
            print("=== BULK SST: erreur stockage ===")
            print("session_id:", session_id)
            print("trainee_id:", trainee_id)
            print("filename:", original_name)
            traceback.print_exc()

            # ✅ et on renvoie un message utile côté UI
            failed.append({"filename": original_name, "reason": f"erreur stockage: {str(e)}"})
            continue

        trainee.setdefault("deliverables", {})
        trainee["deliverables"]["carte_sst"] = token
        trainee["updated_at"] = _now_iso()

          # ✅ Envoi mail + SMS (comme l'import manuel deliverables)
        try:
            link = f"{PUBLIC_STUDENT_PORTAL_BASE.rstrip('/')}/espace/{trainee.get('public_token','')}"
            label = DELIVERABLE_LABELS["carte_sst"]
        
            first_name = (trainee.get("first_name") or "").strip() or "Madame, Monsieur"
            formation_type = formation_label(_session_get(s, "training_type", ""))
            dstart = fr_date(_session_get(s, "date_start", ""))
            dend = fr_date(_session_get(s, "date_end", ""))
        
            extra_line = (
                "🩺 Votre carte SST est disponible sur votre espace en ligne. "
                "Nous vous remettrons également un exemplaire papier en main propre "
                "(attention : aucun duplicata ne sera délivré). "
                "Conservez-la précieusement, elle peut être demandée par un employeur."
            )
        
            subject = f"{label} disponible – Intégrale Academy"  # ✅ FIX ICI
        
            html = mail_layout(f"""
              <h2 style="text-align:center">✅ {label} disponible</h2>
        
              <p>Bonjour <strong>{first_name}</strong>,</p>
        
              <p>
                Nous avons le plaisir de vous informer que votre <strong>{label}</strong>
                est désormais disponible dans votre espace stagiaire.
              </p>
        
              <p style='margin-top:10px;font-weight:700'>{extra_line}</p>
        
              <div style="background:#f3f4f6;border:1px solid #e5e7eb;border-radius:12px;padding:14px;margin:16px 0">
                <p style="margin:0 0 10px 0">
                  <strong>📌 Formation :</strong> {formation_type}
                  {" — <strong>Dates :</strong> " + dstart + " au " + dend if (dstart or dend) else ""}
                </p>
        
                <p style="margin:0">
                  <strong>📍 Accéder à votre espace stagiaire :</strong><br>
                  <a href="{link}" style="color:#1f8f4a;text-decoration:none;font-weight:bold">{link}</a>
                </p>
              </div>
        
              <p style="text-align:center;margin-top:18px">
                <a href="{link}"
                   style="display:inline-block;background:#1f8f4a;color:white;padding:12px 18px;border-radius:10px;
                          text-decoration:none;font-weight:bold">
                  👉 Accéder à mon espace stagiaire
                </a>
              </p>
        
              <p style="margin-top:22px">
                Pour toute question, vous pouvez nous contacter au <strong>04 22 47 07 68</strong>.
              </p>
        
              <p style="margin-top:22px">
                Bien cordialement,<br>
                <strong>Clément VAILLANT</strong><br>
                Directeur Intégrale Academy
              </p>
            """)
        
            sms_name = (trainee.get("first_name") or "").strip()
            sms = (
                f"Intégrale Academy ✅ {sms_name + ', ' if sms_name else ''}"
                f"Votre {label} est disponible sur votre Espace Stagiaire : {link} "
                f"A bientôt, la Team Intégrale Academy"
            )
        
            if (trainee.get("email") or "").strip():
                brevo_send_email(trainee.get("email",""), subject, html)
            if (trainee.get("phone") or "").strip():
                brevo_send_sms(trainee.get("phone",""), sms)
        
        except Exception as e:
            print("=== BULK SST: erreur envoi mail/sms ===", repr(e))



        added.append({
            "filename": original_name,
            "trainee_id": trainee_id,
            "trainee_name": f"{trainee.get('first_name','')} {trainee.get('last_name','')}".strip()
        })

    # persist
    s["trainees"] = trainees
    s.pop("stagiaires", None)
    save_data(data)

    return jsonify({
        "ok": True,
        "received": received,
        "added_count": len(added),
        "added": added,
        "failed": failed
    })


@app.post("/api/sessions/<session_id>/diplome/bulk_upload")
@admin_login_required
@admin_write_required
def api_diplome_bulk_upload(session_id: str):
    import traceback

    data = load_data()
    s = find_session(data, session_id)
    if not s:
        return jsonify({"ok": False, "error": "session_not_found"}), 404

    files = request.files.getlist("files")
    if not files:
        return jsonify({"ok": False, "error": "no_files"}), 400

    trainees = _session_trainees_list(s)

    received = 0
    added = []
    failed = []

    for f in files:
        if not f or not f.filename:
            continue

        received += 1
        original_name = f.filename
        ext = _safe_ext(original_name)

        if ext not in (".pdf", ".jpg", ".jpeg", ".png"):
            failed.append({"filename": original_name, "reason": "extension non autorisée"})
            continue

        trainee, reason = _match_trainee_from_filename(trainees, original_name)
        if not trainee:
            failed.append({"filename": original_name, "reason": reason or "non rattaché"})
            continue

        trainee_id = trainee.get("id") or trainee.get("trainee_id") or trainee.get("personal_id")
        if not trainee_id:
            failed.append({"filename": original_name, "reason": "trainee_id introuvable"})
            continue

        # ✅ si déjà un diplôme, on n'écrase pas + on n'envoie pas
        existing = ((trainee.get("deliverables") or {}).get("diplome") or "").strip()
        if existing:
            failed.append({"filename": original_name, "reason": "déjà un diplôme existant (non remplacé)"})
            continue

        try:
            try:
                f.stream.seek(0)
            except Exception:
                pass

            stored = _store_file(session_id, trainee_id, "deliverables", f)
            token = _tokenize_path(stored)
        except Exception as e:
            print("=== BULK DIPLOME: erreur stockage ===")
            traceback.print_exc()
            failed.append({"filename": original_name, "reason": f"erreur stockage: {str(e)}"})
            continue

        trainee.setdefault("deliverables", {})
        trainee["deliverables"]["diplome"] = token
        trainee["updated_at"] = _now_iso()

        # ✅ Envoi mail + SMS (mail enrichi + CNAPS + avis Google)
        try:
            link = f"{PUBLIC_STUDENT_PORTAL_BASE.rstrip('/')}/espace/{trainee.get('public_token','')}"
            label = DELIVERABLE_LABELS["diplome"]

            first_name = (trainee.get("first_name") or "").strip() or "Madame, Monsieur"
            subject = f"{label} disponible – Intégrale Academy"

            # --- Détection type formation (APS / A3P / Dirigeant) ---
            formation_type = formation_label(_session_get(s, "training_type", ""))
            ft = (formation_type or "").lower()

            cnaps_title = ""
            cnaps_html = ""

            if "aps" in ft:
                cnaps_title = "📌 Carte professionnelle : aucune démarche"
                cnaps_html = """
                  <p style="margin:0">
                    <strong>Vous n'avez aucune démarche à effectuer pour votre carte professionnelle.</strong>
                    Votre diplôme a été automatiquement transmis au CNAPS qui procède actuellement à une enquête administrative.
                    Dès que l'enquête sera terminée, vous recevrez votre carte professionnelle directement chez vous par courrier postal.
                  </p>
                  <p style="margin:10px 0 0 0">
                    <strong>Pour rappel :</strong> vous ne pouvez pas exercer la profession tant que vous n'avez pas reçu votre carte professionnelle.
                  </p>
                """
            elif "a3p" in ft:
                cnaps_title = "📌 Carte professionnelle : démarche à effectuer sur Téléservices CNAPS"
                cnaps_html = """
                  <p style="margin:0">
                    Vous pouvez à présent procéder à la demande de carte professionnelle depuis l'espace Téléservices du CNAPS.
                  </p>
                  <ul style="margin:10px 0 0 18px; padding:0; line-height:1.6">
                    <li>Si vous êtes déjà agent de sécurité : cliquez sur <strong>"Ma demande concerne une extension de carte professionnelle"</strong>.</li>
                    <li>Si vous n'êtes pas agent de sécurité : cliquez sur <strong>"Ma demande concerne une carte professionnelle"</strong>.</li>
                    <li>Dans les deux cas, complétez la rubrique <strong>"J'ai un NUB"</strong> en indiquant :
                      <strong>votre NOM</strong> (uniquement votre nom, pas votre prénom) et votre <strong>NUB</strong>
                      (les <strong>7 derniers chiffres</strong> de votre numéro d'autorisation préalable ou de votre carte professionnelle).
                    </li>
                    <li>Suivez les étapes et téléchargez les pièces justificatives :
                      <strong>pièce d'identité</strong>, <strong>justificatif de domicile</strong> de moins de 3 mois, et <strong>votre diplôme</strong>.
                    </li>
                  </ul>
                  <p style="margin:12px 0 0 0">
                    <a href="https://depot-teleservices-cnaps.interieur.gouv.fr/"
                       style="color:#1f8f4a;text-decoration:none;font-weight:800">
                      👉 Cliquez ici pour demander votre carte professionnelle
                    </a>
                  </p>
                """
            elif "dirigeant" in ft:
                cnaps_title = "📌 Agrément dirigeant : démarche à effectuer"
                cnaps_html = """
                  <p style="margin:0">
                    Vous pouvez à présent procéder à votre demande d'agrément dirigeant directement depuis le site internet du CNAPS
                    en complétant le formulaire en cliquant ici :
                  </p>
                  <p style="margin:12px 0 0 0">
                    <a href="https://www.cnaps.interieur.gouv.fr/Demarches-en-ligne/Vous-etes-un-particulier/Diriger-une-entreprise-de-securite-privee-un-organisme-de-formation-un-service-interne-de-securite/Diriger-un-organisme-de-formation-une-entreprise-de-securite-privee-un-service-interne-de-securite"
                       style="color:#1f8f4a;text-decoration:none;font-weight:800">
                      👉 Faire ma demande d'agrément dirigeant
                    </a>
                  </p>
                """

            cnaps_block = ""
            if cnaps_html:
                cnaps_block = f"""
                  <div style="background:#fff7ed;border:1px solid #fed7aa;border-radius:14px;padding:14px;margin:16px 0">
                    <div style="font-weight:900;margin:0 0 8px 0">{cnaps_title}</div>
                    <div style="color:#111827;line-height:1.6">{cnaps_html}</div>
                  </div>
                """

            # --- Bloc Avis Google (convaincant, pro, court) ---
            avis_block = """
              <div style="background:#ecfdf5;border:1px solid #bbf7d0;border-radius:14px;padding:14px;margin:18px 0">
                <div style="font-weight:900;margin:0 0 8px 0">⭐ Un avis Google, ça nous aide énormément</div>
                <div style="color:#065f46;line-height:1.6">
                  Si la formation vous a été utile, pouvez-vous prendre <strong>1 minute</strong> pour laisser un avis ?
                  Ça aide les futurs stagiaires à choisir une école sérieuse, et ça nous permet d’améliorer encore notre accompagnement.
                </div>
                <div style="text-align:center;margin-top:12px">
                  <a href="https://g.page/r/CZ0Ug-feyXjHEAE"
                     style="display:inline-block;background:#1f8f4a;color:white;padding:10px 14px;border-radius:10px;
                            text-decoration:none;font-weight:900">
                    👉 Laisser un avis Google
                  </a>
                </div>
              </div>
            """

            html = mail_layout(f"""
              <h2 style="text-align:center">✅ {label} disponible</h2>

              <p>Bonjour <strong>{first_name}</strong>,</p>

              <p style="margin-top:10px;font-weight:800">
                🎉 Félicitations ! Votre diplôme est maintenant disponible dans votre espace stagiaire.
              </p>

              <div style="background:#f3f4f6;border:1px solid #e5e7eb;border-radius:12px;padding:14px;margin:16px 0">
                <p style="margin:0 0 10px 0">
                  <strong>📌 Formation :</strong> {formation_type}
                </p>

                <p style="margin:0">
                  <strong>📍 Accéder à votre espace stagiaire :</strong><br>
                  <a href="{link}" style="color:#1f8f4a;text-decoration:none;font-weight:bold">{link}</a>
                </p>
              </div>

              <p style="text-align:center;margin-top:18px">
                <a href="{link}"
                   style="display:inline-block;background:#1f8f4a;color:white;padding:12px 18px;border-radius:10px;
                          text-decoration:none;font-weight:bold">
                  👉 Accéder à mon espace stagiaire
                </a>
              </p>

              {cnaps_block}

              {avis_block}

              <p style="margin-top:18px">
                Pour toute question, vous pouvez nous contacter au <strong>04 22 47 07 68</strong>.
              </p>

              <p style="margin-top:18px">
                Bien cordialement,<br>
                <strong>Clément VAILLANT</strong><br>
                Directeur Intégrale Academy
              </p>

              <hr style="margin:26px 0;border:none;border-top:1px solid #e5e7eb">

              <p style="font-size:12px;color:#6b7280;text-align:center;line-height:1.6">
                © Intégrale Academy — Merci de votre confiance 💛<br>
                54 chemin du Carreou 83480 PUGET SUR ARGENS / 142 rue de Rivoli 75001 PARIS<br>
                <a href="https://www.integraleacademy.com"
                   style="color:#1f8f4a;text-decoration:none;font-weight:bold">
                  integraleacademy.com
                </a>
              </p>
            """)

            sms_name = (trainee.get("first_name") or "").strip()
            sms = (
                f"Intégrale Academy ✅ {sms_name + ', ' if sms_name else ''}"
                f"votre {label} est disponible sur votre espace : {link} "
                f"(Aide : 04 22 47 07 68)"
            )

            if (trainee.get("email") or "").strip():
                brevo_send_email(trainee.get("email", ""), subject, html)
            if (trainee.get("phone") or "").strip():
                brevo_send_sms(trainee.get("phone", ""), sms)

        except Exception as e:
            print("=== BULK DIPLOME: erreur envoi mail/sms ===", repr(e))

        added.append({
            "filename": original_name,
            "trainee_id": trainee_id,
            "trainee_name": f"{trainee.get('first_name','')} {trainee.get('last_name','')}".strip()
        })

    s["trainees"] = trainees
    s.pop("stagiaires", None)
    save_data(data)

    return jsonify({"ok": True, "received": received, "added_count": len(added), "added": added, "failed": failed})

@app.post("/api/sessions/<session_id>/attestation/bulk_upload")
@admin_login_required
@admin_write_required
def api_attestation_bulk_upload(session_id: str):
    import traceback

    data = load_data()
    s = find_session(data, session_id)
    if not s:
        return jsonify({"ok": False, "error": "session_not_found"}), 404

    files = request.files.getlist("files")
    if not files:
        return jsonify({"ok": False, "error": "no_files"}), 400

    trainees = _session_trainees_list(s)

    received = 0
    added = []
    failed = []

    for f in files:
        if not f or not f.filename:
            continue

        received += 1
        original_name = f.filename
        ext = _safe_ext(original_name)

        if ext not in (".pdf", ".jpg", ".jpeg", ".png"):
            failed.append({"filename": original_name, "reason": "extension non autorisée"})
            continue

        trainee, reason = _match_trainee_from_filename(trainees, original_name)
        if not trainee:
            failed.append({"filename": original_name, "reason": reason or "non rattaché"})
            continue

        trainee_id = trainee.get("id") or trainee.get("trainee_id") or trainee.get("personal_id")
        if not trainee_id:
            failed.append({"filename": original_name, "reason": "trainee_id introuvable"})
            continue

        # ✅ si déjà une attestation, on n'écrase pas + on n'envoie pas
        existing = ((trainee.get("deliverables") or {}).get("attestation_fin_formation") or "").strip()
        if existing:
            failed.append({"filename": original_name, "reason": "déjà une attestation existante (non remplacée)"})
            continue

        try:
            try:
                f.stream.seek(0)
            except Exception:
                pass

            stored = _store_file(session_id, trainee_id, "deliverables", f)
            token = _tokenize_path(stored)
        except Exception as e:
            print("=== BULK ATTESTATION: erreur stockage ===")
            traceback.print_exc()
            failed.append({"filename": original_name, "reason": f"erreur stockage: {str(e)}"})
            continue

        trainee.setdefault("deliverables", {})
        trainee["deliverables"]["attestation_fin_formation"] = token
        trainee["updated_at"] = _now_iso()

        # ✅ mail + sms
        try:
            link = f"{PUBLIC_STUDENT_PORTAL_BASE.rstrip('/')}/espace/{trainee.get('public_token','')}"
            label = DELIVERABLE_LABELS["attestation_fin_formation"]

            first_name = (trainee.get("first_name") or "").strip() or "Madame, Monsieur"
            subject = f"{label} disponible – Intégrale Academy"

            html = mail_layout(f"""
              <h2 style="text-align:center">✅ {label} disponible</h2>
              <p>Bonjour <strong>{first_name}</strong>,</p>
              <p>📄 Votre attestation de fin de formation est disponible dans votre espace stagiaire.</p>
              <p style="text-align:center;margin-top:18px">
                <a href="{link}" style="display:inline-block;background:#1f8f4a;color:white;padding:12px 18px;border-radius:10px;text-decoration:none;font-weight:bold">
                  👉 Accéder à mon espace stagiaire
                </a>
              </p>
            """)

            sms_name = (trainee.get("first_name") or "").strip()
            sms = (
                f"Intégrale Academy ✅ {sms_name + ', ' if sms_name else ''}"
                f"Votre {label} est disponible sur votre Espace Stagiaire : {link} A bientôt, la Team Intégrale Academy"
            )

            if (trainee.get("email") or "").strip():
                brevo_send_email(trainee.get("email",""), subject, html)
            if (trainee.get("phone") or "").strip():
                brevo_send_sms(trainee.get("phone",""), sms)

        except Exception as e:
            print("=== BULK ATTESTATION: erreur envoi mail/sms ===", repr(e))

        added.append({
            "filename": original_name,
            "trainee_id": trainee_id,
            "trainee_name": f"{trainee.get('first_name','')} {trainee.get('last_name','')}".strip()
        })

    s["trainees"] = trainees
    s.pop("stagiaires", None)
    save_data(data)

    return jsonify({"ok": True, "received": received, "added_count": len(added), "added": added, "failed": failed})



# =========================
# VAE DESP - Dossier de faisabilité
# =========================
VAE_DATA_FILE = os.path.join(PERSIST_DIR, "data_vae.json")
_vae_lock = threading.RLock()

def _now_iso_utc() -> str:
    return datetime.datetime.utcnow().replace(microsecond=0).isoformat() + "Z"

def _vae_default_dossier(dossier_id: Optional[str] = None) -> Dict[str, Any]:
    now = _now_iso_utc()
    return {
        "id": dossier_id or str(uuid.uuid4()),
        "statut_dossier": "brouillon",
        "nature_demande": "initiale",
        "candidat": {
            "nom_naissance": "", "nom_usage": "", "prenoms": "", "date_naissance": "", "nationalite": "",
            "genre": "", "niveau_formation": "", "niveau_certification": "", "certifications_obtenues": "",
            "adresse": "", "code_postal": "", "ville": "", "telephone": "", "email": "", "statut": "", "convention_collective": "", "objectifs": []
        },
        "certification": {
            "intitule": "DIRIGEANT D’ENTREPRISE DE SÉCURITÉ PRIVÉE",
            "rncp": "40385",
            "certificateur": "SCOTIA FORMATION",
            "option": "NC",
            "parcours_mention": "NC",
            "prerequis": "",
            "vise": "complete",
            "blocs_vises": []
        },
        "experiences": [{"date_debut": "", "duree": "", "description": ""}],
        "blocs_competences": {
            "activite1": {"commentaires": "", "competence1": {"intitule": "", "statut": ""}, "competence2": {"intitule": "", "statut": ""}, "competence3": {"intitule": "", "statut": ""}, "competence4": {"intitule": "", "statut": ""}},
            "activite2": {"commentaires": "", "competence1": {"intitule": "", "statut": ""}, "competence2": {"intitule": "", "statut": ""}, "competence3": {"intitule": "", "statut": ""}, "competence4": {"intitule": "", "statut": ""}},
            "activite3": {"commentaires": "", "competence1": {"intitule": "", "statut": ""}, "competence2": {"intitule": "", "statut": ""}, "competence3": {"intitule": "", "statut": ""}, "competence4": {"intitule": "", "statut": ""}},
            "activite4": {"commentaires": "", "competence1": {"intitule": "", "statut": ""}, "competence2": {"intitule": "", "statut": ""}, "competence3": {"intitule": "", "statut": ""}, "competence4": {"intitule": "", "statut": ""}},
            "activite5": {"commentaires": "", "competence1": {"intitule": "", "statut": ""}, "competence2": {"intitule": "", "statut": ""}, "competence3": {"intitule": "", "statut": ""}, "competence4": {"intitule": "", "statut": ""}}
        },
        "parcours_previsionnel": {
            "accompagnement_individuel": {"heures": "", "modalites": ""},
            "accompagnement_collectif": {"heures": "", "modalites": ""},
            "formations_prealables": {"organisme": "", "intitule": "", "objectifs": "", "heures": ""},
            "immersion": {"type": "", "structure": "", "objectifs": "", "heures": ""},
            "autres_actions": ""
        },
        "avis_admin": {
            "decision": "", "motivation": "", "nom_accompagnateur": "", "email": "",
            "telephone": "", "organisme": "", "date": ""
        },
        "engagement": {
            "souhaite_accompagnement": False, "accord_analyse": False,
            "lieu_signature": "", "date_signature": "", "nom_signature": "", "signature_trace": "", "signature_signed_at": "", "commentaires_defavorable": ""
        },
        "created_at": now,
        "updated_at": now,
    }

def _vae_load_all() -> Dict[str, Any]:
    with _vae_lock:
        if not os.path.exists(VAE_DATA_FILE):
            data = {"dossiers": []}
            _vae_save_all(data)
            return data
        try:
            with open(VAE_DATA_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
            if not isinstance(data, dict):
                data = {"dossiers": []}
            if "dossiers" not in data or not isinstance(data["dossiers"], list):
                data["dossiers"] = []
            return data
        except Exception:
            backup = VAE_DATA_FILE + ".corrupt." + str(int(datetime.datetime.utcnow().timestamp()))
            try:
                os.replace(VAE_DATA_FILE, backup)
            except Exception:
                pass
            data = {"dossiers": []}
            _vae_save_all(data)
            return data

def _vae_save_all(data: Dict[str, Any]) -> None:
    _write_json_with_backups(VAE_DATA_FILE, data, _vae_lock)

def _vae_find_dossier(data: Dict[str, Any], dossier_id: str) -> Optional[Dict[str, Any]]:
    for d in data.get("dossiers", []):
        if d.get("id") == dossier_id:
            return d
    return None

def _vae_find_latest_for_trainee(trainee_id: str) -> Optional[Dict[str, Any]]:
    if not trainee_id:
        return None
    data = _vae_load_all()
    dossiers = [
        d for d in data.get("dossiers", [])
        if str((d.get("meta") or {}).get("trainee_id") or "") == str(trainee_id)
    ]
    if not dossiers:
        return None
    dossiers.sort(key=lambda d: d.get("updated_at") or d.get("created_at") or "", reverse=True)
    return dossiers[0]

def _pdf_escape(text: Any) -> str:
    s = str(text or "")
    return s.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")

def _build_simple_pdf(lines: List[str]) -> bytes:
    clean_lines = [line if isinstance(line, str) else str(line) for line in lines]
    if not clean_lines:
        clean_lines = [""]

    objects: List[bytes] = []
    objects.append(b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n")
    objects.append(b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n")
    objects.append(b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >> endobj\n")
    objects.append(b"4 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n")

    y = 805
    content_lines = ["BT", "/F1 9 Tf"]
    for raw in clean_lines:
        parts = [raw[i:i + 110] for i in range(0, len(raw), 110)] or [""]
        for part in parts:
            if y < 40:
                break
            content_lines.append(f"1 0 0 1 30 {y} Tm ({_pdf_escape(part)}) Tj")
            y -= 12
        if y < 40:
            break
    content_lines.append("ET")
    content = ("\n".join(content_lines) + "\n").encode("latin-1", errors="replace")
    objects.append(f"5 0 obj << /Length {len(content)} >> stream\n".encode("ascii") + content + b"endstream\nendobj\n")

    pdf = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for obj in objects:
        offsets.append(len(pdf))
        pdf.extend(obj)

    xref_pos = len(pdf)
    total = len(objects) + 1
    pdf.extend(f"xref\n0 {total}\n".encode("ascii"))
    pdf.extend(b"0000000000 65535 f \n")
    for off in offsets[1:]:
        pdf.extend(f"{off:010d} 00000 n \n".encode("ascii"))
    pdf.extend(f"trailer << /Size {total} /Root 1 0 R >>\nstartxref\n{xref_pos}\n%%EOF".encode("ascii"))
    return bytes(pdf)

def _vae_dossier_to_lines(dossier: Dict[str, Any]) -> List[str]:
    candidat = dossier.get("candidat", {})
    certification = dossier.get("certification", {})
    parcours = dossier.get("parcours_previsionnel", {})
    avis = dossier.get("avis_admin", {})
    engagement = dossier.get("engagement", {})
    objectifs = ", ".join(candidat.get("objectifs") or [])
    blocs_vises = ", ".join(certification.get("blocs_vises") or [])

    lines = [
        "DOSSIER DE FAISABILITE VAE DESP",
        f"ID dossier: {dossier.get('id')}",
        f"Statut dossier: {dossier.get('statut_dossier')}",
        "",
        "1. Nature de la demande",
        f"Nature: {dossier.get('nature_demande')}",
        "",
        "2. Informations generales sur le candidat",
        f"Nom de naissance: {candidat.get('nom_naissance')}",
        f"Nom d'usage: {candidat.get('nom_usage')}",
        f"Prenoms: {candidat.get('prenoms')}",
        f"Date de naissance: {candidat.get('date_naissance')}",
        f"Nationalite: {candidat.get('nationalite')}",
        f"Adresse: {candidat.get('adresse')}",
        f"Code postal: {candidat.get('code_postal')}",
        f"Ville: {candidat.get('ville')}",
        f"Telephone: {candidat.get('telephone')}",
        f"Email: {candidat.get('email')}",
        f"Objectifs: {objectifs}",
        "",
        "3. Certification professionnelle visee",
        f"Intitule: {certification.get('intitule')}",
        f"RNCP: {certification.get('rncp')}",
        f"Certificateur: {certification.get('certificateur')}",
        f"Parcours/mention: {certification.get('parcours_mention')}",
        f"Type de visee: {certification.get('vise')}",
        f"Blocs vises: {blocs_vises}",
        "",
        "4. Experiences professionnelles ou personnelles",
    ]

    for idx, exp in enumerate(dossier.get("experiences") or [], start=1):
        lines.extend([
            f"Experience {idx}: debut={exp.get('date_debut')} duree={exp.get('duree')}",
            f"Description: {exp.get('description')}",
        ])

    lines.extend(["", "5. Positionnement par competences"])
    blocs = dossier.get("blocs_competences", {})
    for i in range(1, 6):
        act = blocs.get(f"activite{i}", {})
        lines.append(f"Activite {i}:")
        for j in range(1, 5):
            comp = act.get(f"competence{j}", {}) if isinstance(act, dict) else {}
            lines.append(f"  - Competence {j}: intitule={comp.get('intitule')} statut={comp.get('statut')}")
        lines.append(f"  - Commentaires: {(act or {}).get('commentaires') if isinstance(act, dict) else ''}")

    lines.extend([
        "",
        "6. Parcours previsionnel",
        f"Accompagnement individuel (heures): {(parcours.get('accompagnement_individuel') or {}).get('heures')}",
        f"Accompagnement collectif (heures): {(parcours.get('accompagnement_collectif') or {}).get('heures')}",
        f"Formations prealables: {(parcours.get('formations_prealables') or {}).get('organisme')} / {(parcours.get('formations_prealables') or {}).get('intitule')}",
        f"Immersion: {(parcours.get('immersion') or {}).get('structure')}",
        f"Autres actions: {parcours.get('autres_actions')}",
        "",
        "7. Formulaire d'avis de faisabilite",
        f"Decision: {avis.get('decision')}",
        f"Motivation: {avis.get('motivation')}",
        f"Accompagnateur: {avis.get('nom_accompagnateur')} ({avis.get('email')}, {avis.get('telephone')})",
        "",
        "8. Accord pour l'analyse de la faisabilite",
        f"Souhaite accompagnement: {'Oui' if engagement.get('souhaite_accompagnement') else 'Non'}",
        f"Commentaires si avis défavorable: {engagement.get('commentaires_defavorable')}",
        f"Accord analyse: {'Oui' if engagement.get('accord_analyse') else 'Non'}",
        f"Signature: {engagement.get('nom_signature')} le {engagement.get('date_signature')} a {engagement.get('lieu_signature')}",
        f"Trace signature: {engagement.get('signature_trace')} ({engagement.get('signature_signed_at')})",
    ])
    return lines

def _merge_dict(base: Dict[str, Any], incoming: Dict[str, Any]) -> None:
    for k, v in incoming.items():
        if isinstance(v, dict) and isinstance(base.get(k), dict):
            _merge_dict(base[k], v)
        else:
            base[k] = v

def _validate_vae_for_submit(dossier: Dict[str, Any]) -> List[str]:
    errors: List[str] = []

    if (dossier.get("nature_demande") or "") != "initiale":
        errors.append("Nature de la demande : seule la valeur Initiale est autorisée")

    candidat = dossier.get("candidat", {})
    required_fields = {
        "nom_naissance": "Nom de naissance",
        "prenoms": "Prénom(s)",
        "date_naissance": "Date de naissance",
        "email": "Adresse email",
    }
    for key, label in required_fields.items():
        if not (str(candidat.get(key) or "").strip()):
            errors.append(f"1ère étape (Informations candidat) : {label} manquant")

    certification = dossier.get("certification", {})
    if (certification.get("vise") or "") != "complete":
        errors.append("La certification visée doit être la certification professionnelle dans son intégralité")

    experiences = dossier.get("experiences") if isinstance(dossier.get("experiences"), list) else []
    has_filled_experience = any(
        isinstance(exp, dict)
        and str(exp.get("date_debut") or "").strip()
        and str(exp.get("duree") or "").strip()
        and str(exp.get("description") or "").strip()
        for exp in experiences
    )
    if not has_filled_experience:
        errors.append("3ème étape (Expériences du candidat) : au moins une expérience doit être renseignée")

    blocs_competences = dossier.get("blocs_competences", {})
    for activity_idx in range(1, 6):
        activity = blocs_competences.get(f"activite{activity_idx}", {})
        for competence_idx in range(1, 5):
            competence = activity.get(f"competence{competence_idx}", {})
            if not str(competence.get("intitule") or "").strip():
                errors.append(
                    f"4ème étape (Tableau de positionnement) : intitulé manquant pour Activité {activity_idx}, compétence {competence_idx}"
                )
            if not str(competence.get("statut") or "").strip():
                errors.append(
                    f"4ème étape (Tableau de positionnement) : activité manquante pour Activité {activity_idx}, compétence {competence_idx}"
                )
        if not str((activity or {}).get("commentaires") or "").strip():
            errors.append(
                f"4ème étape (Tableau de positionnement) : commentaire manquant pour Activité {activity_idx}"
            )

    if (candidat.get("statut") or "") != "salarie_prive" and str(candidat.get("convention_collective") or "").strip():
        errors.append("1ère étape (Informations candidat) : la convention collective doit rester vide hors salarié du secteur privé")

    engagement = dossier.get("engagement", {})
    if not bool(engagement.get("accord_analyse")):
        errors.append("7ème étape (Accord d'analyse) : vous devez accepter l'analyse du dossier")
    if not str(engagement.get("lieu_signature") or "").strip():
        errors.append("7ème étape (Accord d'analyse) : lieu de signature manquant")
    if not str(engagement.get("date_signature") or "").strip():
        errors.append("7ème étape (Accord d'analyse) : date de signature manquante")
    if not str(engagement.get("nom_signature") or "").strip():
        errors.append("7ème étape (Accord d'analyse) : nom et prénom du signataire manquants")
    if not str(engagement.get("signature_trace") or "").strip() or not str(engagement.get("signature_signed_at") or "").strip():
        errors.append("7ème étape (Accord d'analyse) : signature électronique obligatoire")
    return errors

def _vae_create_and_redirect_for_trainee_token(trainee_token: str):
    trainee_token = (trainee_token or '').strip()
    linked_trainee_id = ''
    linked_session_id = ''

    if trainee_token:
        data_main = load_data()
        s, t = find_session_and_trainee_by_token(data_main, trainee_token)
        if s and t:
            training_type = _session_get(s, "training_type", "")
            if not required_docs_are_deposited(t, training_type):
                abort(403)
            linked_trainee_id = str(t.get('id') or '')
            linked_session_id = str(s.get('id') or '')

    data = _vae_load_all()
    dossier = _vae_default_dossier()
    dossier.setdefault('meta', {})['linkage_id'] = str(uuid.uuid4())
    if trainee_token:
        dossier.setdefault('meta', {})['trainee_token'] = trainee_token
    if linked_trainee_id:
        dossier['meta']['trainee_id'] = linked_trainee_id
        dossier['meta']['session_id'] = linked_session_id

    data.setdefault("dossiers", []).insert(0, dossier)
    _vae_save_all(data)
    return redirect(url_for('vae_wizard', token=dossier['id']))


@app.get('/vae/nouveau/<trainee_token>')
def vae_new_for_trainee(trainee_token: str):
    return _vae_create_and_redirect_for_trainee_token(trainee_token)


@app.get('/vae/nouveau')
def vae_new():
    trainee_token = (request.args.get('trainee_token') or '').strip()
    if not trainee_token:
        trainee_token = _vae_extract_trainee_token_from_referer(request.headers.get('Referer', ''))
    return _vae_create_and_redirect_for_trainee_token(trainee_token)

@app.get('/vae/<token>')
def vae_wizard(token: str):
    data = _vae_load_all()
    dossier = _vae_find_dossier(data, token)
    if not dossier:
        abort(404)
    if (dossier.get('statut_dossier') or '').strip().lower() == 'soumis':
        return redirect(url_for('vae_success', token=token))
    return render_template('vae_wizard.html', dossier=dossier, dossier_json=json.dumps(dossier, ensure_ascii=False))

@app.post('/api/vae/<dossier_id>/save')
@app.patch('/api/vae/<dossier_id>/save')
def api_vae_save(dossier_id: str):
    payload = request.get_json(silent=True) or {}
    if not isinstance(payload, dict):
        return jsonify({"ok": False, "error": "invalid_payload"}), 400

    data = _vae_load_all()
    dossier = _vae_find_dossier(data, dossier_id)
    if not dossier:
        return jsonify({"ok": False, "error": "not_found"}), 404
    if (dossier.get('statut_dossier') or '').strip().lower() == 'soumis':
        return jsonify({"ok": False, "error": "already_submitted"}), 403

    _merge_dict(dossier, payload)

    # Contraintes métier côté serveur
    dossier["nature_demande"] = "initiale"
    candidat = dossier.get("candidat") if isinstance(dossier.get("candidat"), dict) else {}
    if candidat.get("statut") != "salarie_prive":
        candidat["convention_collective"] = ""

    engagement = dossier.get("engagement") if isinstance(dossier.get("engagement"), dict) else {}
    prenoms = str(candidat.get("prenoms") or "").strip()
    nom_usage = str(candidat.get("nom_usage") or "").strip()
    nom_naissance = str(candidat.get("nom_naissance") or "").strip()
    full_name = " ".join([part for part in [prenoms, nom_usage or nom_naissance] if part]).strip()
    if full_name:
        engagement["nom_signature"] = full_name
    if not str(engagement.get("date_signature") or "").strip():
        engagement["date_signature"] = datetime.date.today().isoformat()
    dossier["engagement"] = engagement

    dossier["updated_at"] = _now_iso_utc()
    _vae_save_all(data)
    return jsonify({"ok": True, "id": dossier_id, "updated_at": dossier["updated_at"]})

@app.post('/api/vae/<dossier_id>/submit')
def api_vae_submit(dossier_id: str):
    data = _vae_load_all()
    dossier = _vae_find_dossier(data, dossier_id)
    if not dossier:
        return jsonify({"ok": False, "error": "not_found"}), 404

    errors = _validate_vae_for_submit(dossier)
    if errors:
        return jsonify({"ok": False, "errors": errors}), 400

    dossier["statut_dossier"] = "soumis"
    dossier["updated_at"] = _now_iso_utc()
    _vae_save_all(data)

    meta = dossier.get('meta') or {}
    trainee_id = str(meta.get('trainee_id') or '')
    session_id = str(meta.get('session_id') or '')

    data_main = load_data()
    s = find_session(data_main, session_id) if session_id else None
    t = None

    if s and trainee_id:
        trainees = _session_trainees_list(s)
        t = next((x for x in trainees if str(x.get('id') or '') == trainee_id), None)
    else:
        trainee_token = str(meta.get('trainee_token') or '').strip()
        if trainee_token:
            s, t = find_session_and_trainee_by_token(data_main, trainee_token)
            if s and t:
                trainee_id = str(t.get('id') or '')
                session_id = str(s.get('id') or '')
                dossier.setdefault('meta', {})['trainee_id'] = trainee_id
                dossier.setdefault('meta', {})['session_id'] = session_id
                dossier['updated_at'] = _now_iso_utc()
                _vae_save_all(data)

    if s and t:
        trainees = _session_trainees_list(s)
        current_trainee_id = str(t.get('id') or '')
        t = next((x for x in trainees if str(x.get('id') or '') == current_trainee_id), t)

        previous_status = vae_status_view(t.get('vae_status') or t.get('vae_status_label'))['key']
        view = vae_status_view('livret_1_analysis')
        t['vae_status'] = view['key']
        t['vae_status_label'] = view['label']
        if not isinstance(t.get('vae_action_dates'), dict):
            t['vae_action_dates'] = {}
        if not t['vae_action_dates'].get('livret_1_received'):
            t['vae_action_dates']['livret_1_received'] = fr_date(datetime.datetime.utcnow().strftime('%Y-%m-%d'))
        trainee_display_name = _format_trainee_name(t.get('first_name', ''), t.get('last_name', ''))
        add_admin_notification(
            data_main,
            f"VAE Livret 1️⃣ Déposé par {trainee_display_name}",
            meta={
                'type': 'vae_livret_1_submit',
                'session_id': s.get('id'),
                'trainee_id': t.get('id'),
                'vae_dossier_id': dossier_id,
            },
        )
        s['trainees'] = trainees
        s.pop('stagiaires', None)
        save_data(data_main)

        if previous_status != view['key']:
            _notify_vae_status_change(t, view['key'])
        else:
            print(
                f"[VAE][EMAIL] statut inchangé après soumission livret 1, pas d'email envoyé: "
                f"trainee_id={current_trainee_id!r} status={view['key']!r}"
            )
    else:
        print(
            f"[VAE][EMAIL] liaison session/stagiaire introuvable après soumission livret 1: "
            f"dossier_id={dossier_id!r} session_id={session_id!r} trainee_id={trainee_id!r}"
        )

    return jsonify({"ok": True, "redirect_url": url_for('vae_success', token=dossier_id)})

@app.get('/vae/<token>/succes')
def vae_success(token: str):
    data = _vae_load_all()
    dossier = _vae_find_dossier(data, token)
    if not dossier:
        abort(404)
    trainee_space_url = None
    meta = dossier.get('meta') or {}
    trainee_token = (meta.get('trainee_token') or '').strip()
    if trainee_token:
        trainee_space_url = url_for('public_trainee_space', token=trainee_token)
    return render_template('vae_success.html', dossier=dossier, trainee_space_url=trainee_space_url)

@app.get('/admin/vae')
@admin_login_required
def admin_vae_list():
    data = _vae_load_all()
    dossiers = sorted(data.get("dossiers", []), key=lambda d: d.get("updated_at", ""), reverse=True)
    return render_template('admin_vae_list.html', dossiers=dossiers)

@app.route('/admin/vae/<dossier_id>', methods=['GET', 'POST'])
@admin_login_required
def admin_vae_detail(dossier_id: str):
    data = _vae_load_all()
    dossier = _vae_find_dossier(data, dossier_id)
    if not dossier:
        abort(404)

    if request.method == 'POST':
        action = request.form.get('action', '').strip()
        if action == 'update_avis':
            avis = dossier.setdefault('avis_admin', {})
            avis['decision'] = request.form.get('decision', '').strip()
            avis['motivation'] = request.form.get('motivation', '').strip()
            avis['nom_accompagnateur'] = request.form.get('nom_accompagnateur', '').strip()
            avis['email'] = request.form.get('email', '').strip()
            avis['telephone'] = request.form.get('telephone', '').strip()
            avis['organisme'] = request.form.get('organisme', '').strip()
            avis['date'] = request.form.get('date', '').strip()
        elif action == 'mark_recevable':
            dossier['statut_dossier'] = 'recevable'
        elif action == 'mark_refuse':
            dossier['statut_dossier'] = 'refuse'
        dossier['updated_at'] = _now_iso_utc()
        _vae_save_all(data)
        return redirect(url_for('admin_vae_detail', dossier_id=dossier_id))

    return render_template('admin_vae_detail.html', dossier=dossier)

@app.get('/admin/vae/<dossier_id>/export')
@admin_login_required
def admin_vae_export(dossier_id: str):
    data = _vae_load_all()
    dossier = _vae_find_dossier(data, dossier_id)
    if not dossier:
        abort(404)

    statut_labels = {
        "brouillon": "Brouillon",
        "soumis": "Soumis",
        "recevable": "Recevable",
        "refuse": "Refusé",
    }
    decision_labels = {
        "faisable": "Faisable",
        "faisable_complements": "Faisable avec compléments",
        "non_faisable": "Non faisable",
    }

    return render_template(
        'admin_vae_export.html',
        dossier=dossier,
        statut_labels=statut_labels,
        decision_labels=decision_labels,
        annex_pages=1,
    )

@app.get("/admin/sessions/")
def admin_sessions_slash_redirect():
    return redirect(url_for("admin_sessions"), code=301)



if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)), debug=True)
